"""Domain services — every product-level invariant is enforced here.

Rationale for a service layer (rather than router-level CRUD): rules
like "confirmed revisions are immutable" and "baselines never
re-resolve to latest" are the product's core value. They must be
impossible to bypass from the HTTP edge.
"""
from __future__ import annotations

import csv as _csv
import difflib
import io as _io
import json as _json
import zipfile as _zipfile
from collections import Counter
from datetime import timedelta

from sqlalchemy import delete, or_, select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from . import models as m
from .models import RevisionStatus


class DomainError(Exception):
    """Rule violation the caller must see as 4xx/409."""

    def __init__(self, message: str, status_code: int = 409):
        super().__init__(message)
        self.status_code = status_code


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------


def get_or_404(db: Session, model, id_: str, what: str):
    obj = db.get(model, id_)
    if obj is None:
        raise DomainError(f"{what} not found: {id_}", status_code=404)
    return obj


def require_editable(revision: m.ArtifactRevision) -> None:
    if revision.status != RevisionStatus.DRAFT:
        raise DomainError(
            f"Revision {revision.id} is {revision.status.value} and therefore immutable. "
            "Clone it as a new DRAFT revision to continue editing.",
        )


# ---------------------------------------------------------------------------
# Project
# ---------------------------------------------------------------------------


def create_project(db: Session, *, key: str, name: str, description=None, actor="local-user"):
    project = m.Project(key=key, name=name, description=description, created_by=actor)
    db.add(project)
    db.commit()
    return project


# ---------------------------------------------------------------------------
# Artifact + Revision lifecycle
# ---------------------------------------------------------------------------


def create_artifact(
    db: Session,
    *,
    project_id: str,
    type: m.ArtifactType,
    title: str,
    snapshot: dict | None = None,
    actor="local-user",
) -> m.Artifact:
    project = get_or_404(db, m.Project, project_id, "Project")
    artifact = m.Artifact(project_id=project.id, type=type, title=title, created_by=actor)
    db.add(artifact)
    db.flush()
    revision = m.ArtifactRevision(
        artifact_id=artifact.id,
        revision_number=1,
        status=RevisionStatus.DRAFT,
        snapshot=snapshot or {},
        title=title,
        created_by=actor,
    )
    db.add(revision)
    db.flush()
    artifact.current_draft_revision_id = revision.id
    db.commit()
    return artifact


def next_revision_number(db: Session, artifact_id: str) -> int:
    result = db.execute(
        select(m.ArtifactRevision.revision_number)
        .where(m.ArtifactRevision.artifact_id == artifact_id)
        .order_by(m.ArtifactRevision.revision_number.desc())
        .limit(1)
    ).scalar_one_or_none()
    return (result or 0) + 1


def create_revision(
    db: Session,
    *,
    artifact_id: str,
    snapshot: dict | None = None,
    based_on_revision_id: str | None = None,
    actor="local-user",
    actor_id: str | None = None,
) -> m.ArtifactRevision:
    """Clone-as-new-revision. Confirmed history is never touched."""
    artifact = get_or_404(db, m.Artifact, artifact_id, "Artifact")
    if based_on_revision_id:
        parent = get_or_404(
            db, m.ArtifactRevision, based_on_revision_id, "Revision"
        )
        if parent.artifact_id != artifact_id:
            raise DomainError("based_on_revision belongs to a different artifact")
        base_snapshot = dict(parent.snapshot or {})
        base_title = parent.title
    else:
        latest = (
            db.execute(
                select(m.ArtifactRevision)
                .where(m.ArtifactRevision.artifact_id == artifact_id)
                .order_by(m.ArtifactRevision.revision_number.desc())
                .limit(1)
            )
            .scalar_one_or_none()
        )
        base_snapshot = dict(latest.snapshot or {}) if latest else {}
        base_title = latest.title if latest else artifact.title
        based_on_revision_id = latest.id if latest else None

    revision = m.ArtifactRevision(
        artifact_id=artifact_id,
        revision_number=next_revision_number(db, artifact_id),
        status=RevisionStatus.DRAFT,
        based_on_revision_id=based_on_revision_id,
        snapshot=snapshot if snapshot is not None else base_snapshot,
        title=base_title,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(revision)
    db.flush()
    # A new draft supersedes any previous DRAFT pointer but old drafts
    # themselves become SUPERSEDED only when a newer revision is confirmed.
    if revision.based_on and revision.based_on.status == RevisionStatus.DRAFT:
        revision.based_on.status = RevisionStatus.SUPERSEDED
    artifact.current_draft_revision_id = revision.id
    db.commit()
    return revision


def update_revision_snapshot(
    db: Session, revision_id: str, snapshot: dict, title: str | None = None
) -> m.ArtifactRevision:
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    require_editable(revision)
    revision.snapshot = snapshot
    if title is not None:
        revision.title = title
        revision.artifact.title = title
    db.commit()
    return revision


# Allowed transitions: DRAFT→IN_REVIEW→(CONFIRMED|DRAFT), CONFIRMED→SUPERSEDED,
# and ARCHIVED from CONFIRMED/SUPERSEDED. CONFIRMED is never edited.
_ALLOWED_TRANSITIONS = {
    RevisionStatus.DRAFT: {RevisionStatus.IN_REVIEW, RevisionStatus.ARCHIVED},
    RevisionStatus.IN_REVIEW: {RevisionStatus.DRAFT, RevisionStatus.CONFIRMED, RevisionStatus.ARCHIVED},
    RevisionStatus.CONFIRMED: {RevisionStatus.SUPERSEDED, RevisionStatus.ARCHIVED},
    RevisionStatus.SUPERSEDED: {RevisionStatus.ARCHIVED},
    RevisionStatus.ARCHIVED: set(),
}


def transition_revision(db: Session, revision_id: str, to_status: RevisionStatus) -> m.ArtifactRevision:
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    if to_status not in _ALLOWED_TRANSITIONS[revision.status]:
        raise DomainError(
            f"Illegal transition {revision.status.value} → {to_status.value}"
        )
    revision.status = to_status
    db.commit()
    return revision


def submit_for_review(db: Session, revision_id: str) -> m.ArtifactRevision:
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    if revision.status != RevisionStatus.DRAFT:
        raise DomainError("Only DRAFT revisions can be submitted for review")
    revision.status = RevisionStatus.IN_REVIEW
    db.commit()
    return revision


def confirm_revision(
    db: Session,
    revision_id: str,
    *,
    actor="local-user",
    comment: str | None = None,
    evidence: dict | None = None,
    supersede_confirmed: bool = True,
    actor_id: str | None = None,
) -> tuple[m.ArtifactRevision, m.Confirmation]:
    """Confirm = freeze. Atomic: technical design is snapshotted into the
    revision in the same transaction; any failure rolls everything back so
    a half-confirmed state is impossible.

    If an older CONFIRMED revision of the same artifact exists, it becomes
    SUPERSEDED (still readable, still bound in old baselines).
    """
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    if revision.status not in (RevisionStatus.IN_REVIEW, RevisionStatus.DRAFT):
        raise DomainError(
            f"Cannot confirm a revision in status {revision.status.value}"
        )
    try:
        # Technical-design artifacts freeze their bound designs at confirm time.
        if revision.artifact.type in (
            m.ArtifactType.DR,
            m.ArtifactType.DATABASE_SCHEMA,
            m.ArtifactType.ARCHITECTURE,
        ):
            snapshot = dict(revision.snapshot or {})
            snapshot["technical_design"] = snapshot_technical_design(
                db, revision.artifact.project_id
            )
            revision.snapshot = snapshot

        revision.status = RevisionStatus.CONFIRMED
        revision.confirmed_at = m.utcnow()
        revision.confirmed_by = actor

        if supersede_confirmed:
            siblings = db.execute(
                select(m.ArtifactRevision).where(
                    m.ArtifactRevision.artifact_id == revision.artifact_id,
                    m.ArtifactRevision.id != revision.id,
                    m.ArtifactRevision.status == RevisionStatus.CONFIRMED,
                )
            ).scalars().all()
            for sib in siblings:
                sib.status = RevisionStatus.SUPERSEDED

        confirmation = m.Confirmation(
            project_id=revision.artifact.project_id,
            artifact_revision_id=revision.id,
            confirmed_by=actor,
            comment=comment,
            evidence=evidence,
            actor_id=actor_id,
        )
        db.add(confirmation)
        db.commit()
    except Exception:
        db.rollback()
        raise
    return revision, confirmation


# ---------------------------------------------------------------------------
# Technical design snapshot (auto-frozen into DR revisions at confirmation)
# ---------------------------------------------------------------------------


def flow_snapshot(db: Session, flow_id: str) -> dict:
    flow = get_or_404(db, m.ProcessFlow, flow_id, "ProcessFlow")
    steps = db.execute(
        select(m.ProcessStep)
        .where(m.ProcessStep.flow_id == flow_id)
        .order_by(m.ProcessStep.position)
    ).scalars().all()
    transitions = db.execute(
        select(m.ProcessTransition).where(m.ProcessTransition.flow_id == flow_id)
    ).scalars().all()
    return {
        "name": flow.name,
        "description": flow.description,
        "steps": {s.semantic_id: {"name": s.name, "step_type": s.step_type, "position": s.position} for s in steps},
        "transitions": {
            t.semantic_id: {"from": t.from_step_semantic_id, "to": t.to_step_semantic_id,
                            "label": t.label, "condition": t.condition}
            for t in transitions
        },
    }


def api_endpoint_snapshot(api: m.APIEndpoint) -> dict:
    return {
        "method": api.method,
        "path": api.path,
        "summary": api.summary,
        "description": api.description,
        "authentication": api.authentication,
        "parameters": [
            {"name": p.name, "location": p.location, "data_type": p.data_type,
             "required": p.required, "description": p.description}
            for p in api.parameters
        ],
        "request_fields": [
            {"name": f.name, "data_type": f.data_type, "required": f.required, "description": f.description}
            for f in api.request_fields
        ],
        "response_fields": [
            {"status_code": f.status_code, "name": f.name, "data_type": f.data_type, "description": f.description}
            for f in api.response_fields
        ],
        "error_responses": [
            {"status_code": e.status_code, "message": e.message, "description": e.description}
            for e in api.error_responses
        ],
        "request_spec": api.request_spec,
        "response_spec": api.response_spec,
    }


def snapshot_technical_design(db: Session, project_id: str) -> dict:
    """Freeze the exact current structured designs for a project.

    The result is embedded into a confirmed DR revision snapshot, so a
    historical export can always reproduce the design as it was then.
    """
    designs: dict = {}

    schemas = db.execute(
        select(m.DatabaseSchema).where(m.DatabaseSchema.project_id == project_id)
    ).scalars().all()
    designs["db_schemas"] = {s.semantic_id: db_design_snapshot(db, s.id) for s in schemas}

    flows = db.execute(
        select(m.ProcessFlow).where(m.ProcessFlow.project_id == project_id)
    ).scalars().all()
    designs["flows"] = {f.semantic_id: flow_snapshot(db, f.id) for f in flows}

    apis = db.execute(
        select(m.APIEndpoint).where(m.APIEndpoint.project_id == project_id)
    ).scalars().all()
    designs["api_endpoints"] = {a.semantic_id: api_endpoint_snapshot(a) for a in apis}

    diagrams = db.execute(
        select(m.ArchitectureDiagram).where(m.ArchitectureDiagram.project_id == project_id)
    ).scalars().all()
    designs["architecture"] = {}
    for d in diagrams:
        nodes = db.execute(
            select(m.ArchitectureNode).where(m.ArchitectureNode.diagram_id == d.id)
        ).scalars().all()
        edges = db.execute(
            select(m.ArchitectureEdge).where(m.ArchitectureEdge.diagram_id == d.id)
        ).scalars().all()
        designs["architecture"][d.semantic_id] = {
            "name": d.name,
            "nodes": {n.semantic_id: {"name": n.name, "node_type": n.node_type, "technology": n.technology, "environment": n.environment} for n in nodes},
            "edges": {e.semantic_id: {"from": e.from_node_semantic_id, "to": e.to_node_semantic_id, "label": e.label} for e in edges},
        }

    return designs


# ---------------------------------------------------------------------------
# Baseline
# ---------------------------------------------------------------------------


def create_baseline(
    db: Session,
    *,
    project_id: str,
    name: str,
    description: str | None = None,
    artifact_revision_ids: list[str],
    actor="local-user",
    actor_id: str | None = None,
) -> m.Baseline:
    """Freeze the exact artifact→revision pairs given at creation time.

    The stored binding rows are never re-resolved afterwards — that is
    the whole point. A later v8 of a child artifact does not change a
    baseline that bound v7.
    """
    get_or_404(db, m.Project, project_id, "Project")
    if not artifact_revision_ids:
        raise DomainError("A baseline must bind at least one revision")

    seen: dict[str, m.ArtifactRevision] = {}
    for rid in artifact_revision_ids:
        rev = get_or_404(db, m.ArtifactRevision, rid, "Revision")
        if rev.status != RevisionStatus.CONFIRMED:
            raise DomainError(
                f"Revision {rid} is {rev.status.value}; only CONFIRMED revisions "
                "may be frozen into a baseline"
            )
        if rev.artifact.project_id != project_id:
            raise DomainError(f"Revision {rid} belongs to another project")
        if rev.artifact_id in seen:
            raise DomainError("Each artifact may appear at most once per baseline")
        seen[rev.artifact_id] = rev

    baseline = m.Baseline(
        project_id=project_id, name=name, description=description, created_by=actor, actor_id=actor_id
    )
    db.add(baseline)
    db.flush()
    for artifact_id, rev in seen.items():
        semantic_object = db.execute(
            select(m.SemanticObject).where(
                m.SemanticObject.project_id == project_id,
                m.SemanticObject.entity_ref == artifact_id,
            )
        ).scalar_one_or_none()
        db.add(
            m.BaselineBinding(
                baseline_id=baseline.id,
                artifact_id=artifact_id,
                artifact_revision_id=rev.id,
                semantic_object_id=semantic_object.semantic_id if semantic_object else None,
                semantic_object_type=(
                    semantic_object.object_type.value if semantic_object else None
                ),
            )
        )
    db.commit()
    return baseline


def resolve_baseline(db: Session, baseline_id: str) -> m.Baseline:
    baseline = get_or_404(db, m.Baseline, baseline_id, "Baseline")
    db.expire(baseline, ["bindings"])  # always read the frozen rows
    return baseline


# ---------------------------------------------------------------------------
# SemanticObject + TraceLink
# ---------------------------------------------------------------------------


def ensure_semantic_object(
    db: Session,
    *,
    project_id: str,
    semantic_id: str,
    object_type: m.SemanticObjectType,
    display_name: str,
    entity_ref: str | None = None,
    metadata: dict | None = None,
) -> m.SemanticObject:
    obj = db.execute(
        select(m.SemanticObject).where(
            m.SemanticObject.project_id == project_id,
            m.SemanticObject.semantic_id == semantic_id,
        )
    ).scalar_one_or_none()
    if obj:
        obj.display_name = display_name  # display names may change
        if entity_ref:
            obj.entity_ref = entity_ref
        if metadata:
            obj.metadata_json = metadata
        db.commit()
        return obj
    obj = m.SemanticObject(
        project_id=project_id,
        semantic_id=semantic_id,
        object_type=object_type,
        display_name=display_name,
        entity_ref=entity_ref,
        metadata_json=metadata,
    )
    db.add(obj)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        obj = db.execute(
            select(m.SemanticObject).where(
                m.SemanticObject.project_id == project_id,
                m.SemanticObject.semantic_id == semantic_id,
            )
        ).scalar_one()
    return obj


def create_trace_link(
    db: Session,
    *,
    project_id: str,
    source_semantic_id: str,
    target_semantic_id: str,
    relation_type: m.TraceRelationType,
    revision_context: str | None = None,
    actor="local-user",
) -> m.TraceLink:
    for sid in (source_semantic_id, target_semantic_id):
        exists = db.execute(
            select(m.SemanticObject.id).where(
                m.SemanticObject.project_id == project_id,
                m.SemanticObject.semantic_id == sid,
            )
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(
                f"Unknown semantic object '{sid}'. Traces may only connect "
                "registered semantic objects — never pixel positions or titles.",
                status_code=422,
            )
    link = m.TraceLink(
        project_id=project_id,
        source_semantic_id=source_semantic_id,
        target_semantic_id=target_semantic_id,
        relation_type=relation_type,
        revision_context=revision_context,
        created_by=actor,
    )
    db.add(link)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        raise DomainError("Trace link already exists")
    return link


def trace_graph(db: Session, project_id: str) -> dict:
    """Nodes + edges for the traceability explorer. Only stored links are shown."""
    nodes = db.execute(
        select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
    ).scalars().all()
    edges = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
    ).scalars().all()
    return {
        "nodes": [
            {
                "semantic_id": n.semantic_id,
                "object_type": n.object_type.value,
                "display_name": n.display_name,
            }
            for n in nodes
        ],
        "edges": [
            {
                "source": e.source_semantic_id,
                "target": e.target_semantic_id,
                "relation": e.relation_type.value,
                "revision_context": e.revision_context,
            }
            for e in edges
        ],
    }


def impact_of(db: Session, project_id: str, semantic_id: str) -> dict:
    """1-hop upstream/downstream impact using trace links only."""
    outgoing = db.execute(
        select(m.TraceLink).where(
            m.TraceLink.project_id == project_id,
            m.TraceLink.source_semantic_id == semantic_id,
        )
    ).scalars().all()
    incoming = db.execute(
        select(m.TraceLink).where(
            m.TraceLink.project_id == project_id,
            m.TraceLink.target_semantic_id == semantic_id,
        )
    ).scalars().all()
    return {
        "semantic_id": semantic_id,
        "downstream": [
            {"semantic_id": l.target_semantic_id, "relation": l.relation_type.value}
            for l in outgoing
        ],
        "upstream": [
            {"semantic_id": l.source_semantic_id, "relation": l.relation_type.value}
            for l in incoming
        ],
    }


# ---------------------------------------------------------------------------
# Annotation
# ---------------------------------------------------------------------------


def create_annotation(
    db: Session,
    *,
    project_id: str,
    anchor_object_type: str,
    anchor_semantic_id: str,
    content: str,
    type: m.AnnotationType = m.AnnotationType.COMMENT,
    artifact_revision_id: str | None = None,
    canvas_x: float | None = None,
    canvas_y: float | None = None,
    drawing_payload: dict | None = None,
    thread_id: str | None = None,
    actor="local-user",
    actor_id: str | None = None,
) -> m.Annotation:
    anchored = db.execute(
        select(m.SemanticObject.id).where(
            m.SemanticObject.project_id == project_id,
            m.SemanticObject.semantic_id == anchor_semantic_id,
        )
    ).scalar_one_or_none()
    if not anchored:
        raise DomainError(
            f"Cannot anchor annotation to unknown semantic object '{anchor_semantic_id}'. "
            "Coordinates are optional placement data, never the anchor.",
            status_code=422,
        )
    if thread_id is not None:
        thread = get_or_404(db, m.CommentThread, thread_id, "CommentThread")
        if thread.project_id != project_id:
            raise DomainError("Thread belongs to a different project")
    annotation = m.Annotation(
        project_id=project_id,
        artifact_revision_id=artifact_revision_id,
        anchor_object_type=anchor_object_type,
        anchor_semantic_id=anchor_semantic_id,
        canvas_x=canvas_x,
        canvas_y=canvas_y,
        type=type,
        content=content,
        drawing_payload=drawing_payload,
        thread_id=thread_id,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(annotation)
    db.commit()
    return annotation


def set_annotation_status(
    db: Session, annotation_id: str, status: m.AnnotationStatus
) -> m.Annotation:
    annotation = get_or_404(db, m.Annotation, annotation_id, "Annotation")
    valid = {s.value for s in m.AnnotationStatus}
    if status.value not in valid:
        raise DomainError("Invalid annotation status")
    annotation.status = status
    db.commit()
    return annotation


# ---------------------------------------------------------------------------
# Requirements
# ---------------------------------------------------------------------------


def next_requirement_code(db: Session, project_id: str) -> str:
    count = (
        db.execute(
            select(m.Requirement.id).where(m.Requirement.project_id == project_id)
        )
        .scalars()
        .all()
    )
    return f"REQ-{len(count) + 1:04d}"


def create_requirement(
    db: Session,
    *,
    project_id: str,
    title: str,
    description=None,
    source_type=None,
    source_reference=None,
    priority=None,
    actor="local-user",
) -> m.Requirement:
    code = next_requirement_code(db, project_id)
    requirement = m.Requirement(
        project_id=project_id,
        code=code,
        title=title,
        description=description,
        source_type=source_type,
        source_reference=source_reference,
        priority=priority,
        created_by=actor,
    )
    db.add(requirement)
    db.flush()
    ensure_semantic_object(
        db,
        project_id=project_id,
        semantic_id=code,
        object_type=m.SemanticObjectType.REQUIREMENT,
        display_name=title,
        entity_ref=requirement.id,
    )
    db.commit()
    return requirement


# ---------------------------------------------------------------------------
# ChangeRequest
# ---------------------------------------------------------------------------


def next_cr_code(db: Session, project_id: str) -> str:
    count = (
        db.execute(
            select(m.ChangeRequest.id).where(m.ChangeRequest.project_id == project_id)
        )
        .scalars()
        .all()
    )
    return f"CR-{len(count) + 1:04d}"


def create_change_request(
    db: Session,
    *,
    project_id: str,
    requested_change: str,
    affected_semantic_ids: list[str],
    reason=None,
    requested_by="local-user",
    target_release=None,
    schedule_impact=None,
    commercial_impact=None,
    actor="local-user",
    actor_id: str | None = None,
) -> m.ChangeRequest:
    for sid in affected_semantic_ids:
        exists = db.execute(
            select(m.SemanticObject.id).where(
                m.SemanticObject.project_id == project_id,
                m.SemanticObject.semantic_id == sid,
            )
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(
                f"Change request references unknown semantic object '{sid}'",
                status_code=422,
            )
    cr = m.ChangeRequest(
        project_id=project_id,
        code=next_cr_code(db, project_id),
        requested_by=requested_by,
        reason=reason,
        requested_change=requested_change,
        target_release=target_release,
        schedule_impact=schedule_impact,
        commercial_impact=commercial_impact,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(cr)
    db.flush()
    for sid in affected_semantic_ids:
        db.add(m.ChangeRequestLink(change_request_id=cr.id, semantic_id=sid))
    db.commit()
    return cr


def implement_change_request(
    db: Session,
    change_request_id: str,
    *,
    artifact_revision_map: dict[str, str] | None = None,
    actor="local-user",
) -> dict:
    """Mark a CR IMPLEMENTED by pointing it at the new revision(s) it spawned.

    The CR itself never mutates old confirmed baselines — it only
    references the freshly created draft/confirmed revisions.
    """
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    if cr.status == m.ChangeRequestStatus.IMPLEMENTED:
        raise DomainError("Change request already implemented")
    spawned: list[dict] = []
    for artifact_id, snapshot in (artifact_revision_map or {}).items():
        rev = create_revision(
            db, artifact_id=artifact_id, snapshot=snapshot, actor=actor
        )
        ensure_semantic_object(
            db,
            project_id=cr.project_id,
            semantic_id=f"rev_{rev.id}",
            object_type=m.SemanticObjectType.DOCUMENT_SECTION,
            display_name=f"{cr.code} revision {rev.revision_number}",
            entity_ref=rev.id,
        )
        db.add(
            m.TraceLink(
                project_id=cr.project_id,
                source_semantic_id=cr.code,
                target_semantic_id=f"rev_{rev.id}",
                relation_type=m.TraceRelationType.GENERATED_FROM,
                created_by=actor,
            )
        )
        spawned.append({"artifact_id": artifact_id, "revision_id": rev.id, "revision_number": rev.revision_number})
    cr.status = m.ChangeRequestStatus.IMPLEMENTED
    db.commit()
    return {"change_request": cr, "spawned_revisions": spawned}


def change_request_detail(db: Session, change_request_id: str) -> dict:
    """Full CR view: affected objects, deterministic impact, spawned revisions."""
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    links = [l.semantic_id for l in cr.links]

    affected = []
    for sid in links:
        so = db.execute(
            select(m.SemanticObject).where(
                m.SemanticObject.project_id == cr.project_id,
                m.SemanticObject.semantic_id == sid,
            )
        ).scalar_one_or_none()
        affected.append({
            "semantic_id": sid,
            "object_type": so.object_type.value if so else None,
            "display_name": so.display_name if so else None,
        })

    impact = {sid: impact_of(db, cr.project_id, sid) for sid in links}

    spawned = []
    traces = db.execute(
        select(m.TraceLink).where(
            m.TraceLink.project_id == cr.project_id,
            m.TraceLink.source_semantic_id == cr.code,
            m.TraceLink.relation_type == m.TraceRelationType.GENERATED_FROM,
        )
    ).scalars().all()
    for t in traces:
        so = db.execute(
            select(m.SemanticObject).where(
                m.SemanticObject.project_id == cr.project_id,
                m.SemanticObject.semantic_id == t.target_semantic_id,
            )
        ).scalar_one_or_none()
        if so and so.entity_ref:
            rev = db.get(m.ArtifactRevision, so.entity_ref)
            if rev:
                spawned.append({
                    "revision_id": rev.id,
                    "revision_number": rev.revision_number,
                    "artifact_id": rev.artifact_id,
                    "artifact_title": rev.artifact.title,
                    "status": rev.status.value,
                    "based_on_revision_id": rev.based_on_revision_id,
                })

    return {
        "id": cr.id,
        "code": cr.code,
        "project_id": cr.project_id,
        "requested_by": cr.requested_by,
        "reason": cr.reason,
        "requested_change": cr.requested_change,
        "status": cr.status.value,
        "target_release": cr.target_release,
        "schedule_impact": cr.schedule_impact,
        "commercial_impact": cr.commercial_impact,
        "created_at": cr.created_at.isoformat(),
        "created_by": cr.created_by,
        "affected": affected,
        "impact": impact,
        "spawned_revisions": spawned,
    }


# ---------------------------------------------------------------------------
# Database design (structured model; diagram is a view over it)
# ---------------------------------------------------------------------------


def create_schema(
    db: Session, *, project_id: str, name: str, semantic_id: str, description=None, actor="local-user"
) -> m.DatabaseSchema:
    schema = m.DatabaseSchema(
        project_id=project_id, name=name, semantic_id=semantic_id, description=description,
        created_by=actor,
    )
    db.add(schema)
    db.flush()
    ensure_semantic_object(
        db,
        project_id=project_id,
        semantic_id=semantic_id,
        object_type=m.SemanticObjectType.DB_SCHEMA,
        display_name=name,
        entity_ref=schema.id,
    )
    db.commit()
    return schema


def create_table(
    db: Session, *, schema_id: str, name: str, semantic_id: str | None = None, description=None
) -> m.DatabaseTable:
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    semantic_id = semantic_id or f"tbl_{name}"
    table = m.DatabaseTable(
        schema_id=schema_id, semantic_id=semantic_id, name=name, description=description
    )
    db.add(table)
    db.flush()
    ensure_semantic_object(
        db,
        project_id=schema.project_id,
        semantic_id=semantic_id,
        object_type=m.SemanticObjectType.DB_TABLE,
        display_name=name,
        entity_ref=table.id,
    )
    db.commit()
    return table


def create_field(
    db: Session,
    *,
    table_id: str,
    name: str,
    data_type: str,
    semantic_id: str | None = None,
    length=None,
    nullable=False,
    default=None,
    primary_key=False,
    foreign_key=False,
    reference=None,
    description=None,
    remark=None,
) -> m.DatabaseField:
    table = get_or_404(db, m.DatabaseTable, table_id, "Table")
    semantic_id = semantic_id or f"fld_{table.name}_{name}"
    position = (
        db.execute(
            select(m.DatabaseField.id).where(m.DatabaseField.table_id == table_id)
        )
        .scalars()
        .all()
    )
    field = m.DatabaseField(
        table_id=table_id,
        semantic_id=semantic_id,
        name=name,
        data_type=data_type,
        length=length,
        nullable=nullable,
        default=default,
        primary_key=primary_key,
        foreign_key=foreign_key,
        reference=reference,
        description=description,
        remark=remark,
        position=len(position),
    )
    db.add(field)
    db.flush()
    ensure_semantic_object(
        db,
        project_id=table.schema.project_id,
        semantic_id=semantic_id,
        object_type=m.SemanticObjectType.DB_FIELD,
        display_name=f"{table.name}.{name}",
        entity_ref=field.id,
    )
    db.commit()
    return field


def create_relation(
    db: Session,
    *,
    schema_id: str,
    from_field_semantic_id: str,
    to_field_semantic_id: str,
    relation_type="MANY_TO_ONE",
) -> m.DatabaseRelation:
    for sid in (from_field_semantic_id, to_field_semantic_id):
        exists = db.execute(
            select(m.DatabaseField.id).where(m.DatabaseField.semantic_id == sid)
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(f"Unknown field semantic id '{sid}'", status_code=422)
    rel = m.DatabaseRelation(
        schema_id=schema_id,
        semantic_id=f"rel_{from_field_semantic_id}__{to_field_semantic_id}",
        from_field_semantic_id=from_field_semantic_id,
        to_field_semantic_id=to_field_semantic_id,
        relation_type=relation_type,
    )
    db.add(rel)
    db.flush()
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    ensure_semantic_object(
        db,
        project_id=schema.project_id,
        semantic_id=rel.semantic_id,
        object_type=m.SemanticObjectType.DB_RELATION,
        display_name=f"{from_field_semantic_id} → {to_field_semantic_id}",
        entity_ref=rel.id,
    )
    db.commit()
    return rel


def data_dictionary(db: Session, schema_id: str) -> list[dict]:
    """A pure view over the structured model — no separate truth stored."""
    get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    tables = db.execute(
        select(m.DatabaseTable).where(m.DatabaseTable.schema_id == schema_id)
    ).scalars().all()
    table_sid = {t.id: t.semantic_id for t in tables}
    table_name = {t.id: t.name for t in tables}
    fields = db.execute(
        select(m.DatabaseField).where(m.DatabaseField.table_id.in_(table_sid))
    ).scalars().all()
    result = []
    for f in fields:
        result.append(
            {
                "table": table_name[f.table_id],
                "table_semantic_id": table_sid[f.table_id],
                "field": f.name,
                "field_semantic_id": f.semantic_id,
                "data_type": f.data_type,
                "length": f.length,
                "nullable": f.nullable,
                "default": f.default,
                "primary_key": f.primary_key,
                "foreign_key": f.foreign_key,
                "reference": f.reference,
                "description": f.description,
                "remark": f.remark,
            }
        )
    return result


# ---------------------------------------------------------------------------
# Document workspace (rich UR/DR content, section semantic identity)
# ---------------------------------------------------------------------------

_DOC_BLOCK_KINDS = {"heading", "paragraph", "bullet_list", "numbered_list", "table", "code"}


def _doc_text(content: dict) -> str:
    """Plain text derived from a ProseMirror/Tiptap doc JSON (deterministic)."""
    if not isinstance(content, dict):
        return ""
    parts: list[str] = []

    def walk(node) -> None:
        if not isinstance(node, dict):
            return
        ntype = node.get("type")
        if ntype == "text":
            parts.append(node.get("text") or "")
            return
        for child in node.get("content") or []:
            walk(child)
        if ntype in ("paragraph", "heading", "codeBlock", "blockquote", "listItem", "tableRow"):
            parts.append("\n")

    walk(content)
    return "\n".join(line.rstrip() for line in "".join(parts).splitlines())


def _blocks_to_doc(blocks: list[dict]) -> dict:
    """Convert the legacy P1 block model into a ProseMirror/Tiptap doc JSON."""
    content: list[dict] = []
    for blk in blocks:
        kind = blk.get("kind") or "paragraph"
        if kind == "heading":
            content.append({"type": "heading", "attrs": {"level": blk.get("level") or 2},
                            "content": [{"type": "text", "text": blk.get("text") or ""}]})
        elif kind == "bullet_list":
            content.append({"type": "bulletList", "content": [
                {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": item}]}]}
                for item in blk.get("items", [])
            ]})
        elif kind == "numbered_list":
            content.append({"type": "orderedList", "content": [
                {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": item}]}]}
                for item in blk.get("items", [])
            ]})
        elif kind == "table":
            rows = []
            if blk.get("header"):
                rows.append({"type": "tableRow", "content": [
                    {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": h}]}]}
                    for h in blk["header"]
                ]})
            for row in blk.get("rows", []):
                rows.append({"type": "tableRow", "content": [
                    {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": c}]}]}
                    for c in row
                ]})
            content.append({"type": "table", "content": rows})
        elif kind == "code":
            content.append({"type": "codeBlock", "content": [{"type": "text", "text": blk.get("text") or ""}]})
        else:
            content.append({"type": "paragraph", "content": [{"type": "text", "text": blk.get("text") or ""}]})
    return {"type": "doc", "content": content}


def _normalise_sections(artifact_id: str, sections: list[dict]) -> list[dict]:
    """Assign a stable section id where missing and normalise the content model.

    Section identity lives inside the snapshot and is preserved by the
    editor across edits. Legacy P1 `blocks` are migrated to a structured
    ProseMirror doc JSON; `plain_text` is always derived, never the truth.
    """
    out: list[dict] = []
    used: set[str] = set()
    counter = 0
    for sec in sections:
        sec = dict(sec)
        sid = sec.get("id") or f"docsec_{artifact_id}_{counter}"
        while sid in used:
            counter += 1
            sid = f"docsec_{artifact_id}_{counter}"
        used.add(sid)
        sec["id"] = sid
        sec.setdefault("heading", f"Section {len(out) + 1}")
        if sec.get("content") is None and sec.get("blocks"):
            sec["content"] = _blocks_to_doc(sec.get("blocks", []))
        sec.pop("blocks", None)
        if not sec.get("content"):
            sec["content"] = {"type": "doc", "content": [{"type": "paragraph"}]}
        if not sec.get("plain_text"):
            sec["plain_text"] = f"{sec.get('heading') or ''}\n{_doc_text(sec['content'])}"
        out.append(sec)
        counter += 1
    return out


def save_document(
    db: Session,
    *,
    revision_id: str,
    sections: list[dict],
    title: str | None = None,
    actor: str = "local-user",
) -> m.ArtifactRevision:
    """Persist rich document content for a DRAFT revision.

    Each section is registered as a DOCUMENT_SECTION semantic object so
    comments and traces can bind to it — never to a DOM position.
    Confirmed revisions are immutable and are rejected here.
    """
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    require_editable(revision)
    artifact = revision.artifact
    sections = _normalise_sections(artifact.id, sections)
    for sec in sections:
        ensure_semantic_object(
            db,
            project_id=artifact.project_id,
            semantic_id=sec["id"],
            object_type=m.SemanticObjectType.DOCUMENT_SECTION,
            display_name=sec.get("heading") or "Untitled section",
            entity_ref=artifact.id,
        )
    snapshot = dict(revision.snapshot or {})
    snapshot["sections"] = sections
    revision.snapshot = snapshot
    if title is not None:
        revision.title = title
        artifact.title = title
    db.commit()
    return revision


def get_document(db: Session, revision_id: str) -> dict:
    """Return the revision plus its structured sections."""
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    snapshot = revision.snapshot or {}
    sections = snapshot.get("sections")
    if not isinstance(sections, list):
        # Legacy P0 snapshots (e.g. {"sections": [{"id", "note"}]}) become
        # minimal paragraphs so nothing is lost on first edit.
        legacy = snapshot.get("sections") or snapshot.get("content") or []
        sections = [
            {
                "id": f"docsec_{revision.artifact_id}_{i}",
                "heading": sec.get("id") or sec.get("title") or f"Section {i + 1}",
                "blocks": [
                    {"kind": "paragraph", "text": sec.get("note") or sec.get("text") or ""}
                ],
            }
            for i, sec in enumerate(legacy)
        ] if isinstance(legacy, list) else []
    sections = _normalise_sections(revision.artifact_id, sections)
    return {
        "revision_id": revision.id,
        "artifact_id": revision.artifact_id,
        "revision_number": revision.revision_number,
        "status": revision.status.value,
        "title": revision.title,
        "artifact_type": revision.artifact.type.value,
        "based_on_revision_id": revision.based_on_revision_id,
        "created_by": revision.created_by,
        "confirmed_by": revision.confirmed_by,
        "confirmed_at": revision.confirmed_at.isoformat() if revision.confirmed_at else None,
        "editable": revision.editable,
        "sections": sections,
    }


# ---------------------------------------------------------------------------
# Review workflow (threads, summaries, activity timeline)
# ---------------------------------------------------------------------------


def create_thread(
    db: Session,
    *,
    project_id: str,
    title: str | None = None,
    actor: str = "local-user",
) -> m.CommentThread:
    get_or_404(db, m.Project, project_id, "Project")
    thread = m.CommentThread(project_id=project_id, title=title)
    db.add(thread)
    db.commit()
    return thread


def list_threads(db: Session, project_id: str) -> list[dict]:
    threads = db.execute(
        select(m.CommentThread)
        .where(m.CommentThread.project_id == project_id)
        .order_by(m.CommentThread.created_at.desc())
    ).scalars().all()
    out = []
    for t in threads:
        anns = db.execute(
            select(m.Annotation)
            .where(m.Annotation.thread_id == t.id)
            .order_by(m.Annotation.created_at)
        ).scalars().all()
        out.append({
            "id": t.id,
            "title": t.title,
            "resolved": t.resolved,
            "open_count": sum(1 for a in anns if a.status != m.AnnotationStatus.RESOLVED),
            "total": len(anns),
            "created_at": t.created_at.isoformat(),
            "annotations": [
                {
                    "id": a.id,
                    "anchor_semantic_id": a.anchor_semantic_id,
                    "type": a.type.value,
                    "status": a.status.value,
                    "content": a.content,
                    "created_by": a.created_by,
                    "created_at": a.created_at.isoformat(),
                }
                for a in anns
            ],
        })
    return out


def annotations_summary(db: Session, project_id: str) -> dict:
    anns = db.execute(
        select(m.Annotation).where(m.Annotation.project_id == project_id)
    ).scalars().all()
    by_status = Counter(a.status.value for a in anns)
    by_type = Counter(a.type.value for a in anns)
    by_anchor = Counter(a.anchor_semantic_id for a in anns)
    return {
        "total": len(anns),
        "open": sum(n for s, n in by_status.items() if s != "RESOLVED"),
        "resolved": by_status.get("RESOLVED", 0),
        "by_status": dict(by_status),
        "by_type": dict(by_type),
        "by_anchor": dict(by_anchor),
    }


def _iso(dt) -> str | None:
    return dt.isoformat() if dt else None


def timeline(db: Session, project_id: str, semantic_id: str | None = None) -> list[dict]:
    """Deterministic activity timeline, derived from real records.

    No separate activity table: events are reconstructed from revisions,
    confirmations, annotations, baselines and change requests.
    """
    events: list[dict] = []

    revisions = db.execute(
        select(m.ArtifactRevision)
        .join(m.Artifact, m.ArtifactRevision.artifact_id == m.Artifact.id)
        .where(m.Artifact.project_id == project_id)
    ).scalars().all()
    for r in revisions:
        events.append({
            "at": _iso(r.created_at), "kind": "revision_created",
            "actor": r.created_by, "label": f"{r.artifact.title} r{r.revision_number}",
            "revision_id": r.id, "semantic_id": None,
        })
        if r.confirmed_at:
            events.append({
                "at": _iso(r.confirmed_at), "kind": "revision_confirmed",
                "actor": r.confirmed_by, "label": f"{r.artifact.title} r{r.revision_number}",
                "revision_id": r.id, "semantic_id": None,
            })

    confirmations = db.execute(
        select(m.Confirmation)
        .where(m.Confirmation.project_id == project_id)
    ).scalars().all()
    for c in confirmations:
        events.append({
            "at": _iso(c.confirmed_at), "kind": "confirmation",
            "actor": c.confirmed_by, "label": c.comment or "confirmed",
            "revision_id": c.artifact_revision_id, "semantic_id": None,
        })

    annotations = db.execute(
        select(m.Annotation)
        .where(m.Annotation.project_id == project_id)
    ).scalars().all()
    for a in annotations:
        events.append({
            "at": _iso(a.created_at), "kind": f"annotation_{a.type.value.lower()}",
            "actor": a.created_by, "label": a.content[:120],
            "revision_id": a.artifact_revision_id, "semantic_id": a.anchor_semantic_id,
        })

    baselines = db.execute(
        select(m.Baseline).where(m.Baseline.project_id == project_id)
    ).scalars().all()
    for b in baselines:
        events.append({
            "at": _iso(b.created_at), "kind": "baseline_created",
            "actor": b.created_by, "label": b.name,
            "revision_id": None, "semantic_id": None,
        })

    crs = db.execute(
        select(m.ChangeRequest).where(m.ChangeRequest.project_id == project_id)
    ).scalars().all()
    for cr in crs:
        events.append({
            "at": _iso(cr.created_at), "kind": "change_request_created",
            "actor": cr.created_by, "label": f"{cr.code} — {cr.requested_change[:120]}",
            "revision_id": None, "semantic_id": cr.code,
        })

    if semantic_id:
        events = [
            e for e in events
            if e["semantic_id"] == semantic_id or e["revision_id"] is not None
        ]

    events.sort(key=lambda e: e["at"] or "")
    return events


# ---------------------------------------------------------------------------
# Database designer (CRUD over the structured model) + ERD layout
# ---------------------------------------------------------------------------

# Semantic identity of a field is fixed at creation. Renaming a field
# (or its table) never changes the semantic id — only display names
# change. This is what lets traces/annotations survive design edits.

_FIELD_EDITABLE = {
    "name", "data_type", "length", "nullable", "default",
    "primary_key", "foreign_key", "reference", "description", "remark",
}


def rename_table(db: Session, table_id: str, name: str) -> m.DatabaseTable:
    table = get_or_404(db, m.DatabaseTable, table_id, "Table")
    table.name = name
    ensure_semantic_object(
        db,
        project_id=table.schema.project_id,
        semantic_id=table.semantic_id,
        object_type=m.SemanticObjectType.DB_TABLE,
        display_name=name,
        entity_ref=table.id,
    )
    db.commit()
    return table


def update_field(db: Session, field_id: str, **changes) -> m.DatabaseField:
    field = get_or_404(db, m.DatabaseField, field_id, "Field")
    for key, value in changes.items():
        if key not in _FIELD_EDITABLE:
            raise DomainError(f"Cannot update field attribute '{key}'", status_code=422)
        setattr(field, key, value)
    ensure_semantic_object(
        db,
        project_id=field.table.schema.project_id,
        semantic_id=field.semantic_id,
        object_type=m.SemanticObjectType.DB_FIELD,
        display_name=f"{field.table.name}.{field.name}",
        entity_ref=field.id,
    )
    db.commit()
    return field


def delete_field(db: Session, field_id: str) -> None:
    field = get_or_404(db, m.DatabaseField, field_id, "Field")
    db.execute(
        delete(m.DatabaseRelation).where(
            or_(
                m.DatabaseRelation.from_field_semantic_id == field.semantic_id,
                m.DatabaseRelation.to_field_semantic_id == field.semantic_id,
            )
        )
    )
    db.delete(field)
    db.commit()


def delete_table(db: Session, table_id: str) -> None:
    table = get_or_404(db, m.DatabaseTable, table_id, "Table")
    field_ids = [f.semantic_id for f in table.fields]
    if field_ids:
        db.execute(
            delete(m.DatabaseRelation).where(
                or_(
                    m.DatabaseRelation.from_field_semantic_id.in_(field_ids),
                    m.DatabaseRelation.to_field_semantic_id.in_(field_ids),
                )
            )
        )
    db.delete(table)  # fields cascade via ORM relationship
    db.commit()


def delete_relation(db: Session, relation_id: str) -> None:
    relation = get_or_404(db, m.DatabaseRelation, relation_id, "Relation")
    db.delete(relation)
    db.commit()


def save_erd_layout(db: Session, schema_id: str, layout: dict) -> m.DatabaseSchema:
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    schema.layout = layout or {}
    db.commit()
    return schema


def get_erd_layout(db: Session, schema_id: str) -> dict:
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    return schema.layout or {}


def db_design_snapshot(db: Session, schema_id: str) -> dict:
    """Canonical, semantic-id-keyed snapshot of the structured DB design.

    Used by the semantic diff (P1-E): keys are stable semantic ids, so a
    diff over two snapshots reports ADDED/REMOVED/CHANGED objects rather
    than positional noise.
    """
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    tables = db.execute(
        select(m.DatabaseTable).where(m.DatabaseTable.schema_id == schema_id)
    ).scalars().all()
    table_sid = {t.id: t.semantic_id for t in tables}
    tables_out: dict[str, dict] = {
        t.semantic_id: {"name": t.name, "description": t.description, "fields": {}}
        for t in tables
    }
    fields = db.execute(
        select(m.DatabaseField).where(m.DatabaseField.table_id.in_(table_sid))
    ).scalars().all()
    for f in fields:
        tables_out[table_sid[f.table_id]]["fields"][f.semantic_id] = {
            "name": f.name,
            "data_type": f.data_type,
            "length": f.length,
            "nullable": f.nullable,
            "default": f.default,
            "primary_key": f.primary_key,
            "foreign_key": f.foreign_key,
            "reference": f.reference,
            "description": f.description,
            "remark": f.remark,
        }
    relations = {
        r.semantic_id: {
            "from": r.from_field_semantic_id,
            "to": r.to_field_semantic_id,
            "type": r.relation_type,
        }
        for r in db.execute(
            select(m.DatabaseRelation).where(m.DatabaseRelation.schema_id == schema_id)
        ).scalars()
    }
    return {"tables": tables_out, "relations": relations}


# ---------------------------------------------------------------------------
# Revision compare + semantic diff
# ---------------------------------------------------------------------------


def _section_text(section: dict) -> str:
    lines: list[str] = [f"## {section.get('heading') or ''}"]
    if section.get("content") is not None:
        lines.append(_doc_text(section.get("content")))
    else:  # legacy blocks fallback
        for blk in section.get("blocks", []):
            kind = blk.get("kind")
            if kind == "heading":
                lines.append(f"{'#' * (blk.get('level') or 2)} {blk.get('text') or ''}")
            elif kind in ("bullet_list", "numbered_list"):
                for item in blk.get("items", []):
                    lines.append(f"- {item}")
            elif kind == "table":
                lines.append(" | ".join(blk.get("header", [])))
                for row in blk.get("rows", []):
                    lines.append(" | ".join(row))
            elif kind == "code":
                lines.append("```")
                lines.extend((blk.get("text") or "").splitlines())
                lines.append("```")
            else:
                lines.append(blk.get("text") or "")
    return "\n".join(lines)


def text_diff(a: str, b: str) -> list[dict]:
    """Line-level diff (insert/delete/equal) using difflib."""
    sm = difflib.SequenceMatcher(a=a.splitlines(), b=b.splitlines())
    out: list[dict] = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            out.append({"op": "equal", "lines": a.splitlines()[i1:i2]})
        elif tag == "replace":
            out.append({"op": "delete", "lines": a.splitlines()[i1:i2]})
            out.append({"op": "insert", "lines": b.splitlines()[j1:j2]})
        elif tag == "delete":
            out.append({"op": "delete", "lines": a.splitlines()[i1:i2]})
        elif tag == "insert":
            out.append({"op": "insert", "lines": b.splitlines()[j1:j2]})
    return out


def document_diff(db: Session, rev_a_id: str, rev_b_id: str) -> list[dict]:
    """Semantic document diff keyed by stable section ids."""
    a = {s["id"]: s for s in get_document(db, rev_a_id)["sections"]}
    b = {s["id"]: s for s in get_document(db, rev_b_id)["sections"]}
    changes: list[dict] = []
    for sid in sorted(set(a) - set(b)):
        changes.append({"kind": "REMOVED", "object": "SECTION", "semantic_id": sid, "label": a[sid].get("heading") or sid})
    for sid in sorted(set(b) - set(a)):
        changes.append({"kind": "ADDED", "object": "SECTION", "semantic_id": sid, "label": b[sid].get("heading") or sid})
    for sid in sorted(set(a) & set(b)):
        ta, tb = _section_text(a[sid]), _section_text(b[sid])
        if ta != tb:
            changes.append({
                "kind": "CHANGED", "object": "SECTION", "semantic_id": sid,
                "label": b[sid].get("heading") or sid, "text_diff": text_diff(ta, tb),
            })
    return changes


def semantic_db_diff(a: dict, b: dict) -> list[dict]:
    """Semantic diff between two canonical DB design snapshots.

    Keys are stable semantic ids, so changes are reported as ADDED /
    REMOVED / CHANGED objects and attributes — never positional noise.
    """
    changes: list[dict] = []
    a_t, b_t = a.get("tables", {}), b.get("tables", {})
    for sid in sorted(set(a_t) - set(b_t)):
        changes.append({"kind": "REMOVED", "object": "TABLE", "semantic_id": sid, "label": a_t[sid].get("name", sid)})
    for sid in sorted(set(b_t) - set(a_t)):
        changes.append({"kind": "ADDED", "object": "TABLE", "semantic_id": sid, "label": b_t[sid].get("name", sid)})
    for sid in sorted(set(a_t) & set(b_t)):
        ta, tb = a_t[sid], b_t[sid]
        if ta.get("name") != tb.get("name"):
            changes.append({"kind": "CHANGED", "object": "TABLE", "semantic_id": sid, "attribute": "name", "before": ta.get("name"), "after": tb.get("name")})
        fa, fb = ta.get("fields", {}), tb.get("fields", {})
        for fid in sorted(set(fa) - set(fb)):
            changes.append({"kind": "REMOVED", "object": "FIELD", "semantic_id": fid, "label": fa[fid].get("name", fid)})
        for fid in sorted(set(fb) - set(fa)):
            changes.append({"kind": "ADDED", "object": "FIELD", "semantic_id": fid, "label": fb[fid].get("name", fid)})
        for fid in sorted(set(fa) & set(fb)):
            for attr in sorted(set(fa[fid]) | set(fb[fid])):
                if fa[fid].get(attr) != fb[fid].get(attr):
                    changes.append({
                        "kind": "CHANGED", "object": "FIELD", "semantic_id": fid,
                        "attribute": attr, "before": fa[fid].get(attr), "after": fb[fid].get(attr),
                    })

    a_r, b_r = a.get("relations", {}), b.get("relations", {})
    for rid in sorted(set(a_r) - set(b_r)):
        changes.append({"kind": "REMOVED", "object": "RELATION", "semantic_id": rid})
    for rid in sorted(set(b_r) - set(a_r)):
        changes.append({"kind": "ADDED", "object": "RELATION", "semantic_id": rid})
    for rid in sorted(set(a_r) & set(b_r)):
        if a_r[rid].get("type") != b_r[rid].get("type"):
            changes.append({"kind": "CHANGED", "object": "RELATION", "semantic_id": rid, "attribute": "type", "before": a_r[rid].get("type"), "after": b_r[rid].get("type")})
    return changes


def semantic_flow_diff(a: dict, b: dict) -> list[dict]:
    """Approval/process step counts compared by stable step ids where available."""
    changes: list[dict] = []
    fa = a.get("flows", {})
    fb = b.get("flows", {})
    for fid in sorted(set(fa) - set(fb)):
        changes.append({"kind": "REMOVED", "object": "FLOW", "semantic_id": fid})
    for fid in sorted(set(fb) - set(fa)):
        changes.append({"kind": "ADDED", "object": "FLOW", "semantic_id": fid})
    for fid in sorted(set(fa) & set(fb)):
        sa = fa[fid].get("steps", [])
        sb = fb[fid].get("steps", [])
        if len(sa) != len(sb):
            changes.append({"kind": "CHANGED", "object": "FLOW", "semantic_id": fid, "attribute": "steps", "before": len(sa), "after": len(sb)})
    return changes


def semantic_diff(a: dict, b: dict) -> list[dict]:
    """Combined semantic diff (DB design + flows) over stable ids."""
    return semantic_db_diff(a, b) + semantic_flow_diff(a, b)


def snapshot_database_into_revision(db: Session, revision_id: str, schema_id: str) -> m.ArtifactRevision:
    """Embed the current canonical DB design into a DRAFT revision snapshot.

    This creates versioned DB data inside the document revision so that a
    later semantic diff can compare two points in time by stable ids.
    Confirmed revisions are immutable and rejected.
    """
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    require_editable(revision)
    snapshot = dict(revision.snapshot or {})
    snapshot["database"] = db_design_snapshot(db, schema_id)
    revision.snapshot = snapshot
    db.commit()
    return revision


def revision_diff(db: Session, rev_a_id: str, rev_b_id: str) -> dict:
    """Full comparison of two revisions: document + DB semantic diff."""
    ra = get_or_404(db, m.ArtifactRevision, rev_a_id, "Revision")
    rb = get_or_404(db, m.ArtifactRevision, rev_b_id, "Revision")
    sa, sb = ra.snapshot or {}, rb.snapshot or {}
    db_diff = semantic_diff(sa.get("database", {}), sb.get("database", {}))
    return {
        "a": {"id": ra.id, "revision_number": ra.revision_number, "status": ra.status.value},
        "b": {"id": rb.id, "revision_number": rb.revision_number, "status": rb.status.value},
        "document_diff": document_diff(db, rev_a_id, rev_b_id),
        "database_diff": db_diff,
    }


# ---------------------------------------------------------------------------
# Deterministic impact analysis (graph/rule based, no AI)
# ---------------------------------------------------------------------------


def _graph_adjacency(edges) -> tuple[dict, dict]:
    out: dict[str, list[tuple[str, str]]] = {}
    inc: dict[str, list[tuple[str, str]]] = {}
    for e in edges:
        out.setdefault(e.source_semantic_id, []).append((e.target_semantic_id, e.relation_type.value))
        inc.setdefault(e.target_semantic_id, []).append((e.source_semantic_id, e.relation_type.value))
    return out, inc


def impact_paths(db: Session, project_id: str, semantic_id: str, max_depth: int = 3) -> dict:
    """Bounded transitive impact with relation-path explanation.

    Only TraceLink rows are traversed — nothing is inferred.
    """
    edges = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
    ).scalars().all()
    out, inc = _graph_adjacency(edges)

    def walk(adj):
        results: list[list[dict]] = []
        queue: list[tuple[str, list[dict]]] = [(semantic_id, [])]
        while queue:
            node, path = queue.pop(0)
            for nxt, rel in adj.get(node, []):
                new_path = path + [{"semantic_id": nxt, "relation": rel}]
                results.append(new_path)
                if len(new_path) < max_depth:
                    queue.append((nxt, new_path))
        return results

    return {"downstream": walk(out), "upstream": walk(inc)}


def impact_analysis(db: Session, project_id: str, semantic_id: str, max_depth: int = 3) -> dict:
    direct = impact_of(db, project_id, semantic_id)
    paths = impact_paths(db, project_id, semantic_id, max_depth=max_depth)
    return {
        "semantic_id": semantic_id,
        "max_depth": max_depth,
        "direct": direct,
        "paths": paths,
    }


# ---------------------------------------------------------------------------
# Impact analysis v2 — named change sets + rule-based severity
# ---------------------------------------------------------------------------

CHANGE_TYPES = {"MODIFIED", "ADDED", "REMOVED", "RENAMED"}
_SEVERITY_BY_DEPTH = {1: "DIRECT", 2: "HIGH", 3: "MEDIUM"}
_HIGH_BLAST_RADIUS = {
    m.SemanticObjectType.DB_SCHEMA.value,
    m.SemanticObjectType.DB_TABLE.value,
    m.SemanticObjectType.DB_FIELD.value,
    m.SemanticObjectType.DB_RELATION.value,
}


def _severity(depth: int, object_type: str | None, change_type: str | None) -> str:
    base = _SEVERITY_BY_DEPTH.get(depth, "LOW")
    order = ["LOW", "MEDIUM", "HIGH", "DIRECT"]
    idx = order.index(base)
    if object_type in _HIGH_BLAST_RADIUS:
        idx = min(idx + 1, len(order) - 1)
    if change_type == "REMOVED" and object_type in {
        m.SemanticObjectType.DB_TABLE.value,
        m.SemanticObjectType.DB_SCHEMA.value,
    }:
        idx = min(idx + 1, len(order) - 1)
    return order[idx]


def create_change_set(
    db: Session,
    *,
    project_id: str,
    name: str,
    description: str | None = None,
    items: list[dict] | None = None,
    actor="local-user",
    actor_id: str | None = None,
) -> dict:
    get_or_404(db, m.Project, project_id, "Project")
    cs = m.ChangeSet(
        project_id=project_id, name=name, description=description,
        created_by=actor, actor_id=actor_id,
    )
    db.add(cs)
    db.flush()
    built: list[dict] = []
    for item in items or []:
        ct = (item.get("change_type") or "MODIFIED").upper()
        if ct not in CHANGE_TYPES:
            raise DomainError(f"invalid change type: {ct}")
        ci = m.ChangeItem(
            change_set_id=cs.id, semantic_id=item["semantic_id"],
            change_type=ct, rationale=item.get("rationale"),
        )
        db.add(ci)
        built.append({"semantic_id": item["semantic_id"], "change_type": ct})
    db.commit()
    db.refresh(cs)
    return {"id": cs.id, "project_id": cs.project_id, "name": cs.name,
            "description": cs.description, "items": built,
            "created_at": cs.created_at.isoformat(), "created_by": cs.created_by}


def list_change_sets(db: Session, project_id: str | None = None) -> list[dict]:
    q = select(m.ChangeSet).order_by(m.ChangeSet.created_at.desc())
    if project_id:
        q = q.where(m.ChangeSet.project_id == project_id)
    out = []
    for cs in db.execute(q).scalars().all():
        items = db.execute(
            select(m.ChangeItem).where(m.ChangeItem.change_set_id == cs.id)
        ).scalars().all()
        out.append({
            "id": cs.id, "project_id": cs.project_id, "name": cs.name,
            "description": cs.description, "created_at": cs.created_at.isoformat(),
            "created_by": cs.created_by,
            "items": [{"semantic_id": i.semantic_id, "change_type": i.change_type,
                       "rationale": i.rationale} for i in items],
        })
    return out


def impact_analysis_v2(
    db: Session, project_id: str, semantic_id: str, max_depth: int = 4
) -> dict:
    """Rule-based impact over transitive trace links.

    Severity: DIRECT (1 hop) > HIGH (2) > MEDIUM (3) > LOW (4+).
    DB schema/table/field/relation and releases bump severity one level;
    REMOVED of a DB table/schema bumps again. Every result carries the
    explicit relation path used, so nothing is inferred silently.
    """
    edges = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
    ).scalars().all()
    out_adj, inc_adj = _graph_adjacency(edges)

    types = {
        so.semantic_id: so.object_type.value
        for so in db.execute(
            select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
        ).scalars().all()
    }

    def walk(adj, direction: str) -> list[dict]:
        results: dict[str, dict] = {}
        queue: list[tuple[str, list[dict], int]] = [(semantic_id, [], 0)]
        while queue:
            node, path, depth = queue.pop(0)
            for nxt, rel in adj.get(node, []):
                if depth + 1 > max_depth:
                    continue
                new_path = path + [{"semantic_id": nxt, "relation": rel}]
                sev = _severity(depth + 1, types.get(nxt), None)
                prev = results.get(nxt)
                if prev is None or _severity_order(sev) > _severity_order(prev["severity"]):
                    results[nxt] = {
                        "semantic_id": nxt,
                        "object_type": types.get(nxt),
                        "severity": sev,
                        "depth": depth + 1,
                        "direction": direction,
                        "path": new_path,
                    }
                queue.append((nxt, new_path, depth + 1))
        return sorted(results.values(), key=lambda r: _severity_order(r["severity"]), reverse=True)

    return {
        "semantic_id": semantic_id,
        "object_type": types.get(semantic_id),
        "max_depth": max_depth,
        "affected": walk(out_adj, "downstream") + walk(inc_adj, "upstream"),
    }


def _severity_order(sev: str) -> int:
    return {"LOW": 1, "MEDIUM": 2, "HIGH": 3, "DIRECT": 4}.get(sev, 0)


# ---------------------------------------------------------------------------
# Project memory — semantic context for the right-hand panel
# ---------------------------------------------------------------------------

_DB_OBJECT_TYPES = {
    m.SemanticObjectType.DB_SCHEMA,
    m.SemanticObjectType.DB_TABLE,
    m.SemanticObjectType.DB_FIELD,
    m.SemanticObjectType.DB_RELATION,
}


def semantic_context(db: Session, project_id: str, semantic_id: str) -> dict:
    """Structured context for one semantic object. Never fabricates."""
    so = db.execute(
        select(m.SemanticObject).where(
            m.SemanticObject.project_id == project_id,
            m.SemanticObject.semantic_id == semantic_id,
        )
    ).scalar_one_or_none()
    if so is None:
        raise DomainError(f"Unknown semantic object '{semantic_id}'", status_code=404)

    out: dict = {
        "semantic_id": so.semantic_id,
        "object_type": so.object_type.value,
        "display_name": so.display_name,
        "entity_ref": so.entity_ref,
        "status": None,
        "confirmed": None,
        "revision": None,
        "evidence": [],
    }

    if so.object_type == m.SemanticObjectType.REQUIREMENT and so.entity_ref:
        req = db.get(m.Requirement, so.entity_ref)
        if req:
            out["status"] = req.status.value
            out["confirmed"] = req.status.value == m.RequirementStatus.CONFIRMED.value
            out["owner"] = {"type": "Requirement", "code": req.code, "priority": req.priority}
    elif so.object_type in _DB_OBJECT_TYPES:
        out["status"] = "live"
        out["owner"] = {"type": "Database design", "note": "working model — frozen at DR revision snapshot, not here"}
    elif so.object_type == m.SemanticObjectType.DOCUMENT_SECTION and so.entity_ref:
        artifact = db.get(m.Artifact, so.entity_ref)
        if artifact:
            revs = db.execute(
                select(m.ArtifactRevision)
                .where(m.ArtifactRevision.artifact_id == artifact.id)
                .order_by(m.ArtifactRevision.revision_number.desc())
            ).scalars().all()
            holders = [
                r for r in revs
                if any(s.get("id") == semantic_id for s in (r.snapshot or {}).get("sections") or [])
            ]
            if holders:
                latest = holders[0]
                out["status"] = latest.status.value
                out["confirmed"] = latest.status.value == m.RevisionStatus.CONFIRMED.value
                out["revision"] = {
                    "id": latest.id,
                    "revision_number": latest.revision_number,
                    "artifact_title": artifact.title,
                    "created_by": latest.created_by,
                    "created_at": latest.created_at.isoformat(),
                    "confirmed_by": latest.confirmed_by,
                    "confirmed_at": latest.confirmed_at.isoformat() if latest.confirmed_at else None,
                }
                out["evidence"] = [
                    {
                        "confirmed_by": c.confirmed_by,
                        "confirmed_at": c.confirmed_at.isoformat(),
                        "comment": c.comment,
                        "evidence": c.evidence,
                    }
                    for c in db.execute(
                        select(m.Confirmation).where(
                            m.Confirmation.artifact_revision_id == latest.id
                        )
                    ).scalars()
                ]

    elif so.object_type == m.SemanticObjectType.PROCESS_FLOW and so.entity_ref:
        flow = db.get(m.ProcessFlow, so.entity_ref)
        if flow:
            out["status"] = "live"
            out["owner"] = {"type": "Process flow", "name": flow.name, "description": flow.description}
    elif so.object_type == m.SemanticObjectType.PROCESS_STEP and so.entity_ref:
        step = db.get(m.ProcessStep, so.entity_ref)
        if step:
            out["status"] = "live"
            out["owner"] = {"type": "Process step", "name": step.name, "step_type": step.step_type, "flow": step.flow.name}
    elif so.object_type == m.SemanticObjectType.API_ENDPOINT and so.entity_ref:
        api = db.get(m.APIEndpoint, so.entity_ref)
        if api:
            out["status"] = "live"
            out["owner"] = {"type": "API endpoint", "method": api.method, "path": api.path, "summary": api.summary}
    elif so.object_type == m.SemanticObjectType.ARCHITECTURE_NODE and so.entity_ref:
        node = db.get(m.ArchitectureNode, so.entity_ref)
        if node:
            out["status"] = "live"
            out["owner"] = {"type": "Architecture node", "name": node.name, "node_type": node.node_type, "technology": node.technology}
    elif so.object_type == m.SemanticObjectType.DECISION and so.entity_ref:
        d = db.get(m.Decision, so.entity_ref)
        if d:
            out["status"] = "recorded"
            out["owner"] = {"type": "Decision", "title": d.title, "decided_by": d.decided_by, "content": d.content}
    elif so.object_type == m.SemanticObjectType.ASSUMPTION and so.entity_ref:
        a = db.get(m.Assumption, so.entity_ref)
        if a:
            out["status"] = a.status
            out["owner"] = {"type": "Assumption", "content": a.content}
    elif so.object_type == m.SemanticObjectType.CLARIFICATION and so.entity_ref:
        c = db.get(m.Clarification, so.entity_ref)
        if c:
            out["status"] = "resolved" if c.resolved else "open"
            out["owner"] = {"type": "Clarification", "question": c.question, "answer": c.answer}

    # annotations on this object
    anns = db.execute(
        select(m.Annotation).where(
            m.Annotation.project_id == project_id,
            m.Annotation.anchor_semantic_id == semantic_id,
        )
    ).scalars().all()
    out["annotations"] = {
        "total": len(anns),
        "open": sum(1 for a in anns if a.status != m.AnnotationStatus.RESOLVED),
        "by_type": dict(Counter(a.type.value for a in anns)),
    }

    return out


# ---------------------------------------------------------------------------
# Semantic search (favours semantic objects over file names)
# ---------------------------------------------------------------------------


def search_semantic(db: Session, project_id: str, query: str, limit: int = 60) -> list[dict]:
    q = (query or "").strip().lower()
    if not q:
        return []
    results: list[dict] = []

    sos = db.execute(
        select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
    ).scalars().all()
    for so in sos:
        hay = f"{so.semantic_id} {so.display_name or ''}".lower()
        if q in hay:
            results.append({
                "kind": "semantic_object",
                "semantic_id": so.semantic_id,
                "object_type": so.object_type.value,
                "title": so.display_name or so.semantic_id,
            })

    anns = db.execute(
        select(m.Annotation).where(m.Annotation.project_id == project_id)
    ).scalars().all()
    for a in anns:
        if q in (a.content or "").lower():
            results.append({
                "kind": "annotation",
                "semantic_id": a.anchor_semantic_id,
                "object_type": a.type.value,
                "title": (a.content or "")[:120],
                "status": a.status.value,
            })

    crs = db.execute(
        select(m.ChangeRequest).where(m.ChangeRequest.project_id == project_id)
    ).scalars().all()
    for cr in crs:
        if q in cr.code.lower() or q in (cr.requested_change or "").lower():
            results.append({
                "kind": "change_request",
                "semantic_id": cr.code,
                "object_type": "CHANGE_REQUEST",
                "title": (cr.requested_change or "")[:120],
                "status": cr.status.value,
            })

    return results[:limit]


# ---------------------------------------------------------------------------
# Process flow designer (structured; diagram is a view over this)
# ---------------------------------------------------------------------------

FLOW_STEP_TYPES = {"START", "ACTION", "DECISION", "APPROVAL", "SYSTEM", "MANUAL", "END"}


def create_flow(
    db: Session, *, project_id: str, name: str, semantic_id: str, description=None, actor="local-user"
) -> m.ProcessFlow:
    flow = m.ProcessFlow(project_id=project_id, semantic_id=semantic_id, name=name, description=description)
    db.add(flow)
    db.flush()
    ensure_semantic_object(
        db, project_id=project_id, semantic_id=semantic_id,
        object_type=m.SemanticObjectType.PROCESS_FLOW, display_name=name, entity_ref=flow.id,
    )
    db.commit()
    return flow


def add_flow_step(
    db: Session, *, flow_id: str, name: str, step_type: str = "ACTION",
    semantic_id: str | None = None, description=None,
) -> m.ProcessStep:
    flow = get_or_404(db, m.ProcessFlow, flow_id, "ProcessFlow")
    if step_type not in FLOW_STEP_TYPES:
        raise DomainError(f"Unknown step type '{step_type}'", status_code=422)
    semantic_id = semantic_id or f"flow_step_{name.lower().replace(' ', '_').replace('-', '_')}"
    existing = db.execute(
        select(m.ProcessStep.id).where(
            m.ProcessStep.flow_id == flow_id, m.ProcessStep.semantic_id == semantic_id
        )
    ).scalar_one_or_none()
    if existing:
        raise DomainError(f"Step semantic id '{semantic_id}' already exists in this flow")
    position = (
        db.execute(select(m.ProcessStep.id).where(m.ProcessStep.flow_id == flow_id)).scalars().all()
    )
    step = m.ProcessStep(
        flow_id=flow_id, semantic_id=semantic_id, name=name, step_type=step_type,
        description=description, position=len(position),
    )
    db.add(step)
    db.flush()
    ensure_semantic_object(
        db, project_id=flow.project_id, semantic_id=semantic_id,
        object_type=m.SemanticObjectType.PROCESS_STEP, display_name=name, entity_ref=step.id,
    )
    db.commit()
    return step


def add_flow_transition(
    db: Session, *, flow_id: str, from_step_semantic_id: str, to_step_semantic_id: str,
    label: str | None = None, condition: str | None = None,
) -> m.ProcessTransition:
    flow = get_or_404(db, m.ProcessFlow, flow_id, "ProcessFlow")
    for sid in (from_step_semantic_id, to_step_semantic_id):
        exists = db.execute(
            select(m.ProcessStep.id).where(
                m.ProcessStep.flow_id == flow_id, m.ProcessStep.semantic_id == sid
            )
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(f"Unknown step '{sid}' in this flow", status_code=422)
    semantic_id = f"flow_transition_{from_step_semantic_id}__{to_step_semantic_id}"
    t = m.ProcessTransition(
        flow_id=flow_id, semantic_id=semantic_id,
        from_step_semantic_id=from_step_semantic_id, to_step_semantic_id=to_step_semantic_id,
        label=label, condition=condition,
    )
    db.add(t)
    db.commit()
    return t


def delete_flow_step(db: Session, step_id: str) -> None:
    step = get_or_404(db, m.ProcessStep, step_id, "ProcessStep")
    flow_id = step.flow_id
    db.execute(
        delete(m.ProcessTransition).where(
            or_(
                m.ProcessTransition.from_step_semantic_id == step.semantic_id,
                m.ProcessTransition.to_step_semantic_id == step.semantic_id,
            )
        )
    )
    db.delete(step)
    db.commit()


def delete_flow_transition(db: Session, transition_id: str) -> None:
    t = get_or_404(db, m.ProcessTransition, transition_id, "ProcessTransition")
    db.delete(t)
    db.commit()


def save_flow_layout(db: Session, flow_id: str, layout: dict) -> m.ProcessFlow:
    flow = get_or_404(db, m.ProcessFlow, flow_id, "ProcessFlow")
    flow.layout = layout or {}
    db.commit()
    return flow


def list_flows(db: Session, project_id: str) -> list[dict]:
    flows = db.execute(
        select(m.ProcessFlow).where(m.ProcessFlow.project_id == project_id)
    ).scalars().all()
    out = []
    for f in flows:
        steps = db.execute(
            select(m.ProcessStep).where(m.ProcessStep.flow_id == f.id).order_by(m.ProcessStep.position)
        ).scalars().all()
        transitions = db.execute(
            select(m.ProcessTransition).where(m.ProcessTransition.flow_id == f.id)
        ).scalars().all()
        out.append({
            "id": f.id, "semantic_id": f.semantic_id, "name": f.name,
            "description": f.description, "layout": f.layout or {},
            "steps": [
                {"id": s.id, "semantic_id": s.semantic_id, "name": s.name,
                 "step_type": s.step_type, "position": s.position, "description": s.description}
                for s in steps
            ],
            "transitions": [
                {"id": t.id, "semantic_id": t.semantic_id,
                 "from": t.from_step_semantic_id, "to": t.to_step_semantic_id,
                 "label": t.label, "condition": t.condition}
                for t in transitions
            ],
        })
    return out


# ---------------------------------------------------------------------------
# API design workspace (structured, not free-text)
# ---------------------------------------------------------------------------


def _slug_path(path: str) -> str:
    return path.strip("/").replace("/", "_").replace("{", "").replace("}", "").replace("-", "_")


def create_api_endpoint(
    db: Session, *, project_id: str, method: str, path: str, summary=None,
    semantic_id: str | None = None, description=None, authentication="NONE", actor="local-user",
) -> m.APIEndpoint:
    semantic_id = semantic_id or f"api_{method.lower()}_{_slug_path(path)}"
    existing = db.execute(
        select(m.APIEndpoint.id).where(
            m.APIEndpoint.project_id == project_id, m.APIEndpoint.semantic_id == semantic_id
        )
    ).scalar_one_or_none()
    if existing:
        raise DomainError(f"API semantic id '{semantic_id}' already exists")
    ep = m.APIEndpoint(
        project_id=project_id, semantic_id=semantic_id, method=method.upper(), path=path,
        summary=summary, description=description, authentication=authentication,
    )
    db.add(ep)
    db.flush()
    ensure_semantic_object(
        db, project_id=project_id, semantic_id=semantic_id,
        object_type=m.SemanticObjectType.API_ENDPOINT, display_name=f"{method.upper()} {path}",
        entity_ref=ep.id,
    )
    db.commit()
    return ep


def update_api_endpoint(db: Session, endpoint_id: str, **changes) -> m.APIEndpoint:
    ep = get_or_404(db, m.APIEndpoint, endpoint_id, "APIEndpoint")
    for k, v in changes.items():
        if hasattr(ep, k):
            setattr(ep, k, v)
    ensure_semantic_object(
        db, project_id=ep.project_id, semantic_id=ep.semantic_id,
        object_type=m.SemanticObjectType.API_ENDPOINT,
        display_name=f"{ep.method} {ep.path}", entity_ref=ep.id,
    )
    db.commit()
    return ep


def _add_api_child(db, model, *, endpoint_id, **fields):
    get_or_404(db, m.APIEndpoint, endpoint_id, "APIEndpoint")
    child = model(endpoint_id=endpoint_id, **fields)
    db.add(child)
    db.commit()
    return child


def _delete_api_child(db, model, child_id):
    child = get_or_404(db, model, child_id, model.__name__)
    db.delete(child)
    db.commit()


def list_api_endpoints(db: Session, project_id: str) -> list[dict]:
    eps = db.execute(
        select(m.APIEndpoint).where(m.APIEndpoint.project_id == project_id).order_by(m.APIEndpoint.path)
    ).scalars().all()
    out = []
    for ep in eps:
        out.append({
            "id": ep.id, "semantic_id": ep.semantic_id, "method": ep.method, "path": ep.path,
            "summary": ep.summary, "description": ep.description, "authentication": ep.authentication,
            "parameters": [
                {"id": p.id, "name": p.name, "location": p.location, "data_type": p.data_type,
                 "required": p.required, "description": p.description} for p in ep.parameters
            ],
            "request_fields": [
                {"id": f.id, "name": f.name, "data_type": f.data_type, "required": f.required,
                 "description": f.description} for f in ep.request_fields
            ],
            "response_fields": [
                {"id": f.id, "status_code": f.status_code, "name": f.name, "data_type": f.data_type,
                 "description": f.description} for f in ep.response_fields
            ],
            "error_responses": [
                {"id": e.id, "status_code": e.status_code, "message": e.message,
                 "description": e.description} for e in ep.error_responses
            ],
        })
    return out


# ---------------------------------------------------------------------------
# OpenAPI 3.x import/export (structured API design as the source of truth)
# ---------------------------------------------------------------------------


def _parse_openapi(text: str) -> dict:
    """Parse a JSON or YAML OpenAPI 3.x document into a raw spec dict."""
    import yaml as _yaml

    try:
        spec = _json.loads(text)
    except Exception:
        try:
            spec = _yaml.safe_load(text)
        except Exception as exc:
            raise DomainError(f"OpenAPI document is neither valid JSON nor YAML: {exc}", status_code=422)
    if not isinstance(spec, dict) or "paths" not in spec or not isinstance(spec["paths"], dict):
        raise DomainError("OpenAPI document must be an object with a 'paths' object", status_code=422)
    return spec


def _resolve_schema_type(schema: dict) -> str:
    if not isinstance(schema, dict):
        return "string"
    if "type" in schema and isinstance(schema["type"], str):
        return schema["type"]
    if "$ref" in schema:
        return schema["$ref"].split("/")[-1]
    if "anyOf" in schema or "oneOf" in schema:
        return "union"
    return "object"


def openapi_to_endpoints(spec: dict) -> list[dict]:
    """Normalize an OpenAPI 3.x spec into Document Again endpoint dicts."""
    out = []
    for path, path_item in spec.get("paths", {}).items():
        if not isinstance(path_item, dict):
            continue
        for method in ("get", "post", "put", "delete", "patch", "head", "options"):
            op = path_item.get(method)
            if not isinstance(op, dict):
                continue
            parameters = []
            for p in op.get("parameters", []) or []:
                if isinstance(p, dict):
                    sch = p.get("schema") or {}
                    parameters.append({
                        "name": p.get("name", ""), "location": p.get("in", "query"),
                        "data_type": _resolve_schema_type(sch), "required": bool(p.get("required")),
                        "description": p.get("description"),
                    })
            request_fields = []
            if "requestBody" in op and isinstance(op["requestBody"], dict):
                content = op["requestBody"].get("content", {})
                for ctype, media in content.items():
                    schema = (media or {}).get("schema") or {}
                    props = schema.get("properties", {}) if isinstance(schema, dict) else {}
                    for fname, fsch in props.items():
                        if isinstance(fsch, dict):
                            request_fields.append({
                                "name": fname, "data_type": _resolve_schema_type(fsch),
                                "required": fname in (schema.get("required") or []),
                                "description": fsch.get("description"),
                            })
            response_fields = []
            error_responses = []
            for status, resp in (op.get("responses") or {}).items():
                if not isinstance(resp, dict):
                    continue
                if status.startswith("2"):
                    content = resp.get("content", {})
                    for ctype, media in content.items():
                        schema = (media or {}).get("schema") or {}
                        props = schema.get("properties", {}) if isinstance(schema, dict) else {}
                        for fname, fsch in props.items():
                            if isinstance(fsch, dict):
                                response_fields.append({
                                    "status_code": status, "name": fname,
                                    "data_type": _resolve_schema_type(fsch),
                                    "description": fsch.get("description"),
                                })
                else:
                    error_responses.append({
                        "status_code": status,
                        "message": resp.get("description") or f"HTTP {status}",
                        "description": resp.get("description"),
                    })
            out.append({
                "method": method.upper(), "path": path, "summary": op.get("summary"),
                "description": op.get("description"),
                "authentication": "BEARER" if ("security" in op and op["security"]) else "NONE",
                "parameters": parameters, "request_fields": request_fields,
                "response_fields": response_fields, "error_responses": error_responses,
            })
    return out


def preview_openapi_import(db: Session, project_id: str, text: str) -> dict:
    """Diff a parsed OpenAPI spec against existing endpoints (no writes)."""
    incoming = openapi_to_endpoints(_parse_openapi(text))
    existing = list_api_endpoints(db, project_id)
    existing_keys = {e["semantic_id"] for e in existing}
    report = {"added": [], "changed": [], "removed": [], "unchanged": [], "conflicts": []}
    for ep in incoming:
        sid = f"api_{ep['method'].lower()}_{_slug_path(ep['path'])}"
        if sid in existing_keys:
            current = next(e for e in existing if e["semantic_id"] == sid)
            if (current["method"] == ep["method"] and current["path"] == ep["path"]
                    and current["summary"] == ep["summary"]):
                report["unchanged"].append(sid)
            else:
                report["changed"].append(sid)
        else:
            report["added"].append(sid)
    incoming_keys = {f"api_{ep['method'].lower()}_{_slug_path(ep['path'])}" for ep in incoming}
    for e in existing:
        if e["semantic_id"] not in incoming_keys:
            report["removed"].append(e["semantic_id"])
    return {"project_id": project_id, **report}


def import_openapi(db: Session, project_id: str, text: str, actor="local-user") -> dict:
    """Apply an OpenAPI spec: create missing endpoints, update changed ones.

    Removed endpoints are never deleted (historical truth); they are listed
    in the report so a human can decide.
    """
    spec = _parse_openapi(text)
    preview = preview_openapi_import(db, project_id, text)
    applied = []
    for ep in openapi_to_endpoints(spec):
        sid = f"api_{ep['method'].lower()}_{_slug_path(ep['path'])}"
        existing = db.execute(
            select(m.APIEndpoint).where(
                m.APIEndpoint.project_id == project_id, m.APIEndpoint.semantic_id == sid
            )
        ).scalar_one_or_none()
        if existing is None:
            created = create_api_endpoint(
                db, project_id=project_id, method=ep["method"], path=ep["path"],
                summary=ep["summary"], semantic_id=sid, description=ep["description"],
                authentication=ep["authentication"], actor=actor,
            )
            _apply_api_children(db, created, ep)
            applied.append({"semantic_id": sid, "action": "ADDED"})
        else:
            update_api_endpoint(
                db, existing.id, method=ep["method"], path=ep["path"],
                summary=ep["summary"], description=ep["description"],
                authentication=ep["authentication"],
            )
            _replace_api_children(db, existing, ep)
            applied.append({"semantic_id": sid, "action": "UPDATED"})
    return {"project_id": project_id, "applied": applied,
            "removed_in_spec": preview["removed"], "conflicts": preview["conflicts"]}


def _apply_api_children(db, ep, ep_dict):
    for p in ep_dict["parameters"]:
        _add_api_child(db, m.ApiParameter, endpoint_id=ep.id, **p)
    for f in ep_dict["request_fields"]:
        _add_api_child(db, m.ApiRequestField, endpoint_id=ep.id, **f)
    for f in ep_dict["response_fields"]:
        _add_api_child(db, m.ApiResponseField, endpoint_id=ep.id, **f)
    for e in ep_dict["error_responses"]:
        _add_api_child(db, m.ApiErrorResponse, endpoint_id=ep.id, **e)


def _replace_api_children(db, ep, ep_dict):
    for model in (m.ApiParameter, m.ApiRequestField, m.ApiResponseField, m.ApiErrorResponse):
        for child in db.execute(
            select(model).where(model.endpoint_id == ep.id)
        ).scalars().all():
            db.delete(child)
    db.flush()
    _apply_api_children(db, ep, ep_dict)


def export_openapi(db: Session, revision_id: str) -> dict:
    """Reconstruct an OpenAPI 3.0 document from a revision's API snapshot."""
    rev = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    snapshot = rev.snapshot or {}
    api_map = snapshot.get("technical_design", {}).get("api_endpoints") or {}
    if isinstance(api_map, dict):
        endpoints = list(api_map.values())
    else:
        endpoints = api_map
    paths: dict = {}
    for ep in endpoints:
        method = (ep.get("method") or "get").lower()
        path = ep.get("path") or "/"
        op = paths.setdefault(path, {})
        op[method] = {
            "summary": ep.get("summary"),
            "description": ep.get("description"),
            "parameters": [
                {"name": p.get("name"), "in": p.get("location", "query"),
                 "required": p.get("required", False),
                 "schema": {"type": p.get("data_type", "string")},
                 "description": p.get("description")}
                for p in (ep.get("parameters") or [])
            ],
            "responses": {
                str(f.get("status_code", "200")): {
                    "description": f.get("description") or "",
                    "content": {"application/json": {"schema": {
                        "type": "object", "properties": {
                            f.get("name", "value"): {"type": f.get("data_type", "string")}
                        }
                    }}}
                }
                for f in (ep.get("response_fields") or [])
            } or {"200": {"description": "OK"}},
        }
    return {"openapi": "3.0.0", "info": {"title": "Document Again — exported API design",
                                         "version": "1.0"}, "paths": paths}


# ---------------------------------------------------------------------------
# Architecture design workspace
# ---------------------------------------------------------------------------

ARCH_NODE_TYPES = {
    "USER", "CLIENT", "SERVICE", "DATABASE", "QUEUE", "STORAGE",
    "EXTERNAL_SYSTEM", "NETWORK_ZONE", "CLOUD_SERVICE",
}


def create_architecture_diagram(
    db: Session, *, project_id: str, name: str, semantic_id: str, description=None, actor="local-user"
) -> m.ArchitectureDiagram:
    d = m.ArchitectureDiagram(project_id=project_id, semantic_id=semantic_id, name=name, description=description)
    db.add(d)
    db.commit()
    return d


def add_architecture_node(
    db: Session, *, diagram_id: str, name: str, semantic_id: str,
    node_type: str = "SERVICE", description=None, technology=None, environment=None, metadata=None,
) -> m.ArchitectureNode:
    diagram = get_or_404(db, m.ArchitectureDiagram, diagram_id, "ArchitectureDiagram")
    if node_type not in ARCH_NODE_TYPES:
        raise DomainError(f"Unknown node type '{node_type}'", status_code=422)
    node = m.ArchitectureNode(
        diagram_id=diagram_id, semantic_id=semantic_id, name=name, node_type=node_type,
        description=description, technology=technology, environment=environment, metadata_json=metadata,
    )
    db.add(node)
    db.flush()
    ensure_semantic_object(
        db, project_id=diagram.project_id, semantic_id=semantic_id,
        object_type=m.SemanticObjectType.ARCHITECTURE_NODE, display_name=name, entity_ref=node.id,
    )
    db.commit()
    return node


def add_architecture_edge(
    db: Session, *, diagram_id: str, from_node_semantic_id: str, to_node_semantic_id: str, label=None,
) -> m.ArchitectureEdge:
    diagram = get_or_404(db, m.ArchitectureDiagram, diagram_id, "ArchitectureDiagram")
    for sid in (from_node_semantic_id, to_node_semantic_id):
        exists = db.execute(
            select(m.ArchitectureNode.id).where(
                m.ArchitectureNode.diagram_id == diagram_id,
                m.ArchitectureNode.semantic_id == sid,
            )
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(f"Unknown node '{sid}' in this diagram", status_code=422)
    edge = m.ArchitectureEdge(
        diagram_id=diagram_id,
        semantic_id=f"edge_{from_node_semantic_id}__{to_node_semantic_id}",
        from_node_semantic_id=from_node_semantic_id, to_node_semantic_id=to_node_semantic_id, label=label,
    )
    db.add(edge)
    db.commit()
    return edge


def delete_architecture_node(db: Session, node_id: str) -> None:
    node = get_or_404(db, m.ArchitectureNode, node_id, "ArchitectureNode")
    db.execute(
        delete(m.ArchitectureEdge).where(
            or_(
                m.ArchitectureEdge.from_node_semantic_id == node.semantic_id,
                m.ArchitectureEdge.to_node_semantic_id == node.semantic_id,
            )
        )
    )
    db.delete(node)
    db.commit()


def delete_architecture_edge(db: Session, edge_id: str) -> None:
    edge = get_or_404(db, m.ArchitectureEdge, edge_id, "ArchitectureEdge")
    db.delete(edge)
    db.commit()


def save_architecture_layout(db: Session, diagram_id: str, layout: dict) -> m.ArchitectureDiagram:
    diagram = get_or_404(db, m.ArchitectureDiagram, diagram_id, "ArchitectureDiagram")
    diagram.layout = layout or {}
    db.commit()
    return diagram


def list_architecture_diagrams(db: Session, project_id: str) -> list[dict]:
    diagrams = db.execute(
        select(m.ArchitectureDiagram).where(m.ArchitectureDiagram.project_id == project_id)
    ).scalars().all()
    out = []
    for d in diagrams:
        nodes = db.execute(select(m.ArchitectureNode).where(m.ArchitectureNode.diagram_id == d.id)).scalars().all()
        edges = db.execute(select(m.ArchitectureEdge).where(m.ArchitectureEdge.diagram_id == d.id)).scalars().all()
        out.append({
            "id": d.id, "semantic_id": d.semantic_id, "name": d.name, "description": d.description,
            "layout": d.layout or {},
            "nodes": [
                {"id": n.id, "semantic_id": n.semantic_id, "name": n.name, "node_type": n.node_type,
                 "description": n.description, "technology": n.technology, "environment": n.environment}
                for n in nodes
            ],
            "edges": [
                {"id": e.id, "semantic_id": e.semantic_id, "from": e.from_node_semantic_id,
                 "to": e.to_node_semantic_id, "label": e.label}
                for e in edges
            ],
        })
    return out


# ---------------------------------------------------------------------------
# Decision / Assumption / Clarification project-memory surfaces
# ---------------------------------------------------------------------------


def _next_code(db: Session, project_id: str, prefix: str, model) -> str:
    count = db.execute(select(model.id).where(model.project_id == project_id)).scalars().all()
    return f"{prefix}-{len(count) + 1:04d}"


def create_decision(
    db: Session, *, project_id: str, title: str, content: str,
    decided_by="local-user", related_semantic_ids: list[str] | None = None, actor="local-user",
    actor_id: str | None = None,
) -> m.Decision:
    code = _next_code(db, project_id, "DEC", m.Decision)
    d = m.Decision(project_id=project_id, semantic_id=code, title=title, content=content, decided_by=decided_by, actor_id=actor_id)
    db.add(d)
    db.flush()
    ensure_semantic_object(
        db, project_id=project_id, semantic_id=code, object_type=m.SemanticObjectType.DECISION,
        display_name=title, entity_ref=d.id,
    )
    for sid in (related_semantic_ids or []):
        create_trace_link(db, project_id=project_id, source_semantic_id=code, target_semantic_id=sid, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
    db.commit()
    return d


def create_assumption(db: Session, *, project_id: str, content: str, related_semantic_ids=None, actor="local-user") -> m.Assumption:
    code = _next_code(db, project_id, "ASM", m.Assumption)
    a = m.Assumption(project_id=project_id, semantic_id=code, content=content, created_by=actor)
    db.add(a)
    db.flush()
    ensure_semantic_object(db, project_id=project_id, semantic_id=code, object_type=m.SemanticObjectType.ASSUMPTION, display_name=code, entity_ref=a.id)
    for sid in (related_semantic_ids or []):
        create_trace_link(db, project_id=project_id, source_semantic_id=code, target_semantic_id=sid, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
    db.commit()
    return a


def create_clarification(db: Session, *, project_id: str, question: str, answer=None, related_semantic_ids=None, actor="local-user") -> m.Clarification:
    code = _next_code(db, project_id, "CLR", m.Clarification)
    c = m.Clarification(project_id=project_id, question=question, answer=answer, asked_by=actor, resolved=answer is not None)
    db.add(c)
    db.flush()
    ensure_semantic_object(db, project_id=project_id, semantic_id=code, object_type=m.SemanticObjectType.CLARIFICATION, display_name=code, entity_ref=c.id)
    for sid in (related_semantic_ids or []):
        create_trace_link(db, project_id=project_id, source_semantic_id=code, target_semantic_id=sid, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
    db.commit()
    return c


def list_project_memory(db: Session, project_id: str) -> dict:
    decisions = db.execute(select(m.Decision).where(m.Decision.project_id == project_id)).scalars().all()
    assumptions = db.execute(select(m.Assumption).where(m.Assumption.project_id == project_id)).scalars().all()
    clarifications = db.execute(select(m.Clarification).where(m.Clarification.project_id == project_id)).scalars().all()
    links = db.execute(select(m.TraceLink).where(m.TraceLink.project_id == project_id)).scalars().all()

    def related(sid):
        return [l.target_semantic_id for l in links if l.source_semantic_id == sid and l.relation_type == m.TraceRelationType.REFERENCES]

    return {
        "decisions": [
            {"id": d.id, "code": d.semantic_id, "title": d.title, "content": d.content,
             "decided_by": d.decided_by, "decided_at": d.decided_at.isoformat(), "related": related(d.semantic_id)}
            for d in decisions
        ],
        "assumptions": [
            {"id": a.id, "code": a.semantic_id, "content": a.content, "status": a.status,
             "created_by": a.created_by, "related": related(a.semantic_id)}
            for a in assumptions
        ],
        "clarifications": [
            {"id": c.id, "code": c.semantic_id, "question": c.question, "answer": c.answer,
             "asked_by": c.asked_by, "resolved": c.resolved, "related": related(c.semantic_id)}
            for c in clarifications
        ],
    }


def promote_annotation(db: Session, *, annotation_id: str, to_kind: str, actor="local-user") -> dict:
    """Promote a comment/annotation into a first-class project-memory record.

    Provenance is retained: the source annotation id / thread is recorded
    and the annotation is marked resolved.
    """
    ann = get_or_404(db, m.Annotation, annotation_id, "Annotation")
    provenance = {
        "source": "annotation", "annotation_id": ann.id, "thread_id": ann.thread_id,
        "author": ann.created_by, "anchor": ann.anchor_semantic_id,
    }
    source_label = f"Comment Thread #{ann.thread_id}" if ann.thread_id else f"Annotation #{ann.id}"

    if to_kind == "change_request":
        cr = create_change_request(
            db, project_id=ann.project_id, requested_change=ann.content,
            affected_semantic_ids=[ann.anchor_semantic_id], reason=f"Promoted from {source_label}",
            actor=actor,
        )
        result = {"kind": "change_request", "code": cr.code, "provenance": provenance}
    elif to_kind == "decision":
        code = _next_code(db, ann.project_id, "DEC", m.Decision)
        d = m.Decision(project_id=ann.project_id, semantic_id=code, title=(ann.content or code)[:80], content=ann.content, decided_by=actor)
        db.add(d)
        db.flush()
        ensure_semantic_object(db, project_id=ann.project_id, semantic_id=code, object_type=m.SemanticObjectType.DECISION,
                               display_name=(ann.content or code)[:80], entity_ref=d.id, metadata={"provenance": provenance})
        create_trace_link(db, project_id=ann.project_id, source_semantic_id=code, target_semantic_id=ann.anchor_semantic_id, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
        result = {"kind": "decision", "code": code, "provenance": provenance}
    elif to_kind == "assumption":
        code = _next_code(db, ann.project_id, "ASM", m.Assumption)
        a = m.Assumption(project_id=ann.project_id, semantic_id=code, content=ann.content, created_by=actor)
        db.add(a)
        db.flush()
        ensure_semantic_object(db, project_id=ann.project_id, semantic_id=code, object_type=m.SemanticObjectType.ASSUMPTION,
                               display_name=code, entity_ref=a.id, metadata={"provenance": provenance})
        create_trace_link(db, project_id=ann.project_id, source_semantic_id=code, target_semantic_id=ann.anchor_semantic_id, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
        result = {"kind": "assumption", "code": code, "provenance": provenance}
    elif to_kind == "clarification":
        code = _next_code(db, ann.project_id, "CLR", m.Clarification)
        c = m.Clarification(project_id=ann.project_id, question=ann.content, asked_by=actor, resolved=False)
        db.add(c)
        db.flush()
        ensure_semantic_object(db, project_id=ann.project_id, semantic_id=code, object_type=m.SemanticObjectType.CLARIFICATION,
                               display_name=code, entity_ref=c.id, metadata={"provenance": provenance})
        create_trace_link(db, project_id=ann.project_id, source_semantic_id=code, target_semantic_id=ann.anchor_semantic_id, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
        result = {"kind": "clarification", "code": code, "provenance": provenance}
    else:
        raise DomainError(f"Unknown promotion target '{to_kind}'", status_code=422)

    ann.status = m.AnnotationStatus.RESOLVED
    db.commit()
    return result


def record_actor(db: Session, actor_id: str, display_name: str, tenant_id: str | None = None, source: str = "LOCAL") -> None:
    """Cache a resolved actor identity (idempotent upsert)."""
    existing = db.get(m.ActorIdentity, actor_id)
    if existing:
        existing.display_name = display_name
        if tenant_id:
            existing.tenant_id = tenant_id
        existing.source = source
        existing.resolved_at = m.utcnow()
    else:
        db.add(m.ActorIdentity(actor_id=actor_id, display_name=display_name, tenant_id=tenant_id, source=source))
    db.commit()


# ---------------------------------------------------------------------------
# Ecosystem event + outbox (durable, idempotent delivery)
# ---------------------------------------------------------------------------

ECOSYSTEM_EVENT_TYPES = {
    "REQUIREMENT_BASELINED", "DESIGN_BASELINED", "CHANGE_REQUEST_APPROVED",
    "DESIGN_CHANGED", "EXECUTION_REQUESTED", "QA_VALIDATION_REQUESTED",
    "QA_RESULT_RECEIVED", "RELEASE_LINKED",
}

OUTBOX_STATUSES = {"PENDING", "SENT", "ACKNOWLEDGED", "FAILED"}


def emit_event(
    db: Session,
    *,
    event_type: str,
    project_id: str,
    payload: dict | None = None,
    target_services: list[str] | None = None,
    correlation_id: str | None = None,
    source_object_id: str | None = None,
    source_revision: str | None = None,
    actor_id: str | None = None,
    tenant_id: str | None = None,
) -> m.EcosystemEvent:
    """Record an ecosystem event and enqueue durable delivery, atomically.

    The domain change and its outbox rows commit together; delivery never
    depends on a fire-and-forget HTTP call.
    """
    if event_type not in ECOSYSTEM_EVENT_TYPES:
        raise DomainError(f"Unknown event type '{event_type}'", status_code=422)
    correlation_id = correlation_id or m.new_id("corr")
    event = m.EcosystemEvent(
        event_type=event_type, project_id=project_id, tenant_id=tenant_id,
        source_service="document-again", source_object_id=source_object_id,
        source_revision=source_revision, actor_id=actor_id,
        payload_version="1.0", payload=payload, correlation_id=correlation_id,
    )
    db.add(event)
    db.flush()
    for ts in (target_services or []):
        db.add(m.OutboxEvent(event_id=event.id, target_service=ts, correlation_id=correlation_id))
    db.commit()
    return event


def due_outbox(db: Session, limit: int = 50) -> list[m.OutboxEvent]:
    now = m.utcnow()
    rows = db.execute(
        select(m.OutboxEvent).where(
            m.OutboxEvent.status.in_(["PENDING", "FAILED"]),
            (m.OutboxEvent.next_attempt_at.is_(None)) | (m.OutboxEvent.next_attempt_at <= now),
        ).order_by(m.OutboxEvent.created_at).limit(limit)
    ).scalars().all()
    return rows


def _backoff_seconds(attempt_count: int) -> int:
    return min(5 * (2 ** max(attempt_count, 1)), 300)


def mark_outbox_sent(db: Session, outbox_id: str, external_reference: str | None = None) -> m.OutboxEvent:
    out = get_or_404(db, m.OutboxEvent, outbox_id, "OutboxEvent")
    out.status = "SENT"
    out.delivered_at = m.utcnow()
    out.attempt_count += 1
    out.last_error = None
    out.external_reference = external_reference
    db.commit()
    return out


def mark_outbox_acknowledged(db: Session, outbox_id: str) -> m.OutboxEvent:
    out = get_or_404(db, m.OutboxEvent, outbox_id, "OutboxEvent")
    out.status = "ACKNOWLEDGED"
    out.delivered_at = m.utcnow()
    db.commit()
    return out


def mark_outbox_failed(db: Session, outbox_id: str, error: str) -> m.OutboxEvent:
    out = get_or_404(db, m.OutboxEvent, outbox_id, "OutboxEvent")
    out.status = "FAILED"
    out.attempt_count += 1
    out.last_error = error[:2000]
    out.next_attempt_at = m.utcnow() + timedelta(seconds=_backoff_seconds(out.attempt_count))
    db.commit()
    return out


def deliver_due_events(db: Session, deliver_fn, limit: int = 50) -> dict:
    """Dispatch due outbox events through a delivery callback.

    `deliver_fn(outbox_event) -> str | None` returns an external reference
    on success. Failures are recorded with a backoff for retry; the unique
    (event_id, target_service) constraint keeps delivery idempotent.
    """
    delivered = 0
    failed = 0
    for out in due_outbox(db, limit):
        try:
            ext_ref = deliver_fn(out)
            out.status = "SENT"
            out.delivered_at = m.utcnow()
            out.attempt_count += 1
            out.last_error = None
            if ext_ref:
                out.external_reference = ext_ref
            delivered += 1
        except Exception as exc:  # noqa: BLE001 — record and retry later
            out.status = "FAILED"
            out.attempt_count += 1
            out.last_error = str(exc)[:2000]
            out.next_attempt_at = m.utcnow() + timedelta(seconds=_backoff_seconds(out.attempt_count))
            failed += 1
    db.commit()
    return {"delivered": delivered, "failed": failed}


def list_ecosystem_events(db: Session, project_id: str | None = None, limit: int = 200) -> list[dict]:
    q = select(m.EcosystemEvent).order_by(m.EcosystemEvent.occurred_at.desc()).limit(limit)
    if project_id:
        q = q.where(m.EcosystemEvent.project_id == project_id)
    rows = db.execute(q).scalars().all()
    return [
        {
            "id": e.id, "event_type": e.event_type, "project_id": e.project_id,
            "source_service": e.source_service, "source_object_id": e.source_object_id,
            "source_revision": e.source_revision, "occurred_at": e.occurred_at.isoformat(),
            "actor_id": e.actor_id, "payload": e.payload, "correlation_id": e.correlation_id,
        }
        for e in rows
    ]


def list_outbox(db: Session, project_id: str | None = None, limit: int = 200) -> list[dict]:
    q = select(m.OutboxEvent).order_by(m.OutboxEvent.created_at.desc()).limit(limit)
    rows = db.execute(q).scalars().all()
    return [
        {
            "id": o.id, "event_id": o.event_id, "target_service": o.target_service,
            "status": o.status, "attempt_count": o.attempt_count, "last_error": o.last_error,
            "delivered_at": o.delivered_at.isoformat() if o.delivered_at else None,
            "external_reference": o.external_reference, "correlation_id": o.correlation_id,
        }
        for o in rows
    ]


# ---------------------------------------------------------------------------
# Ecosystem handoffs (PM / QA) and external references
# ---------------------------------------------------------------------------

EXECUTION_HANDOFF_STATUSES = {"DRAFT", "READY", "SENT", "ACKNOWLEDGED", "FAILED", "CANCELLED"}
QA_HANDOFF_STATUSES = {"DRAFT", "READY", "SENT", "ACKNOWLEDGED", "FAILED", "CANCELLED"}
EXTERNAL_RELATION_TYPES = {"IMPLEMENTED_BY", "VALIDATED_BY", "TRACKED_BY", "RELEASED_IN", "HANDED_OFF_TO"}


def create_execution_handoff(
    db: Session,
    project_id: str,
    baseline_id: str | None = None,
    source_revision_id: str | None = None,
    change_request_id: str | None = None,
    target_service: str = "pm-again",
    actor: str = "local-user",
    actor_id: str | None = None,
    status: str = "DRAFT",
) -> dict:
    """Snapshot the exact baseline context and emit a durable PM handoff.

    The payload holds only immutable references so that PM Again can fetch
    the authoritative design by id; Document Again never hands out a copy of
    mutable execution state.
    """
    if status not in EXECUTION_HANDOFF_STATUSES:
        raise ValueError(f"invalid execution handoff status: {status}")
    baseline_id = baseline_id or _latest_baseline_id(db, project_id)
    source_revision_id = source_revision_id or _latest_design_revision_id(db, project_id)
    payload = {
        "baselineId": baseline_id,
        "sourceRevisionId": source_revision_id,
        "changeRequestId": change_request_id,
        "projectId": project_id,
    }
    correlation_id = f"pm:{project_id}:{baseline_id or source_revision_id or 'adhoc'}"
    handoff = m.ExecutionHandoff(
        project_id=project_id,
        baseline_id=baseline_id,
        source_revision_id=source_revision_id,
        change_request_id=change_request_id,
        target_service=target_service,
        status=status,
        payload_snapshot=payload,
        correlation_id=correlation_id,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(handoff)
    db.flush()
    emit_event(
        db,
        event_type="EXECUTION_REQUESTED",
        project_id=project_id,
        tenant_id=None,
        source_object_id=handoff.id,
        source_revision=source_revision_id,
        actor_id=actor_id,
        payload=payload,
        correlation_id=correlation_id,
        target_services=[target_service],
    )
    db.commit()
    db.refresh(handoff)
    return _execution_handoff_dict(handoff)


def create_qa_validation_handoff(
    db: Session,
    project_id: str,
    baseline_id: str | None = None,
    requirement_ids: list[str] | None = None,
    semantic_object_ids: list[str] | None = None,
    design_revision_ids: list[str] | None = None,
    target_release: str | None = None,
    target_service: str = "qa-again",
    actor: str = "local-user",
    actor_id: str | None = None,
    status: str = "DRAFT",
) -> dict:
    if status not in QA_HANDOFF_STATUSES:
        raise ValueError(f"invalid qa handoff status: {status}")
    baseline_id = baseline_id or _latest_baseline_id(db, project_id)
    payload = {
        "baselineId": baseline_id,
        "requirementIds": requirement_ids or [],
        "semanticObjectIds": semantic_object_ids or [],
        "designRevisionIds": design_revision_ids or [],
        "targetRelease": target_release,
        "projectId": project_id,
    }
    correlation_id = f"qa:{project_id}:{baseline_id or 'adhoc'}"
    handoff = m.QAValidationHandoff(
        project_id=project_id,
        baseline_id=baseline_id,
        requirement_ids=requirement_ids or [],
        semantic_object_ids=semantic_object_ids or [],
        design_revision_ids=design_revision_ids or [],
        target_release=target_release,
        target_service=target_service,
        status=status,
        payload_snapshot=payload,
        correlation_id=correlation_id,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(handoff)
    db.flush()
    emit_event(
        db,
        event_type="QA_VALIDATION_REQUESTED",
        project_id=project_id,
        tenant_id=None,
        source_object_id=handoff.id,
        source_revision=None,
        actor_id=actor_id,
        payload=payload,
        correlation_id=correlation_id,
        target_services=[target_service],
    )
    db.commit()
    db.refresh(handoff)
    return _qa_handoff_dict(handoff)


def mark_handoff_status(
    db: Session, handoff_id: str, kind: str, status: str, external_reference: str | None = None
) -> dict:
    model = m.ExecutionHandoff if kind == "execution" else m.QAValidationHandoff
    valid = EXECUTION_HANDOFF_STATUSES if kind == "execution" else QA_HANDOFF_STATUSES
    if status not in valid:
        raise ValueError(f"invalid {kind} handoff status: {status}")
    row = db.get(model, handoff_id)
    if row is None:
        raise KeyError(f"{kind} handoff not found: {handoff_id}")
    row.status = status
    if external_reference:
        row.external_reference = external_reference
    if status in {"SENT", "ACKNOWLEDGED"} and row.delivered_at is None:
        row.delivered_at = m.utcnow()
    db.commit()
    db.refresh(row)
    return _execution_handoff_dict(row) if kind == "execution" else _qa_handoff_dict(row)


def list_execution_handoffs(db: Session, project_id: str | None = None) -> list[dict]:
    q = select(m.ExecutionHandoff).order_by(m.ExecutionHandoff.created_at.desc())
    if project_id:
        q = q.where(m.ExecutionHandoff.project_id == project_id)
    return [_execution_handoff_dict(h) for h in db.execute(q).scalars().all()]


def list_qa_handoffs(db: Session, project_id: str | None = None) -> list[dict]:
    q = select(m.QAValidationHandoff).order_by(m.QAValidationHandoff.created_at.desc())
    if project_id:
        q = q.where(m.QAValidationHandoff.project_id == project_id)
    return [_qa_handoff_dict(h) for h in db.execute(q).scalars().all()]


def _execution_handoff_dict(h: m.ExecutionHandoff) -> dict:
    return {
        "id": h.id, "project_id": h.project_id, "baseline_id": h.baseline_id,
        "source_revision_id": h.source_revision_id, "change_request_id": h.change_request_id,
        "target_service": h.target_service, "status": h.status,
        "external_reference": h.external_reference, "payload_snapshot": h.payload_snapshot,
        "correlation_id": h.correlation_id, "created_at": h.created_at.isoformat(),
        "created_by": h.created_by, "actor_id": h.actor_id,
        "delivered_at": h.delivered_at.isoformat() if h.delivered_at else None,
    }


def _qa_handoff_dict(h: m.QAValidationHandoff) -> dict:
    return {
        "id": h.id, "project_id": h.project_id, "baseline_id": h.baseline_id,
        "requirement_ids": h.requirement_ids, "semantic_object_ids": h.semantic_object_ids,
        "design_revision_ids": h.design_revision_ids, "target_release": h.target_release,
        "target_service": h.target_service, "status": h.status,
        "external_reference": h.external_reference, "payload_snapshot": h.payload_snapshot,
        "correlation_id": h.correlation_id, "created_at": h.created_at.isoformat(),
        "created_by": h.created_by, "actor_id": h.actor_id,
        "delivered_at": h.delivered_at.isoformat() if h.delivered_at else None,
    }


def _latest_baseline_id(db: Session, project_id: str) -> str | None:
    row = db.execute(
        select(m.Baseline.id)
        .where(m.Baseline.project_id == project_id)
        .order_by(m.Baseline.created_at.desc())
        .limit(1)
    ).scalar_one_or_none()
    return row


def _latest_design_revision_id(db: Session, project_id: str) -> str | None:
    row = db.execute(
        select(m.ArtifactRevision.id)
        .join(m.Artifact, m.Artifact.id == m.ArtifactRevision.artifact_id)
        .where(m.Artifact.project_id == project_id)
        .order_by(m.ArtifactRevision.created_at.desc())
        .limit(1)
    ).scalar_one_or_none()
    return row


def create_external_reference(
    db: Session,
    project_id: str,
    semantic_id: str,
    service: str,
    external_id: str,
    relation_type: str = "TRACKED_BY",
    object_type: str | None = None,
    url: str | None = None,
    metadata: dict | None = None,
) -> dict:
    if relation_type not in EXTERNAL_RELATION_TYPES:
        raise ValueError(f"invalid external relation type: {relation_type}")
    existing = db.execute(
        select(m.ExternalReference).where(
            m.ExternalReference.project_id == project_id,
            m.ExternalReference.service == service,
            m.ExternalReference.external_id == external_id,
        )
    ).scalar_one_or_none()
    if existing:
        existing.relation_type = relation_type
        existing.semantic_id = semantic_id
        existing.object_type = object_type or existing.object_type
        existing.url = url or existing.url
        existing.metadata_json = metadata or existing.metadata_json
        db.commit()
        db.refresh(existing)
        return _external_reference_dict(existing)
    ref = m.ExternalReference(
        project_id=project_id,
        semantic_id=semantic_id,
        service=service,
        external_id=external_id,
        relation_type=relation_type,
        object_type=object_type,
        url=url,
        metadata_json=metadata,
    )
    db.add(ref)
    db.commit()
    db.refresh(ref)
    return _external_reference_dict(ref)


def list_external_references(
    db: Session, project_id: str | None = None, semantic_id: str | None = None
) -> list[dict]:
    q = select(m.ExternalReference).order_by(m.ExternalReference.created_at.desc())
    if project_id:
        q = q.where(m.ExternalReference.project_id == project_id)
    if semantic_id:
        q = q.where(m.ExternalReference.semantic_id == semantic_id)
    return [_external_reference_dict(r) for r in db.execute(q).scalars().all()]


def _external_reference_dict(r: m.ExternalReference) -> dict:
    return {
        "id": r.id, "project_id": r.project_id, "semantic_id": r.semantic_id,
        "relation_type": r.relation_type, "service": r.service, "external_id": r.external_id,
        "object_type": r.object_type, "url": r.url, "metadata": r.metadata_json,
        "created_at": r.created_at.isoformat(),
    }


# ---------------------------------------------------------------------------
# Reproducible export (from versioned structured state, never live state)
# ---------------------------------------------------------------------------


def _safe_filename(name: str) -> str:
    """ASCII-only filename for Content-Disposition headers."""
    keep = []
    for ch in name:
        if ch.isascii() and ch.isalnum() or ch in ".-_":
            keep.append(ch)
        else:
            keep.append("_")
    return "".join(keep)


def export_metadata(db: Session, revision_id: str) -> dict:
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    baseline = db.execute(
        select(m.Baseline).join(m.BaselineBinding, m.BaselineBinding.baseline_id == m.Baseline.id)
        .where(m.BaselineBinding.artifact_revision_id == revision_id)
    ).scalars().first()
    return {
        "project": revision.artifact.project.name,
        "project_key": revision.artifact.project.key,
        "artifact_type": revision.artifact.type.value,
        "artifact_title": revision.artifact.title,
        "revision_number": revision.revision_number,
        "status": revision.status.value,
        "confirmed_by": revision.confirmed_by,
        "confirmed_at": revision.confirmed_at.isoformat() if revision.confirmed_at else None,
        "generated_at": m.utcnow().isoformat(),
        "baseline_id": baseline.id if baseline else None,
        "baseline_name": baseline.name if baseline else None,
    }


def _flatten_section(section: dict) -> list[dict]:
    """Flatten a Tiptap doc JSON section into export-friendly blocks."""
    blocks: list[dict] = []

    def text_of(node) -> str:
        if node.get("type") == "text":
            return node.get("text") or ""
        return "".join(text_of(c) for c in node.get("content") or [])

    def walk(node) -> None:
        t = node.get("type")
        if t == "text":
            return
        children = node.get("content") or []
        if t == "heading":
            blocks.append({"kind": "heading", "level": node.get("attrs", {}).get("level", 2), "text": "".join(text_of(c) for c in children)})
            return
        if t == "paragraph":
            blocks.append({"kind": "paragraph", "text": "".join(text_of(c) for c in children)})
            return
        if t in ("bulletList", "orderedList", "taskList"):
            for item in children:
                blocks.append({"kind": "list_item", "ordered": t == "orderedList", "text": "".join(text_of(c) for c in item.get("content") or [])})
            return
        if t == "codeBlock":
            blocks.append({"kind": "code", "text": "".join(text_of(c) for c in children)})
            return
        if t == "table":
            rows = []
            for row in children:
                rows.append(["".join(text_of(cell) for cell in (row.get("content") or []))])
            blocks.append({"kind": "table", "rows": rows})
            return
        for child in children:
            walk(child)

    walk(section.get("content") or {"type": "doc", "content": []})
    return blocks


def _data_dictionary_from_snapshot(snapshot: dict) -> list[dict]:
    """Data dictionary rows derived from a frozen technical-design snapshot."""
    rows = []
    for schema_sid, schema in (snapshot.get("db_schemas") or {}).items():
        for table_sid, table in (schema.get("tables") or {}).items():
            for field_sid, f in (table.get("fields") or {}).items():
                rows.append({
                    "table": table.get("name"), "table_semantic_id": table_sid,
                    "field": f.get("name"), "field_semantic_id": field_sid,
                    "data_type": f.get("data_type"), "length": f.get("length"),
                    "nullable": f.get("nullable"), "default": f.get("default"),
                    "primary_key": f.get("primary_key"), "foreign_key": f.get("foreign_key"),
                    "reference": f.get("reference"), "description": f.get("description"),
                    "remark": f.get("remark"),
                })
    return rows


def export_revision(db: Session, revision_id: str, format: str) -> tuple[bytes, str, str]:
    """Export a revision from its frozen snapshot (historical correctness)."""
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    snapshot = revision.snapshot or {}
    meta = export_metadata(db, revision_id)
    sections = (snapshot.get("sections") or [])

    if format == "json":
        payload = {
            "metadata": meta,
            "sections": sections,
            "technical_design": snapshot.get("technical_design"),
            "database": snapshot.get("database"),
        }
        return _json.dumps(payload, indent=2, default=str).encode(), "application/json", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}.json"

    if format == "csv":
        # Historical: prefer the frozen technical-design snapshot.
        dd = _data_dictionary_from_snapshot(snapshot.get("technical_design") or {})
        if not dd:
            dd = _data_dictionary_from_snapshot({"db_schemas": snapshot.get("database", {}) or {}})
        buf = _io.StringIO()
        writer = _csv.DictWriter(buf, fieldnames=["table", "field", "data_type", "length", "nullable", "default", "primary_key", "foreign_key", "reference", "description", "remark", "field_semantic_id"])
        writer.writeheader()
        for row in dd:
            writer.writerow({k: row.get(k) for k in writer.fieldnames})
        return buf.getvalue().encode(), "text/csv", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}.csv"

    if format == "pdf":
        return _render_pdf(meta, sections), "application/pdf", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}.pdf"

    if format == "svg":
        return _render_erd_svg(snapshot), "image/svg+xml", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}.svg"

    raise DomainError(f"Unknown export format '{format}'", status_code=422)


def _render_pdf(meta: dict, sections: list[dict]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm)
    styles = getSampleStyleSheet()
    styles["Title"].fontName = "Helvetica-Bold"
    story = [
        Paragraph(meta["artifact_title"], styles["Title"]),
        Paragraph(
            f"{meta['artifact_type']} · revision r{meta['revision_number']} · {meta['status']} · "
            f"confirmed by {meta.get('confirmed_by') or '—'} at {meta.get('confirmed_at') or '—'}<br/>"
            f"project {meta['project']} · generated {meta['generated_at']} · baseline {meta.get('baseline_name') or '—'}",
            styles["Normal"],
        ),
        Spacer(1, 6 * mm),
    ]
    for sec in sections:
        story.append(Paragraph(sec.get("heading") or "Untitled", ParagraphStyle("sh", parent=styles["Heading2"], fontSize=13, spaceBefore=8, spaceAfter=4)))
        story.append(Paragraph(f'<font size="7" color="#666">{sec.get("id")}</font>', styles["Normal"]))
        for blk in _flatten_section(sec):
            kind = blk["kind"]
            if kind == "heading":
                story.append(Paragraph(blk["text"], ParagraphStyle("h", parent=styles["Heading3"], fontSize=11, spaceBefore=4, spaceAfter=2)))
            elif kind == "paragraph":
                story.append(Paragraph(blk["text"].replace("\n", "<br/>"), styles["Normal"]))
            elif kind == "list_item":
                story.append(Paragraph(("• " if not blk.get("ordered") else "1. ") + blk["text"], styles["Normal"]))
            elif kind == "code":
                story.append(Preformatted(blk["text"], ParagraphStyle("code", fontName="Courier", fontSize=8, leading=10)))
            elif kind == "table":
                rows = blk["rows"]
                if rows:
                    t = Table([[Paragraph(c, styles["Normal"]) for c in row] for row in rows])
                    t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, "#999"), ("FONTSIZE", (0, 0), (-1, -1), 8)]))
                    story.append(t)
        story.append(Spacer(1, 3 * mm))
    doc.build(story)
    return buf.getvalue()


def _render_erd_svg(snapshot: dict) -> bytes:
    """Simple ERD SVG from a frozen technical-design snapshot."""
    design = snapshot.get("technical_design") or {}
    schemas = design.get("db_schemas") or {}
    tables: list[tuple[str, str, list]] = []  # (name, sid, [(field, type, pk, fk)])
    for schema in schemas.values():
        for tid, t in (schema.get("tables") or {}).items():
            fields = [(f.get("name"), f.get("data_type"), f.get("primary_key"), f.get("foreign_key")) for f in (t.get("fields") or {}).values()]
            tables.append((t.get("name"), tid, fields))
    parts = ['<svg xmlns="http://www.w3.org/2000/svg" width="800" height="600" viewBox="0 0 800 600">']
    x = 40
    y = 40
    for i, (name, sid, fields) in enumerate(tables):
        h = 30 + len(fields) * 18
        if y + h > 560:
            y = 40
            x += 240
        parts.append(f'<rect x="{x}" y="{y}" width="200" height="{h}" rx="6" fill="#161a23" stroke="#6366f1"/>')
        parts.append(f'<text x="{x+8}" y="{y+18}" font-size="13" font-weight="bold" fill="#e2e8f0">{name}</text>')
        for j, (fname, ftype, pk, fk) in enumerate(fields):
            fy = y + 34 + j * 18
            mark = "🔑" if pk else ("🔗" if fk else "")
            parts.append(f'<text x="{x+8}" y="{fy}" font-size="11" fill="#94a3b8">{mark} {fname}: {ftype}</text>')
        y += h + 24
    parts.append("</svg>")
    return "".join(parts).encode()


def export_design_package(db: Session, baseline_id: str) -> bytes:
    """ZIP of export artifacts, all generated from the same baseline context."""
    baseline = get_or_404(db, m.Baseline, baseline_id, "Baseline")
    bindings = db.execute(
        select(m.BaselineBinding).where(m.BaselineBinding.baseline_id == baseline_id)
    ).scalars().all()
    if not bindings:
        raise DomainError("Baseline has no bindings")

    buf = _io.BytesIO()
    with _zipfile.ZipFile(buf, "w", _zipfile.ZIP_DEFLATED) as z:
        meta_rows = []
        for b in bindings:
            rev = db.get(m.ArtifactRevision, b.artifact_revision_id)
            if rev is None:
                continue
            meta = export_metadata(db, rev.id)
            meta_rows.append(meta)
            if rev.artifact.type in (m.ArtifactType.UR, m.ArtifactType.DR):
                pdf, _, name = export_revision(db, rev.id, "pdf")
                z.writestr(name, pdf)
            json_bytes, _, name = export_revision(db, rev.id, "json")
            z.writestr(f"revision-{rev.revision_number}-{rev.artifact.type.value}.json", json_bytes)
            csv_bytes, _, _ = export_revision(db, rev.id, "csv")
            if csv_bytes.strip():
                z.writestr(f"data-dictionary-{rev.artifact.type.value}-r{rev.revision_number}.csv", csv_bytes)
        z.writestr("manifest.json", _json.dumps({"baseline": baseline.name, "baseline_id": baseline.id, "revisions": meta_rows}, indent=2, default=str))
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Export V2 — XLSX / DOCX / traceability matrix / design package directory
# ---------------------------------------------------------------------------


def _traceability_rows(db: Session, project_id: str) -> list[dict]:
    """Traceability matrix rows from current trace links (labelled as live)."""
    types = {
        so.semantic_id: so.object_type.value
        for so in db.execute(
            select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
        ).scalars().all()
    }
    links = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
        .order_by(m.TraceLink.source_semantic_id)
    ).scalars().all()
    return [
        {
            "source": l.source_semantic_id, "source_type": types.get(l.source_semantic_id),
            "relation": l.relation_type.value, "target": l.target_semantic_id,
            "target_type": types.get(l.target_semantic_id),
            "revision_context": l.revision_context,
        }
        for l in links
    ]


def _render_xlsx(meta: dict, sections: list[dict], dd_rows: list[dict], trace_rows: list[dict]) -> bytes:
    import openpyxl
    from openpyxl.styles import Font

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Metadata"
    for i, (k, v) in enumerate(meta.items(), start=1):
        ws.cell(row=i, column=1, value=k).font = Font(bold=True)
        ws.cell(row=i, column=2, value=v)

    if sections:
        ws2 = wb.create_sheet("Document")
        r = 1
        for sec in sections:
            ws2.cell(row=r, column=1, value=sec.get("id") or "").font = Font(bold=True)
            ws2.cell(row=r, column=2, value=sec.get("heading") or "").font = Font(bold=True)
            r += 1
            for blk in _flatten_section(sec):
                ws2.cell(row=r, column=2, value=blk["text"])
                r += 1

    if dd_rows:
        ws3 = wb.create_sheet("Data Dictionary")
        headers = ["table", "field", "data_type", "length", "nullable", "primary_key",
                   "foreign_key", "reference", "description", "remark", "field_semantic_id"]
        ws3.append(headers)
        for row in dd_rows:
            ws3.append([row.get(h) for h in headers])

    if trace_rows:
        ws4 = wb.create_sheet("Traceability")
        ws4.append(["source", "source_type", "relation", "target", "target_type", "revision_context"])
        for row in trace_rows:
            ws4.append([row["source"], row["source_type"], row["relation"],
                        row["target"], row["target_type"], row["revision_context"]])

    buf = _io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _render_docx(meta: dict, sections: list[dict]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(meta["artifact_title"], level=0)
    doc.add_paragraph(
        f"{meta['artifact_type']} · revision r{meta['revision_number']} · {meta['status']} · "
        f"confirmed by {meta.get('confirmed_by') or '—'} at {meta.get('confirmed_at') or '—'} · "
        f"project {meta['project']} · generated {meta['generated_at']}"
    )
    for sec in sections:
        doc.add_heading(sec.get("heading") or "Untitled", level=1)
        for blk in _flatten_section(sec):
            kind = blk["kind"]
            if kind == "heading":
                doc.add_heading(blk["text"], level=2)
            elif kind == "paragraph":
                doc.add_paragraph(blk["text"])
            elif kind == "list_item":
                doc.add_paragraph(blk["text"], style="List Bullet")
            elif kind == "code":
                doc.add_paragraph(blk["text"], style="No Spacing")
            elif kind == "table" and blk["rows"]:
                table = doc.add_table(rows=len(blk["rows"]), cols=len(blk["rows"][0]))
                for i, row in enumerate(blk["rows"]):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_revision_v2(db: Session, revision_id: str, format: str) -> tuple[bytes, str, str]:
    """Export V2: xlsx / docx in addition to the V1 formats (reuse)."""
    if format in ("xlsx", "docx"):
        revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
        snapshot = revision.snapshot or {}
        meta = export_metadata(db, revision_id)
        sections = snapshot.get("sections") or []
        dd = _data_dictionary_from_snapshot(snapshot.get("technical_design") or {})
        if not dd:
            dd = _data_dictionary_from_snapshot({"db_schemas": snapshot.get("database", {}) or {}})
        base = f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}"
        if format == "xlsx":
            trace = _traceability_rows(db, revision.artifact.project_id)
            return _render_xlsx(meta, sections, dd, trace), \
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"{base}.xlsx"
        return _render_docx(meta, sections), \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{base}.docx"
    return export_revision(db, revision_id, format)


def export_design_package_v2(db: Session, baseline_id: str) -> bytes:
    """ZIP export with a clean directory structure (design package V2)."""
    baseline = get_or_404(db, m.Baseline, baseline_id, "Baseline")
    bindings = db.execute(
        select(m.BaselineBinding).where(m.BaselineBinding.baseline_id == baseline_id)
    ).scalars().all()
    if not bindings:
        raise DomainError("Baseline has no bindings")

    buf = _io.BytesIO()
    with _zipfile.ZipFile(buf, "w", _zipfile.ZIP_DEFLATED) as z:
        meta_rows = []
        dd_rows = []
        trace_rows = _traceability_rows(db, baseline.project_id)
        for b in bindings:
            rev = db.get(m.ArtifactRevision, b.artifact_revision_id)
            if rev is None:
                continue
            meta = export_metadata(db, rev.id)
            meta_rows.append(meta)
            snapshot = rev.snapshot or {}
            dd_rows.extend(_data_dictionary_from_snapshot(snapshot.get("technical_design") or {}))
            json_bytes, _, _ = export_revision(db, rev.id, "json")
            z.writestr(f"revisions/{rev.revision_number}-{rev.artifact.type.value}.json", json_bytes)
            if rev.artifact.type in (m.ArtifactType.UR, m.ArtifactType.DR):
                pdf, _, _ = export_revision(db, rev.id, "pdf")
                z.writestr(f"documents/{_safe_filename(meta['artifact_title'])}-r{rev.revision_number}.pdf", pdf)
                docx, _, _ = export_revision_v2(db, rev.id, "docx")
                z.writestr(f"documents/{_safe_filename(meta['artifact_title'])}-r{rev.revision_number}.docx", docx)
            # OpenAPI export from API_DESIGN revisions
            if rev.artifact.type == m.ArtifactType.API_DESIGN:
                api_map = (snapshot.get("technical_design") or {}).get("api_endpoints") or {}
                if api_map:
                    z.writestr(f"api/{_safe_filename(meta['artifact_title'])}.openapi.json",
                               _json.dumps(export_openapi(db, rev.id), indent=2, default=str))
        if dd_rows:
            z.writestr("data-dictionary.xlsx", _render_xlsx(
                export_metadata(db, bindings[0].artifact_revision_id), [], dd_rows, []))
        z.writestr("traceability.xlsx", _render_xlsx(
            export_metadata(db, bindings[0].artifact_revision_id), [], [], trace_rows))
        z.writestr("manifest.json", _json.dumps(
            {"baseline": baseline.name, "baseline_id": baseline.id,
             "directory_structure": ["manifest.json", "documents/", "revisions/",
                                     "api/", "data-dictionary.xlsx", "traceability.xlsx"],
             "revisions": meta_rows}, indent=2, default=str))
    return buf.getvalue()
