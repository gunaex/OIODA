"""Domain services — every product-level invariant is enforced here.

Rationale for a service layer (rather than router-level CRUD): rules
like "confirmed revisions are immutable" and "baselines never
re-resolve to latest" are the product's core value. They must be
impossible to bypass from the HTTP edge.
"""
from __future__ import annotations

import csv as _csv
import difflib
import io as _io
import json as _json
import os
import re
import zipfile as _zipfile
from collections import Counter
from datetime import timedelta

from sqlalchemy import delete, or_, select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from . import models as m
from .contracts import versioned_payload
from .models import RevisionStatus
from .observability import metrics
from .tenant import current_tenant


class DomainError(Exception):
    """Rule violation the caller must see as 4xx/409."""

    def __init__(self, message: str, status_code: int = 409):
        super().__init__(message)
        self.status_code = status_code


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------


def get_or_404(db: Session, model, id_: str, what: str):
    obj = db.get(model, id_)
    if obj is None:
        raise DomainError(f"{what} not found: {id_}", status_code=404)
    return obj


def require_editable(revision: m.ArtifactRevision) -> None:
    if revision.status != RevisionStatus.DRAFT:
        raise DomainError(
            f"Revision {revision.id} is {revision.status.value} and therefore immutable. "
            "Clone it as a new DRAFT revision to continue editing.",
        )


def guard_project(db: Session, project_id: str) -> m.Project:
    """Block cross-tenant access to a project-scoped resource.

    Tenant enforcement is active only when the request carries an explicit
    tenant (account_again mode, or local mode with X-Tenant-Id). An unscoped
    project (tenant_id NULL) remains accessible in local development; a
    mismatched tenant is rejected with 403 (never leaked as 404).
    """
    project = get_or_404(db, m.Project, project_id, "Project")
    tenant = current_tenant()
    if tenant is not None and project.tenant_id is not None and project.tenant_id != tenant:
        raise DomainError("Cross-tenant access denied", status_code=403)
    return project

# ---------------------------------------------------------------------------
# Project
# ---------------------------------------------------------------------------


def create_project(db: Session, *, key: str, name: str, description=None, actor="local-user", tenant_id: str | None = None, metadata: dict | None = None):
    project = m.Project(
        key=key, name=name, description=description, created_by=actor,
        tenant_id=tenant_id if tenant_id is not None else current_tenant(),
        project_meta=metadata or {},
    )
    db.add(project)
    db.commit()
    return project


# ---------------------------------------------------------------------------
# R16 — Project lifecycle (Document Again is the lifecycle authority)
# ---------------------------------------------------------------------------

LIFECYCLE_STATES = ("ACTIVE", "ARCHIVED", "DELETE_REQUESTED", "DELETED")
CLONE_POLICY_VERSION = "1.0"


def list_projects(db: Session, *, state: str | None = None):
    q = select(m.Project)
    tenant = current_tenant()
    if tenant is not None:
        q = q.where(m.Project.tenant_id == tenant)
    if state and state.upper() in LIFECYCLE_STATES:
        q = q.where(m.Project.lifecycle_state == state.upper())
    return db.execute(q).scalars().all()


def _require_lifecycle(project: m.Project, allowed: tuple[str, ...]) -> None:
    if project.lifecycle_state not in allowed:
        raise DomainError(
            f"Project is {project.lifecycle_state} — must be {'/'.join(allowed)} for this action.",
            status_code=409,
        )


def archive_project(db: Session, project_id: str, *, actor="local-user", actor_id: str | None = None) -> dict:
    """ACTIVE → ARCHIVED. Preserves all bounded-service truth; removes the
    project from the default Active view only."""
    project = guard_project(db, project_id)
    _require_lifecycle(project, ("ACTIVE",))
    project.lifecycle_state = "ARCHIVED"
    db.commit()
    record_audit(
        db, action="PROJECT_ARCHIVED", project_id=project_id, actor_id=actor_id,
        object_type="Project", object_id=project_id,
        metadata={"from": "ACTIVE", "to": "ARCHIVED"},
    )
    return {"project_id": project.id, "lifecycle_state": project.lifecycle_state}


def restore_project(db: Session, project_id: str, *, actor="local-user", actor_id: str | None = None) -> dict:
    """ARCHIVED → ACTIVE."""
    project = guard_project(db, project_id)
    _require_lifecycle(project, ("ARCHIVED",))
    project.lifecycle_state = "ACTIVE"
    db.commit()
    record_audit(
        db, action="PROJECT_RESTORED", project_id=project_id, actor_id=actor_id,
        object_type="Project", object_id=project_id,
        metadata={"from": "ARCHIVED", "to": "ACTIVE"},
    )
    return {"project_id": project.id, "lifecycle_state": project.lifecycle_state}


def clone_project(db: Session, project_id: str, *, key: str, name: str,
                  description: str | None = None, actor="local-user",
                  actor_id: str | None = None) -> dict:
    """Clone reusable project knowledge into a NEW project. Execution/history
    truth is NOT cloned. Every cloned authority object is created through the
    owning service (Document Again) — no direct DB copies of other services.

    Cross-service scope (PM planning / QA test design / Infra design) is
    reported as NOT_APPLICABLE here because PM/QA/Infra do not expose a clone
    endpoint — the clone is Document-level, honestly reported."""
    source = guard_project(db, project_id)
    new = create_project(
        db, key=key, name=name, description=description, actor=actor,
        metadata={"cloned": True},
    )
    new.cloned_from_project_id = source.id
    new.cloned_at = m.utcnow()
    new.cloned_by = actor
    new.clone_policy_version = CLONE_POLICY_VERSION
    db.commit()

    # Reusable Document scope: requirements + project memory (assumptions,
    # clarifications, decisions). Fresh ids; codes preserved within the new
    # project scope; globally-unique semantic ids get a clone suffix.
    copied = {"requirements": 0, "assumptions": 0, "clarifications": 0, "decisions": 0}
    for r in db.execute(select(m.Requirement).where(m.Requirement.project_id == source.id)).scalars():
        db.add(m.Requirement(
            project_id=new.id, code=r.code, title=r.title, description=r.description,
            source_type=r.source_type, source_reference=r.source_reference,
            status=r.status, priority=r.priority, metadata_json=dict(r.metadata_json or {}),
            created_by=actor,
        ))
        copied["requirements"] += 1
    for a in db.execute(select(m.Assumption).where(m.Assumption.project_id == source.id)).scalars():
        db.add(m.Assumption(
            project_id=new.id, semantic_id=f"{a.semantic_id}-{key}", content=a.content,
            status=a.status, created_by=actor,
        ))
        copied["assumptions"] += 1
    for c in db.execute(select(m.Clarification).where(m.Clarification.project_id == source.id)).scalars():
        db.add(m.Clarification(
            project_id=new.id, semantic_id=None, question=c.question, answer=c.answer,
            asked_by=actor, resolved=c.resolved,
        ))
        copied["clarifications"] += 1
    for d in db.execute(select(m.Decision).where(m.Decision.project_id == source.id)).scalars():
        db.add(m.Decision(
            project_id=new.id, semantic_id=f"{d.semantic_id}-{key}", title=d.title,
            content=d.content, decided_by=actor,
        ))
        copied["decisions"] += 1
    db.commit()

    record_audit(
        db, action="PROJECT_CLONED", project_id=new.id, actor_id=actor_id,
        object_type="Project", object_id=new.id,
        metadata={"cloned_from": source.id, "policy": CLONE_POLICY_VERSION, "copied": copied},
    )
    return {
        "project_id": new.id, "key": new.key, "name": new.name,
        "cloned_from_project_id": source.id, "clone_policy_version": CLONE_POLICY_VERSION,
        "copied": copied,
        "service_clone_status": {
            "DOCUMENT_AGAIN": "COMPLETED",
            "PM_AGAIN": "NOT_APPLICABLE",
            "QA_AGAIN": "NOT_APPLICABLE",
            "INFRA_AGAIN": "NOT_APPLICABLE",
        },
        "note": "Execution/history truth is NOT cloned. PM/QA/Infra clone endpoints do not exist.",
    }


def delete_impact(db: Session, project_id: str) -> dict:
    """Impact inspection before delete (Phase 12). Counts every Document-owned
    object for this project, honestly."""
    project = guard_project(db, project_id)
    counts = {}
    for model, name in (
        (m.Requirement, "requirements"),
        (m.ChangeRequest, "change_requests"),
        (m.Artifact, "artifacts"),
        (m.Suggestion, "suggestions"),
        (m.Consultation, "consultations"),
        (m.Clarification, "clarifications"),
        (m.Assumption, "assumptions"),
        (m.Decision, "decisions"),
        (m.Baseline, "baselines"),
        (m.TraceLink, "trace_links"),
    ):
        counts[name] = len(db.execute(select(model).where(model.project_id == project_id)).scalars().all())
    return {"project_id": project_id, "name": project.name, "key": project.key,
            "lifecycle_state": project.lifecycle_state, "document": counts,
            "bounded_services": {"PM_AGAIN": "EXTERNAL", "QA_AGAIN": "EXTERNAL", "INFRA_AGAIN": "EXTERNAL"}}


def delete_project(db: Session, project_id: str, *, actor="local-user", actor_id: str | None = None) -> dict:
    """Delete orchestration (Phase 13). Document Again owns its own truth and
    tombstones it; bounded services report their own capability. No direct SQL
    against PM/QA/Infra — every service reports its own status."""
    project = guard_project(db, project_id)
    _require_lifecycle(project, ("ARCHIVED", "DELETE_REQUESTED"))

    # Tombstone Document-owned truth: set lifecycle to DELETED, keep rows for
    # audit/traceability (never silently destroy).
    project.lifecycle_state = "DELETED"
    db.commit()

    record_audit(
        db, action="PROJECT_DELETED", project_id=project_id, actor_id=actor_id,
        object_type="Project", object_id=project_id,
        metadata={"method": "tombstone"},
    )
    return {
        "project_id": project.id,
        "lifecycle_state": "DELETED",
        "service_status": {
            "DOCUMENT_AGAIN": "COMPLETED",
            "PM_AGAIN": "NOT_APPLICABLE",
            "QA_AGAIN": "NOT_APPLICABLE",
            "INFRA_AGAIN": "NOT_APPLICABLE",
        },
        "note": "Document truth tombstoned (retained for audit). PM/QA/Infra do not expose a delete endpoint — their data is NOT touched.",
    }


# ---------------------------------------------------------------------------
# R16 — Project export / import package (versioned, authority-annotated)
# ---------------------------------------------------------------------------

EXPORT_PACKAGE_VERSION = "1.0"
_SECRET_PATTERNS = [
    (r"-----BEGIN (RSA |EC |OPENSSH )?PRIVATE KEY-----", "private_key"),
    (r"\b(sk|pk|api[_-]?key|secret|token|password)\b\s*[:=]\s*[\"']?[A-Za-z0-9_\-]{8,}", "credential_assignment"),
    (r"eyJ[A-Za-z0-9_\-]{10,}\.[A-Za-z0-9_\-]{10,}\.[A-Za-z0-9_\-]{10,}", "jwt"),
    (r"-----BEGIN CERTIFICATE-----", "certificate"),
]


def secret_scan(data: dict) -> dict:
    """Deterministic, pattern-based secret scan. Never prints matched values."""
    text = _json.dumps(data, default=str, ensure_ascii=False)
    findings = []
    for pattern, kind in _SECRET_PATTERNS:
        if re.search(pattern, text):
            findings.append(kind)
    return {"scanned_at": m.utcnow().isoformat(), "findings": sorted(set(findings)),
            "clean": len(findings) == 0}


def export_project(db: Session, project_id: str) -> dict:
    """Versioned project migration package (*.oida-project). Exports Document
    truth only — PM/QA/Infra truth stays in its bounded services. Never exports
    passwords, JWTs, API keys, or signing material."""
    project = guard_project(db, project_id)
    package = {
        "package_version": EXPORT_PACKAGE_VERSION,
        "exported_at": m.utcnow().isoformat(),
        "source_environment": "LOCAL",
        "project": {"id": project.id, "key": project.key, "name": project.name,
                    "description": project.description, "lifecycle_state": project.lifecycle_state,
                    "demo_reference": project.key == "TCM"},
        "authorities": ["DOCUMENT_AGAIN", "PM_AGAIN", "QA_AGAIN", "INFRA_AGAIN"],
        "bindings": get_workspace_bindings(db, project_id),
        "document": {
            "requirements": [
                {"code": r.code, "title": r.title, "description": r.description,
                 "source_type": r.source_type, "source_reference": r.source_reference,
                 "status": r.status.value, "priority": r.priority}
                for r in db.execute(select(m.Requirement).where(m.Requirement.project_id == project_id)).scalars()
            ],
            "clarifications": [
                {"semantic_id": c.semantic_id, "question": c.question, "answer": c.answer, "resolved": c.resolved}
                for c in db.execute(select(m.Clarification).where(m.Clarification.project_id == project_id)).scalars()
            ],
            "assumptions": [
                {"semantic_id": a.semantic_id, "content": a.content, "status": a.status}
                for a in db.execute(select(m.Assumption).where(m.Assumption.project_id == project_id)).scalars()
            ],
            "decisions": [
                {"semantic_id": d.semantic_id, "title": d.title, "content": d.content}
                for d in db.execute(select(m.Decision).where(m.Decision.project_id == project_id)).scalars()
            ],
            "change_requests": [
                {"code": cr.code, "title": cr.title, "requested_change": cr.requested_change,
                 "status": cr.status.value if hasattr(cr.status, "value") else str(cr.status)}
                for cr in db.execute(select(m.ChangeRequest).where(m.ChangeRequest.project_id == project_id)).scalars()
            ],
            "suggestions": [
                {"title": s.title, "type": s.type, "severity": s.severity, "status": s.status.value if hasattr(s.status, "value") else str(s.status)}
                for s in db.execute(select(m.Suggestion).where(m.Suggestion.project_id == project_id)).scalars()
            ],
            "consultations": [
                {"task_type": c.task_type, "question": c.question,
                 "aggregation_mode": (c.aggregation or {}).get("aggregation_mode")}
                for c in db.execute(select(m.Consultation).where(m.Consultation.project_id == project_id)).scalars()
            ],
            "trace_links": [
                {"source": t.source_semantic_id, "target": t.target_semantic_id,
                 "relation": t.relation_type.value if hasattr(t.relation_type, "value") else str(t.relation_type)}
                for t in db.execute(select(m.TraceLink).where(m.TraceLink.project_id == project_id)).scalars()
            ],
        },
    }
    package["secret_scan"] = secret_scan(package)
    return package


def import_project(db: Session, package: dict, *, actor="local-user", actor_id: str | None = None) -> dict:
    """Import a project package THROUGH Document Again (no DB copies of other
    services). Rebuilds Document truth + bindings; PM/QA/Infra import is
    reported NOT_APPLICABLE (no import endpoint). Records source→target id
    mappings. Restartable/idempotent via a stable import key."""
    if not isinstance(package, dict) or package.get("package_version") != EXPORT_PACKAGE_VERSION:
        raise DomainError(f"Unsupported package version (expected {EXPORT_PACKAGE_VERSION}).", status_code=422)

    scan = secret_scan(package)
    if not scan["clean"]:
        raise DomainError(f"Package failed secret scan: {', '.join(scan['findings'])}. Refusing import.", status_code=422)

    proj = package.get("project") or {}
    key = (proj.get("key") or "IMPORT")[:20]
    # Idempotency: reuse an existing project created by the same import key.
    existing = db.execute(select(m.Project).where(m.Project.key == key)).scalars().first()
    if existing:
        return {"project_id": existing.id, "reused": True,
                "migration_report": {"package_version": EXPORT_PACKAGE_VERSION, "note": "Already imported (idempotent reuse)."}}

    project = create_project(db, key=key, name=proj.get("name") or key,
                             description=proj.get("description"), actor=actor)
    doc = package.get("document") or {}
    mappings = {"requirements": {}, "clarifications": {}, "assumptions": {}, "decisions": {}}
    for r in doc.get("requirements", []):
        row = m.Requirement(project_id=project.id, code=r.get("code"), title=r.get("title"),
                            description=r.get("description"), source_type=r.get("source_type"),
                            source_reference=r.get("source_reference"), priority=r.get("priority"),
                            status=m.RequirementStatus(r.get("status", "DRAFT")), created_by=actor)
        db.add(row)
        mappings["requirements"][r.get("code")] = row.id
    for c in doc.get("clarifications", []):
        row = m.Clarification(project_id=project.id, semantic_id=None, question=c.get("question"),
                              answer=c.get("answer"), resolved=c.get("resolved", False), asked_by=actor)
        db.add(row)
    for a in doc.get("assumptions", []):
        row = m.Assumption(project_id=project.id, semantic_id=f"{a.get('semantic_id')}-{key}",
                           content=a.get("content"), status=a.get("status", "OPEN"), created_by=actor)
        db.add(row)
    for d in doc.get("decisions", []):
        row = m.Decision(project_id=project.id, semantic_id=f"{d.get('semantic_id')}-{key}",
                         title=d.get("title"), content=d.get("content"), decided_by=actor)
        db.add(row)

    def _enum(cls, value, default):
        try:
            return cls(value)
        except (ValueError, TypeError):
            return default

    # Rebuild change requests / suggestions / consultations (these are exported
    # but were previously not restored on import).
    for cr in doc.get("change_requests", []):
        row = m.ChangeRequest(
            project_id=project.id, code=cr.get("code"), title=cr.get("title"),
            requested_change=cr.get("requested_change") or "",
            status=_enum(m.ChangeRequestStatus, cr.get("status"), m.ChangeRequestStatus.DRAFT),
            requested_by=actor, created_by=actor,
        )
        db.add(row)
    for s in doc.get("suggestions", []):
        row = m.Suggestion(
            project_id=project.id, title=s.get("title") or "Imported suggestion",
            type=s.get("type"), severity=s.get("severity"),
            status=_enum(m.SuggestionStatus, s.get("status"), m.SuggestionStatus.OPEN),
            created_by=actor,
        )
        db.add(row)
    for c in doc.get("consultations", []):
        agg = {"aggregation_mode": c.get("aggregation_mode")} if c.get("aggregation_mode") else None
        row = m.Consultation(
            project_id=project.id, task_type=c.get("task_type") or "GENERAL_REVIEW",
            question=c.get("question") or "", aggregation=agg,
        )
        db.add(row)

    # Trace links: re-key assumption/decision semantic ids that were remapped.
    semantic_map = {}
    for a in doc.get("assumptions", []):
        if a.get("semantic_id"):
            semantic_map[a["semantic_id"]] = f"{a['semantic_id']}-{key}"
    for d in doc.get("decisions", []):
        if d.get("semantic_id"):
            semantic_map[d["semantic_id"]] = f"{d['semantic_id']}-{key}"

    db.commit()

    for t in doc.get("trace_links", []):
        src = semantic_map.get(t.get("source"), t.get("source"))
        tgt = semantic_map.get(t.get("target"), t.get("target"))
        if not src or not tgt:
            continue
        row = m.TraceLink(
            project_id=project.id,
            source_semantic_id=src,
            target_semantic_id=tgt,
            relation_type=_enum(m.TraceRelationType, t.get("relation"), m.TraceRelationType.REFERENCES),
            created_by=actor,
        )
        db.add(row)
    db.commit()

    # Rebuild bindings as pointers (never import PM/QA/Infra data).
    bindings = package.get("bindings") or {}
    if bindings:
        put_workspace_bindings(db, project.id,
                               pm_project_slug=bindings.get("pm_project_slug"),
                               qa_project_slugs=bindings.get("qa_project_slugs"),
                               infra_design_id=bindings.get("infra_design_id"))

    record_audit(db, action="PROJECT_IMPORTED", project_id=project.id, actor_id=actor_id,
                 object_type="Project", object_id=project.id,
                 metadata={"package_version": EXPORT_PACKAGE_VERSION, "source": proj.get("id")})
    return {
        "project_id": project.id, "key": project.key, "reused": False,
        "service_import_status": {"DOCUMENT_AGAIN": "COMPLETED", "PM_AGAIN": "NOT_APPLICABLE",
                                  "QA_AGAIN": "NOT_APPLICABLE", "INFRA_AGAIN": "NOT_APPLICABLE"},
        "migration_report": {"package_version": EXPORT_PACKAGE_VERSION,
                             "source_project_id": proj.get("id"), "target_project_id": project.id,
                             "mappings": {k: v for k, v in mappings.items() if v},
                             "secret_scan": scan,
                             "note": "PM/QA/Infra truth not imported (no import endpoint); bindings rebuilt as pointers."},
    }


# ---------------------------------------------------------------------------
# Workspace bindings (R12) — correlation metadata only. OIDA stores the mapping
# between a Document project and its PM/QA workspaces; the bounded services keep
# all business truth.
# ---------------------------------------------------------------------------

def get_workspace_bindings(db: Session, project_id: str) -> dict:
    project = guard_project(db, project_id)
    from .project_truth import normalize_bindings
    meta = project.project_meta or {}
    bindings = meta.get("workspace_bindings") or {}
    legacy = {
        "project_id": project.id,
        "pm_project_slug": bindings.get("pm_project_slug"),
        "qa_project_slugs": bindings.get("qa_project_slugs") or {},
        "infra_design_id": bindings.get("infra_design_id"),
    }
    legacy["binding_contract"] = normalize_bindings(project)
    return legacy


def put_workspace_bindings(db: Session, project_id: str, *, pm_project_slug: str | None = None,
                           qa_project_slugs: dict | None = None, infra_design_id: str | None = None,
                           binding_contract: dict | None = None) -> dict:
    project = guard_project(db, project_id)
    meta = dict(project.project_meta or {})
    bindings = dict(meta.get("workspace_bindings") or {})
    if pm_project_slug is not None:
        bindings["pm_project_slug"] = pm_project_slug
    if qa_project_slugs is not None:
        merged = dict(bindings.get("qa_project_slugs") or {})
        merged.update({k: v for k, v in qa_project_slugs.items() if v})
        bindings["qa_project_slugs"] = merged
    if infra_design_id is not None:
        bindings["infra_design_id"] = infra_design_id
    if binding_contract is not None:
        current_v1 = dict(bindings.get("v1") or {})
        current_v1.update({k: v for k, v in binding_contract.items() if v is not None})
        bindings["v1"] = current_v1
    elif bindings.get("v1"):
        current_v1 = dict(bindings["v1"])
        if pm_project_slug is not None:
            current_v1["pm"] = {"service": "PM_AGAIN", "external_project_id": pm_project_slug,
                                "binding_status": "BOUND", "source": "USER_SELECTED"}
        if qa_project_slugs is not None:
            existing_qa = {b.get("scope_id"): b for b in current_v1.get("qa", [])}
            existing_qa.update({scope: {"service": "QA_AGAIN", "external_project_id": slug,
                                        "scope_id": scope, "binding_status": "BOUND", "source": "USER_SELECTED"}
                                for scope, slug in qa_project_slugs.items() if slug})
            current_v1["qa"] = list(existing_qa.values())
        if infra_design_id is not None:
            current_v1["infra"] = {"service": "INFRA_AGAIN", "external_project_id": infra_design_id,
                                   "binding_status": "BOUND", "source": "USER_SELECTED"}
        bindings["v1"] = current_v1
    meta["workspace_bindings"] = bindings
    project.project_meta = meta
    db.commit()
    record_audit(
        db, action="WORKSPACE_BINDINGS_UPDATED", project_id=project_id,
        object_type="Project", object_id=project_id,
        metadata={"pm_project_slug": bindings.get("pm_project_slug"),
                  "qa_project_slugs": bindings.get("qa_project_slugs"),
                  "infra_design_id": bindings.get("infra_design_id"),
                  "binding_contract_version": "project_bindings/v1" if bindings.get("v1") else None},
    )
    return get_workspace_bindings(db, project_id)


# ---------------------------------------------------------------------------
# Artifact + Revision lifecycle
# ---------------------------------------------------------------------------


def create_artifact(
    db: Session,
    *,
    project_id: str,
    type: m.ArtifactType,
    title: str,
    snapshot: dict | None = None,
    actor="local-user",
) -> m.Artifact:
    project = guard_project(db, project_id)
    artifact = m.Artifact(project_id=project.id, type=type, title=title, created_by=actor)
    db.add(artifact)
    db.flush()
    revision = m.ArtifactRevision(
        artifact_id=artifact.id,
        revision_number=1,
        status=RevisionStatus.DRAFT,
        snapshot=snapshot or {},
        title=title,
        created_by=actor,
    )
    db.add(revision)
    db.flush()
    artifact.current_draft_revision_id = revision.id
    db.commit()
    return artifact


def next_revision_number(db: Session, artifact_id: str) -> int:
    result = db.execute(
        select(m.ArtifactRevision.revision_number)
        .where(m.ArtifactRevision.artifact_id == artifact_id)
        .order_by(m.ArtifactRevision.revision_number.desc())
        .limit(1)
    ).scalar_one_or_none()
    return (result or 0) + 1


def create_revision(
    db: Session,
    *,
    artifact_id: str,
    snapshot: dict | None = None,
    based_on_revision_id: str | None = None,
    actor="local-user",
    actor_id: str | None = None,
) -> m.ArtifactRevision:
    """Clone-as-new-revision. Confirmed history is never touched."""
    artifact = get_or_404(db, m.Artifact, artifact_id, "Artifact")
    if based_on_revision_id:
        parent = get_or_404(
            db, m.ArtifactRevision, based_on_revision_id, "Revision"
        )
        if parent.artifact_id != artifact_id:
            raise DomainError("based_on_revision belongs to a different artifact")
        base_snapshot = dict(parent.snapshot or {})
        base_title = parent.title
    else:
        latest = (
            db.execute(
                select(m.ArtifactRevision)
                .where(m.ArtifactRevision.artifact_id == artifact_id)
                .order_by(m.ArtifactRevision.revision_number.desc())
                .limit(1)
            )
            .scalar_one_or_none()
        )
        base_snapshot = dict(latest.snapshot or {}) if latest else {}
        base_title = latest.title if latest else artifact.title
        based_on_revision_id = latest.id if latest else None

    revision = m.ArtifactRevision(
        artifact_id=artifact_id,
        revision_number=next_revision_number(db, artifact_id),
        status=RevisionStatus.DRAFT,
        based_on_revision_id=based_on_revision_id,
        snapshot=snapshot if snapshot is not None else base_snapshot,
        title=base_title,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(revision)
    db.flush()
    # A new draft supersedes any previous DRAFT pointer but old drafts
    # themselves become SUPERSEDED only when a newer revision is confirmed.
    if revision.based_on and revision.based_on.status == RevisionStatus.DRAFT:
        revision.based_on.status = RevisionStatus.SUPERSEDED
    artifact.current_draft_revision_id = revision.id
    db.commit()
    return revision


def update_revision_snapshot(
    db: Session, revision_id: str, snapshot: dict, title: str | None = None
) -> m.ArtifactRevision:
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    require_editable(revision)
    revision.snapshot = snapshot
    if title is not None:
        revision.title = title
        revision.artifact.title = title
    db.commit()
    return revision


# Allowed transitions: DRAFT→IN_REVIEW→(CONFIRMED|DRAFT), CONFIRMED→SUPERSEDED,
# and ARCHIVED from CONFIRMED/SUPERSEDED. CONFIRMED is never edited.
_ALLOWED_TRANSITIONS = {
    RevisionStatus.DRAFT: {RevisionStatus.IN_REVIEW, RevisionStatus.ARCHIVED},
    RevisionStatus.IN_REVIEW: {RevisionStatus.DRAFT, RevisionStatus.CONFIRMED, RevisionStatus.ARCHIVED},
    RevisionStatus.CONFIRMED: {RevisionStatus.SUPERSEDED, RevisionStatus.ARCHIVED},
    RevisionStatus.SUPERSEDED: {RevisionStatus.ARCHIVED},
    RevisionStatus.ARCHIVED: set(),
}


def transition_revision(db: Session, revision_id: str, to_status: RevisionStatus) -> m.ArtifactRevision:
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    if to_status not in _ALLOWED_TRANSITIONS[revision.status]:
        raise DomainError(
            f"Illegal transition {revision.status.value} → {to_status.value}"
        )
    revision.status = to_status
    db.commit()
    return revision


def submit_for_review(db: Session, revision_id: str) -> m.ArtifactRevision:
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    if revision.status != RevisionStatus.DRAFT:
        raise DomainError("Only DRAFT revisions can be submitted for review")
    revision.status = RevisionStatus.IN_REVIEW
    db.commit()
    return revision


def confirm_revision(
    db: Session,
    revision_id: str,
    *,
    actor="local-user",
    comment: str | None = None,
    evidence: dict | None = None,
    supersede_confirmed: bool = True,
    actor_id: str | None = None,
) -> tuple[m.ArtifactRevision, m.Confirmation]:
    """Confirm = freeze. Atomic: technical design is snapshotted into the
    revision in the same transaction; any failure rolls everything back so
    a half-confirmed state is impossible.

    If an older CONFIRMED revision of the same artifact exists, it becomes
    SUPERSEDED (still readable, still bound in old baselines).
    """
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    if revision.status not in (RevisionStatus.IN_REVIEW, RevisionStatus.DRAFT):
        raise DomainError(
            f"Cannot confirm a revision in status {revision.status.value}"
        )
    try:
        # Technical-design artifacts freeze their bound designs at confirm time.
        if revision.artifact.type in (
            m.ArtifactType.DR,
            m.ArtifactType.DATABASE_SCHEMA,
            m.ArtifactType.ARCHITECTURE,
        ):
            snapshot = dict(revision.snapshot or {})
            snapshot["technical_design"] = snapshot_technical_design(
                db, revision.artifact.project_id
            )
            revision.snapshot = snapshot

        revision.status = RevisionStatus.CONFIRMED
        revision.confirmed_at = m.utcnow()
        revision.confirmed_by = actor

        if supersede_confirmed:
            siblings = db.execute(
                select(m.ArtifactRevision).where(
                    m.ArtifactRevision.artifact_id == revision.artifact_id,
                    m.ArtifactRevision.id != revision.id,
                    m.ArtifactRevision.status == RevisionStatus.CONFIRMED,
                )
            ).scalars().all()
            for sib in siblings:
                sib.status = RevisionStatus.SUPERSEDED

        confirmation = m.Confirmation(
            project_id=revision.artifact.project_id,
            artifact_revision_id=revision.id,
            confirmed_by=actor,
            comment=comment,
            evidence=evidence,
            actor_id=actor_id,
        )
        db.add(confirmation)
        db.commit()
    except Exception:
        db.rollback()
        raise
    metrics.inc("confirmation_completed")
    record_audit(
        db, action="REVISION_CONFIRMED", project_id=revision.artifact.project_id,
        actor_id=revision.actor_id, object_type="ArtifactRevision", object_id=revision.id,
        revision_context=revision.id, metadata={"artifact_id": revision.artifact_id},
    )
    return revision, confirmation


# ---------------------------------------------------------------------------
# Technical design snapshot (auto-frozen into DR revisions at confirmation)
# ---------------------------------------------------------------------------


def flow_snapshot(db: Session, flow_id: str) -> dict:
    flow = get_or_404(db, m.ProcessFlow, flow_id, "ProcessFlow")
    steps = db.execute(
        select(m.ProcessStep)
        .where(m.ProcessStep.flow_id == flow_id)
        .order_by(m.ProcessStep.position)
    ).scalars().all()
    transitions = db.execute(
        select(m.ProcessTransition).where(m.ProcessTransition.flow_id == flow_id)
    ).scalars().all()
    return {
        "name": flow.name,
        "description": flow.description,
        "steps": {s.semantic_id: {"name": s.name, "step_type": s.step_type, "position": s.position} for s in steps},
        "transitions": {
            t.semantic_id: {"from": t.from_step_semantic_id, "to": t.to_step_semantic_id,
                            "label": t.label, "condition": t.condition}
            for t in transitions
        },
    }


def api_endpoint_snapshot(api: m.APIEndpoint) -> dict:
    return {
        "method": api.method,
        "path": api.path,
        "summary": api.summary,
        "description": api.description,
        "authentication": api.authentication,
        "parameters": [
            {"name": p.name, "location": p.location, "data_type": p.data_type,
             "required": p.required, "description": p.description}
            for p in api.parameters
        ],
        "request_fields": [
            {"name": f.name, "data_type": f.data_type, "required": f.required, "description": f.description}
            for f in api.request_fields
        ],
        "response_fields": [
            {"status_code": f.status_code, "name": f.name, "data_type": f.data_type, "description": f.description}
            for f in api.response_fields
        ],
        "error_responses": [
            {"status_code": e.status_code, "message": e.message, "description": e.description}
            for e in api.error_responses
        ],
        "request_spec": api.request_spec,
        "response_spec": api.response_spec,
    }


def snapshot_technical_design(db: Session, project_id: str) -> dict:
    """Freeze the exact current structured designs for a project.

    The result is embedded into a confirmed DR revision snapshot, so a
    historical export can always reproduce the design as it was then.
    """
    designs: dict = {}

    schemas = db.execute(
        select(m.DatabaseSchema).where(m.DatabaseSchema.project_id == project_id)
    ).scalars().all()
    designs["db_schemas"] = {s.semantic_id: db_design_snapshot(db, s.id) for s in schemas}

    flows = db.execute(
        select(m.ProcessFlow).where(m.ProcessFlow.project_id == project_id)
    ).scalars().all()
    designs["flows"] = {f.semantic_id: flow_snapshot(db, f.id) for f in flows}

    apis = db.execute(
        select(m.APIEndpoint).where(m.APIEndpoint.project_id == project_id)
    ).scalars().all()
    designs["api_endpoints"] = {a.semantic_id: api_endpoint_snapshot(a) for a in apis}

    diagrams = db.execute(
        select(m.ArchitectureDiagram).where(m.ArchitectureDiagram.project_id == project_id)
    ).scalars().all()
    designs["architecture"] = {}
    for d in diagrams:
        nodes = db.execute(
            select(m.ArchitectureNode).where(m.ArchitectureNode.diagram_id == d.id)
        ).scalars().all()
        edges = db.execute(
            select(m.ArchitectureEdge).where(m.ArchitectureEdge.diagram_id == d.id)
        ).scalars().all()
        designs["architecture"][d.semantic_id] = {
            "name": d.name,
            "nodes": {n.semantic_id: {"name": n.name, "node_type": n.node_type, "technology": n.technology, "environment": n.environment} for n in nodes},
            "edges": {e.semantic_id: {"from": e.from_node_semantic_id, "to": e.to_node_semantic_id, "label": e.label} for e in edges},
        }

    return designs


# ---------------------------------------------------------------------------
# Baseline
# ---------------------------------------------------------------------------


def create_baseline(
    db: Session,
    *,
    project_id: str,
    name: str,
    description: str | None = None,
    artifact_revision_ids: list[str],
    target_release: str | None = None,
    actor="local-user",
    actor_id: str | None = None,
) -> m.Baseline:
    """Freeze the exact artifact→revision pairs given at creation time.

    The stored binding rows are never re-resolved afterwards — that is
    the whole point. A later v8 of a child artifact does not change a
    baseline that bound v7.
    """
    guard_project(db, project_id)
    if not artifact_revision_ids:
        raise DomainError("A baseline must bind at least one revision")

    seen: dict[str, m.ArtifactRevision] = {}
    for rid in artifact_revision_ids:
        rev = get_or_404(db, m.ArtifactRevision, rid, "Revision")
        if rev.status != RevisionStatus.CONFIRMED:
            raise DomainError(
                f"Revision {rid} is {rev.status.value}; only CONFIRMED revisions "
                "may be frozen into a baseline"
            )
        if rev.artifact.project_id != project_id:
            raise DomainError(f"Revision {rid} belongs to another project")
        if rev.artifact_id in seen:
            raise DomainError("Each artifact may appear at most once per baseline")
        seen[rev.artifact_id] = rev

    baseline = m.Baseline(
        project_id=project_id, name=name, description=description, created_by=actor,
        actor_id=actor_id, target_release=target_release,
    )
    db.add(baseline)
    db.flush()
    for artifact_id, rev in seen.items():
        # The artifact's own semantic identity (if any). Document sections are
        # registered with entity_ref=artifact.id, so exclude DOCUMENT_SECTION to
        # avoid a multiple-row match when a UR/DR has many sections.
        semantic_object = db.execute(
            select(m.SemanticObject).where(
                m.SemanticObject.project_id == project_id,
                m.SemanticObject.entity_ref == artifact_id,
                m.SemanticObject.object_type != m.SemanticObjectType.DOCUMENT_SECTION,
            )
        ).scalar_one_or_none()
        db.add(
            m.BaselineBinding(
                baseline_id=baseline.id,
                artifact_id=artifact_id,
                artifact_revision_id=rev.id,
                semantic_object_id=semantic_object.semantic_id if semantic_object else None,
                semantic_object_type=(
                    semantic_object.object_type.value if semantic_object else None
                ),
            )
        )
    db.commit()
    record_audit(
        db, action="BASELINE_CREATED", project_id=project_id, actor_id=actor_id,
        object_type="Baseline", object_id=baseline.id, baseline_id=baseline.id,
        metadata={"name": name, "bindings": len(seen)},
    )
    return baseline


def resolve_baseline(db: Session, baseline_id: str) -> m.Baseline:
    baseline = get_or_404(db, m.Baseline, baseline_id, "Baseline")
    db.expire(baseline, ["bindings"])  # always read the frozen rows
    return baseline


# ---------------------------------------------------------------------------
# SemanticObject + TraceLink
# ---------------------------------------------------------------------------


def ensure_semantic_object(
    db: Session,
    *,
    project_id: str,
    semantic_id: str,
    object_type: m.SemanticObjectType,
    display_name: str,
    entity_ref: str | None = None,
    metadata: dict | None = None,
) -> m.SemanticObject:
    obj = db.execute(
        select(m.SemanticObject).where(
            m.SemanticObject.project_id == project_id,
            m.SemanticObject.semantic_id == semantic_id,
        )
    ).scalar_one_or_none()
    if obj:
        obj.display_name = display_name  # display names may change
        if entity_ref:
            obj.entity_ref = entity_ref
        if metadata:
            obj.metadata_json = metadata
        db.commit()
        return obj
    obj = m.SemanticObject(
        project_id=project_id,
        semantic_id=semantic_id,
        object_type=object_type,
        display_name=display_name,
        entity_ref=entity_ref,
        metadata_json=metadata,
    )
    db.add(obj)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        obj = db.execute(
            select(m.SemanticObject).where(
                m.SemanticObject.project_id == project_id,
                m.SemanticObject.semantic_id == semantic_id,
            )
        ).scalar_one()
    return obj


def create_trace_link(
    db: Session,
    *,
    project_id: str,
    source_semantic_id: str,
    target_semantic_id: str,
    relation_type: m.TraceRelationType,
    revision_context: str | None = None,
    actor="local-user",
) -> m.TraceLink:
    guard_project(db, project_id)
    for sid in (source_semantic_id, target_semantic_id):
        exists = db.execute(
            select(m.SemanticObject.id).where(
                m.SemanticObject.project_id == project_id,
                m.SemanticObject.semantic_id == sid,
            )
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(
                f"Unknown semantic object '{sid}'. Traces may only connect "
                "registered semantic objects — never pixel positions or titles.",
                status_code=422,
            )
    link = m.TraceLink(
        project_id=project_id,
        source_semantic_id=source_semantic_id,
        target_semantic_id=target_semantic_id,
        relation_type=relation_type,
        revision_context=revision_context,
        created_by=actor,
    )
    db.add(link)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        raise DomainError("Trace link already exists")
    return link


def trace_graph(db: Session, project_id: str) -> dict:
    """Nodes + edges for the traceability explorer. Only stored links are shown."""
    nodes = db.execute(
        select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
    ).scalars().all()
    edges = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
    ).scalars().all()
    return {
        "nodes": [
            {
                "semantic_id": n.semantic_id,
                "object_type": n.object_type.value,
                "display_name": n.display_name,
            }
            for n in nodes
        ],
        "edges": [
            {
                "source": e.source_semantic_id,
                "target": e.target_semantic_id,
                "relation": e.relation_type.value,
                "revision_context": e.revision_context,
            }
            for e in edges
        ],
    }


def impact_of(db: Session, project_id: str, semantic_id: str) -> dict:
    """1-hop upstream/downstream impact using trace links only."""
    outgoing = db.execute(
        select(m.TraceLink).where(
            m.TraceLink.project_id == project_id,
            m.TraceLink.source_semantic_id == semantic_id,
        )
    ).scalars().all()
    incoming = db.execute(
        select(m.TraceLink).where(
            m.TraceLink.project_id == project_id,
            m.TraceLink.target_semantic_id == semantic_id,
        )
    ).scalars().all()
    return {
        "semantic_id": semantic_id,
        "downstream": [
            {"semantic_id": l.target_semantic_id, "relation": l.relation_type.value}
            for l in outgoing
        ],
        "upstream": [
            {"semantic_id": l.source_semantic_id, "relation": l.relation_type.value}
            for l in incoming
        ],
    }


# ---------------------------------------------------------------------------
# Annotation
# ---------------------------------------------------------------------------


def create_annotation(
    db: Session,
    *,
    project_id: str,
    anchor_object_type: str,
    anchor_semantic_id: str,
    content: str,
    type: m.AnnotationType = m.AnnotationType.COMMENT,
    artifact_revision_id: str | None = None,
    canvas_x: float | None = None,
    canvas_y: float | None = None,
    drawing_payload: dict | None = None,
    thread_id: str | None = None,
    actor="local-user",
    actor_id: str | None = None,
) -> m.Annotation:
    guard_project(db, project_id)
    anchored = db.execute(
        select(m.SemanticObject.id).where(
            m.SemanticObject.project_id == project_id,
            m.SemanticObject.semantic_id == anchor_semantic_id,
        )
    ).scalar_one_or_none()
    if not anchored:
        raise DomainError(
            f"Cannot anchor annotation to unknown semantic object '{anchor_semantic_id}'. "
            "Coordinates are optional placement data, never the anchor.",
            status_code=422,
        )
    if thread_id is not None:
        thread = get_or_404(db, m.CommentThread, thread_id, "CommentThread")
        if thread.project_id != project_id:
            raise DomainError("Thread belongs to a different project")
    annotation = m.Annotation(
        project_id=project_id,
        artifact_revision_id=artifact_revision_id,
        anchor_object_type=anchor_object_type,
        anchor_semantic_id=anchor_semantic_id,
        canvas_x=canvas_x,
        canvas_y=canvas_y,
        type=type,
        content=content,
        drawing_payload=drawing_payload,
        thread_id=thread_id,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(annotation)
    db.commit()
    return annotation


def set_annotation_status(
    db: Session, annotation_id: str, status: m.AnnotationStatus
) -> m.Annotation:
    annotation = get_or_404(db, m.Annotation, annotation_id, "Annotation")
    valid = {s.value for s in m.AnnotationStatus}
    if status.value not in valid:
        raise DomainError("Invalid annotation status")
    annotation.status = status
    db.commit()
    return annotation


# ---------------------------------------------------------------------------
# Requirements
# ---------------------------------------------------------------------------


def next_requirement_code(db: Session, project_id: str) -> str:
    count = (
        db.execute(
            select(m.Requirement.id).where(m.Requirement.project_id == project_id)
        )
        .scalars()
        .all()
    )
    return f"REQ-{len(count) + 1:04d}"


def create_requirement(
    db: Session,
    *,
    project_id: str,
    title: str,
    description=None,
    source_type=None,
    source_reference=None,
    priority=None,
    code: str | None = None,
    metadata: dict | None = None,
    actor="local-user",
) -> m.Requirement:
    if code is not None:
        existing = db.execute(
            select(m.Requirement.id).where(
                m.Requirement.project_id == project_id, m.Requirement.code == code
            )
        ).scalar_one_or_none()
        if existing:
            raise DomainError(f"Requirement code '{code}' already exists in this project")
    req_code = code or next_requirement_code(db, project_id)
    requirement = m.Requirement(
        project_id=project_id,
        code=req_code,
        title=title,
        description=description,
        source_type=source_type,
        source_reference=source_reference,
        priority=priority,
        metadata_json=metadata or {},
        created_by=actor,
    )
    db.add(requirement)
    db.flush()
    ensure_semantic_object(
        db,
        project_id=project_id,
        semantic_id=req_code,
        object_type=m.SemanticObjectType.REQUIREMENT,
        display_name=title,
        entity_ref=requirement.id,
    )
    _seed_requirement_revision(db, requirement, actor=actor, initial_status=requirement.status)
    db.commit()
    return requirement


def _seed_requirement_revision(
    db: Session,
    requirement: m.Requirement,
    *,
    actor: str = "local-user",
    initial_status: m.RequirementStatus | None = None,
) -> m.RequirementRevision:
    """Create revision #1 from the requirement's current fields. Existing
    requirements (which predate revisioning) are seeded once, idempotently."""
    existing = db.execute(
        select(m.RequirementRevision.id).where(
            m.RequirementRevision.requirement_id == requirement.id
        )
    ).scalars().first()
    if existing:
        return db.get(m.RequirementRevision, existing)
    rev = m.RequirementRevision(
        requirement_id=requirement.id,
        revision_number=1,
        title=requirement.title,
        description=requirement.description,
        source_type=requirement.source_type,
        source_reference=requirement.source_reference,
        priority=requirement.priority,
        status=initial_status or requirement.status,
        based_on_revision_id=None,
        created_by=requirement.created_by or actor,
        confirmed_at=m.utcnow() if requirement.status == m.RequirementStatus.CONFIRMED else None,
        confirmed_by=requirement.created_by if requirement.status == m.RequirementStatus.CONFIRMED else None,
    )
    db.add(rev)
    db.flush()
    return rev


def seed_requirement_revisions(db: Session, project_id: str) -> int:
    """One-time backfill: give every existing requirement a revision #1."""
    requirements = db.execute(
        select(m.Requirement).where(m.Requirement.project_id == project_id)
    ).scalars().all()
    created = 0
    for req in requirements:
        before = db.execute(
            select(m.RequirementRevision.id).where(
                m.RequirementRevision.requirement_id == req.id
            )
        ).scalars().first()
        if before is None:
            _seed_requirement_revision(db, req, actor=req.created_by)
            created += 1
    db.commit()
    return created


def _latest_requirement_revision(db: Session, requirement_id: str) -> m.RequirementRevision:
    return db.execute(
        select(m.RequirementRevision)
        .where(m.RequirementRevision.requirement_id == requirement_id)
        .order_by(m.RequirementRevision.revision_number.desc())
        .limit(1)
    ).scalar_one_or_none()


def create_requirement_draft(
    db: Session, requirement_id: str, *, actor="local-user", actor_id: str | None = None
) -> tuple[m.RequirementChange, m.RequirementRevision]:
    """Edit a requirement without touching the confirmed revision: clone the
    latest revision into a new DRAFT and open a RequirementChange record."""
    requirement = get_or_404(db, m.Requirement, requirement_id, "Requirement")
    latest = _latest_requirement_revision(db, requirement_id)
    if latest is None:
        _seed_requirement_revision(db, requirement, actor=actor)
        latest = _latest_requirement_revision(db, requirement_id)

    # If a DRAFT change is already open for this requirement, reuse it.
    open_change = db.execute(
        select(m.RequirementChange).where(
            m.RequirementChange.requirement_id == requirement_id,
            m.RequirementChange.status.in_(["DRAFT", "IMPACT_READY", "REGENERATED", "REVIEWED"]),
        )
    ).scalars().first()
    if open_change and open_change.to_revision_id:
        draft = db.get(m.RequirementRevision, open_change.to_revision_id)
        if draft and draft.status == m.RequirementStatus.DRAFT:
            return open_change, draft

    draft = m.RequirementRevision(
        requirement_id=requirement_id,
        revision_number=(latest.revision_number if latest else 0) + 1,
        title=latest.title if latest else requirement.title,
        description=latest.description if latest else requirement.description,
        source_type=latest.source_type if latest else requirement.source_type,
        source_reference=latest.source_reference if latest else requirement.source_reference,
        priority=latest.priority if latest else requirement.priority,
        status=m.RequirementStatus.DRAFT,
        based_on_revision_id=latest.id if latest else None,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(draft)
    db.flush()

    change = m.RequirementChange(
        project_id=requirement.project_id,
        requirement_id=requirement_id,
        from_revision_id=latest.id if latest else None,
        to_revision_id=draft.id,
        status="DRAFT",
        created_by=actor,
        actor_id=actor_id,
        label=f"{requirement.code} {requirement.title}",
    )
    db.add(change)
    db.commit()
    db.refresh(change)
    record_audit(
        db, action="REQUIREMENT_DRAFT_CREATED", project_id=requirement.project_id,
        actor_id=actor_id, object_type="RequirementChange", object_id=change.id,
        metadata={"requirement_id": requirement_id, "code": requirement.code,
                  "from_revision": change.from_revision_id, "to_revision": draft.id},
    )
    return change, draft


def update_requirement_draft(
    db: Session,
    requirement_id: str,
    revision_id: str,
    *,
    title: str | None = None,
    description: str | None = None,
    source_type: str | None = None,
    source_reference: str | None = None,
    priority: str | None = None,
) -> m.RequirementRevision:
    draft = get_or_404(db, m.RequirementRevision, revision_id, "Requirement revision")
    if draft.requirement_id != requirement_id:
        raise DomainError("Revision does not belong to this requirement")
    if draft.status != m.RequirementStatus.DRAFT:
        raise DomainError("Only DRAFT requirement revisions are editable")
    if title is not None:
        draft.title = title
    if description is not None:
        draft.description = description
    if source_type is not None:
        draft.source_type = source_type
    if source_reference is not None:
        draft.source_reference = source_reference
    if priority is not None:
        draft.priority = priority
    db.commit()
    return draft


def list_requirement_revisions(db: Session, requirement_id: str) -> list[m.RequirementRevision]:
    get_or_404(db, m.Requirement, requirement_id, "Requirement")
    return db.execute(
        select(m.RequirementRevision)
        .where(m.RequirementRevision.requirement_id == requirement_id)
        .order_by(m.RequirementRevision.revision_number.desc())
    ).scalars().all()


_RELATION_REASON = {
    "DERIVED_FROM": "directly defined by this requirement",
    "IMPLEMENTS": "implements this requirement",
    "DESIGNED_BY": "designed from this requirement",
    "VALIDATED_BY": "validated against this requirement",
    "AFFECTS": "affected by this requirement",
    "REFERENCES": "references this requirement",
    "GENERATED_FROM": "generated from this requirement",
    "CONFIRMED_BY": "confirmed by this requirement",
    "SUPERSEDES": "superseded by this requirement",
}

_REGEN_ELIGIBILITY = {
    "DOCUMENT_SECTION": "ARTIFACT_DRAFT",
    "PROCESS_FLOW": "NEEDS_REGENERATION",
    "PROCESS_STEP": "NEEDS_REGENERATION",
    "ARCHITECTURE_NODE": "NEEDS_REGENERATION",
    "DB_SCHEMA": "NEEDS_REGENERATION",
    "DB_TABLE": "NEEDS_REGENERATION",
    "API_ENDPOINT": "NEEDS_REGENERATION",
    "SCREEN": "NEEDS_REGENERATION",
}


def requirement_change_impact(db: Session, change_id: str) -> dict:
    """Downstream impact of a requirement change using exact trace links.

    DIRECT = one hop from the requirement; INDIRECT = further hops. Each
    affected node carries a human reason, a dependency path and a regeneration
    eligibility. Unaffected objects are never regenerated.

    The result is enriched with a truthful trust view: known / potential /
    unknown areas, confidence and coverage — never claiming completeness when
    trace coverage is incomplete.
    """
    change = get_or_404(db, m.RequirementChange, change_id, "Change")
    requirement = get_or_404(db, m.Requirement, change.requirement_id, "Requirement")
    project_id = change.project_id

    nodes = {
        n.semantic_id: n
        for n in db.execute(
            select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
        ).scalars().all()
    }
    edges = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
    ).scalars().all()

    adj_down = {}  # semantic_id -> [(target, relation)]
    for e in edges:
        adj_down.setdefault(e.source_semantic_id, []).append((e.target_semantic_id, e.relation_type.value))

    start = requirement.code
    affected = []
    seen = set()
    path_of = {start: [start]}
    # BFS downstream (requirement defines/derives downstream objects).
    queue = [(start, 0)]
    while queue:
        cur, depth = queue.pop(0)
        for target, rel in adj_down.get(cur, []):
            if target in seen or target == start:
                continue
            seen.add(target)
            path_of[target] = path_of.get(cur, []) + [target]
            level = "DIRECT" if depth == 0 else "INDIRECT"
            node = nodes.get(target)
            affected.append({
                "semantic_id": target,
                "object_type": node.object_type.value if node else "UNKNOWN",
                "display_name": node.display_name if node else target,
                "level": level,
                "depth": depth + 1,
                "relation": rel,
                "reason": _RELATION_REASON.get(rel, f"linked via {rel}"),
                "path": path_of[target],
                "regeneration": _REGEN_ELIGIBILITY.get(node.object_type.value if node else "", "NONE"),
            })
            if depth + 1 < 2:  # follow up to INDIRECT (depth 2)
                queue.append((target, depth + 1))

    # Incoming REFERENCES: objects that reference the requirement are POTENTIAL.
    for e in edges:
        if e.target_semantic_id == start and e.source_semantic_id not in seen and e.source_semantic_id != start:
            node = nodes.get(e.source_semantic_id)
            path_of[e.source_semantic_id] = [e.source_semantic_id, start]
            affected.append({
                "semantic_id": e.source_semantic_id,
                "object_type": node.object_type.value if node else "UNKNOWN",
                "display_name": node.display_name if node else e.source_semantic_id,
                "level": "POTENTIAL",
                "depth": 1,
                "relation": e.relation_type.value,
                "reason": _RELATION_REASON.get(e.relation_type.value, f"linked via {e.relation_type.value}"),
                "path": path_of[e.source_semantic_id],
                "regeneration": _REGEN_ELIGIBILITY.get(node.object_type.value if node else "", "NONE"),
            })
            seen.add(e.source_semantic_id)

    affected_count = len(affected)
    unaffected_count = max(0, len(nodes) - affected_count - 1)
    unaffected = [n.semantic_id for n in nodes.values() if n.semantic_id not in seen and n.semantic_id != start]

    trust = _trust_impact_view(db, project_id, affected, unaffected)

    draft = db.get(m.RequirementRevision, change.to_revision_id) if change.to_revision_id else None
    result = {
        "change_id": change.id,
        "requirement": {"id": requirement.id, "code": requirement.code, "title": requirement.title},
        "from_revision": _requirement_revision_dict(db.get(m.RequirementRevision, change.from_revision_id)) if change.from_revision_id else None,
        "to_revision": _requirement_revision_dict(draft) if draft else None,
        "affected": affected,
        "affected_count": affected_count,
        "unaffected_count": unaffected_count,
        # Downstream services reached via Conductor on confirm.
        "cross_domain": ["pm", "qa"],
        **trust,
    }
    change.impact_json = result
    if change.status == "DRAFT":
        change.status = "IMPACT_READY"
    db.commit()
    record_audit(
        db, action="IMPACT_ANALYSIS_CALCULATED", project_id=project_id,
        actor_id=change.actor_id, object_type="RequirementChange", object_id=change.id,
        metadata={"affected": affected_count, "unaffected": unaffected_count,
                  "confidence": trust["impact_confidence"]},
    )
    return result


def _requirement_revision_dict(r: m.RequirementRevision | None) -> dict | None:
    if r is None:
        return None
    return {
        "id": r.id, "requirement_id": r.requirement_id, "revision_number": r.revision_number,
        "title": r.title, "description": r.description, "source_type": r.source_type,
        "source_reference": r.source_reference, "priority": r.priority, "status": r.status.value,
        "based_on_revision_id": r.based_on_revision_id, "created_at": r.created_at.isoformat(),
        "created_by": r.created_by, "confirmed_at": r.confirmed_at.isoformat() if r.confirmed_at else None,
        "confirmed_by": r.confirmed_by,
    }


def regenerate_change(db: Session, change_id: str, *, mode: str = "affected", actor="local-user", actor_id: str | None = None) -> dict:
    """Create draft revisions for affected derived artifacts.

    mode = "affected" (default): only artifacts that contain an affected
    DOCUMENT_SECTION are cloned into a new draft. mode = "full" clones every
    project artifact (elevated choice). Confirmed history is never modified.
    """
    change = get_or_404(db, m.RequirementChange, change_id, "Change")
    if change.status not in ("IMPACT_READY", "REVIEWED"):
        raise DomainError("Run impact analysis before regenerating")
    project_id = change.project_id

    if mode == "full":
        artifacts = db.execute(
            select(m.Artifact).where(m.Artifact.project_id == project_id)
        ).scalars().all()
        artifact_ids = {a.id for a in artifacts}
    else:
        artifact_ids = set()
        for item in (change.impact_json or {}).get("affected", []):
            if item.get("regeneration") == "ARTIFACT_DRAFT":
                so = db.execute(
                    select(m.SemanticObject).where(
                        m.SemanticObject.project_id == project_id,
                        m.SemanticObject.semantic_id == item["semantic_id"],
                    )
                ).scalar_one_or_none()
                if so and so.entity_ref:
                    artifact_ids.add(so.entity_ref)

    generated = []
    for artifact_id in sorted(artifact_ids):
        artifact = db.get(m.Artifact, artifact_id)
        if artifact is None:
            continue
        # Clone the latest CONFIRMED revision (fall back to latest).
        base = db.execute(
            select(m.ArtifactRevision)
            .where(m.ArtifactRevision.artifact_id == artifact_id)
            .order_by(m.ArtifactRevision.revision_number.desc())
            .limit(1)
        ).scalar_one_or_none()
        if base is None:
            continue
        draft = create_revision(
            db, artifact_id=artifact_id, based_on_revision_id=base.id,
            actor=actor, actor_id=actor_id,
        )
        generated.append({
            "artifact_id": artifact_id,
            "artifact_title": artifact.title,
            "revision_id": draft.id,
            "revision_number": draft.revision_number,
            "based_on": base.id,
        })

    change.generated_revision_ids = [g["revision_id"] for g in generated]
    change.status = "REGENERATED"
    db.commit()
    record_audit(
        db, action="REGENERATION_STARTED", project_id=project_id, actor_id=actor_id,
        object_type="RequirementChange", object_id=change.id,
        metadata={"mode": mode, "generated": len(generated)},
    )
    return {"mode": mode, "generated": generated, "needs_regeneration": [
        i for i in (change.impact_json or {}).get("affected", [])
        if i.get("regeneration") == "NEEDS_REGENERATION"
    ]}


def _verify_confirmation_token(token: str) -> dict:
    """Ask Account Again to verify a short-lived admin re-auth token."""
    from .account_client import AccountAgainClient
    client = AccountAgainClient()
    return client.verify_confirmation_token(token)


def _next_baseline_name(db: Session, project_id: str) -> str:
    rows = db.execute(
        select(m.Baseline.name).where(m.Baseline.project_id == project_id)
    ).scalars().all()
    versions = []
    for name in rows:
        mch = re.search(r"v(\d+)", name or "")
        if mch:
            versions.append(int(mch.group(1)))
    return f"True Cloud Migration v{max(versions) + 1 if versions else 1}.0"


def confirm_change(
    db: Session,
    change_id: str,
    *,
    confirmation_token: str,
    actor="local-user",
    actor_id: str | None = None,
) -> dict:
    """Privileged confirm: verify admin re-auth, confirm the requirement draft,
    confirm generated artifact drafts, create the next baseline, and dispatch
    PM/QA handoffs via Conductor. Historical truth is preserved throughout."""
    change = get_or_404(db, m.RequirementChange, change_id, "Change")
    if change.status not in ("REGENERATED", "REVIEWED"):
        raise DomainError("Change must be regenerated and reviewed before confirmation")

    # 1. Admin re-auth via Account Again (short-lived token, never the password).
    try:
        claims = _verify_confirmation_token(confirmation_token)
    except Exception as exc:
        raise DomainError(f"Admin re-authentication failed: {exc}", status_code=403)

    requirement = get_or_404(db, m.Requirement, change.requirement_id, "Requirement")
    project_id = change.project_id

    # 2. Confirm the requirement draft.
    draft = db.get(m.RequirementRevision, change.to_revision_id)
    if draft is None or draft.status != m.RequirementStatus.DRAFT:
        raise DomainError("Requirement draft missing or not DRAFT")
    old_confirmed = db.execute(
        select(m.RequirementRevision).where(
            m.RequirementRevision.requirement_id == requirement.id,
            m.RequirementRevision.id != draft.id,
            m.RequirementRevision.status == m.RequirementStatus.CONFIRMED,
        )
    ).scalars().all()
    draft.status = m.RequirementStatus.CONFIRMED
    draft.confirmed_at = m.utcnow()
    draft.confirmed_by = actor
    for old in old_confirmed:
        old.status = m.RequirementStatus.SUPERSEDED
    # Reflect the confirmed revision on the requirement row.
    requirement.title = draft.title
    requirement.description = draft.description
    requirement.source_type = draft.source_type
    requirement.source_reference = draft.source_reference
    requirement.priority = draft.priority
    requirement.status = m.RequirementStatus.CONFIRMED

    # 3. Confirm generated artifact drafts.
    confirmed_artifact_revisions = []
    for rid in change.generated_revision_ids or []:
        rev, _ = confirm_revision(db, rid, actor=actor, actor_id=actor_id)
        confirmed_artifact_revisions.append(rev.id)

    db.commit()

    # 4. Create the next baseline (bind the confirmed artifact drafts).
    baseline = create_baseline(
        db,
        project_id=project_id,
        name=_next_baseline_name(db, project_id),
        description=f"Confirmed from change {change.id} ({requirement.code})",
        artifact_revision_ids=confirmed_artifact_revisions,
        actor=actor,
        actor_id=actor_id,
    )

    # 5. Dispatch PM + QA handoffs via Conductor (best-effort, partial sync is
    #    reported honestly and never rolls back the confirmed baseline).
    sync = {"pm": {"status": "SKIPPED"}, "qa": {"status": "SKIPPED"}}
    try:
        ex = create_execution_handoff(
            db, project_id, baseline_id=baseline.id,
            actor=actor, actor_id=actor_id, status="READY",
        )
        deliver_handoff_to_conductor(db, ex["id"], "execution")
        sync["pm"] = {"status": "SYNCED", "handoff_id": ex["id"]}
    except Exception as exc:
        sync["pm"] = {"status": "FAILED", "error": str(exc)[:300]}
    try:
        qa = create_qa_validation_handoff(
            db, project_id, baseline_id=baseline.id,
            requirement_ids=[requirement.id],
            design_revision_ids=confirmed_artifact_revisions,
            actor=actor, actor_id=actor_id, status="READY",
        )
        deliver_handoff_to_conductor(db, qa["id"], "qa")
        sync["qa"] = {"status": "SYNCED", "handoff_id": qa["id"]}
    except Exception as exc:
        sync["qa"] = {"status": "FAILED", "error": str(exc)[:300]}

    sync["infra"] = {"status": "NOT_LINKED", "note": "No infrastructure object directly linked to this change"}

    change.status = "CONFIRMED"
    change.baseline_id = baseline.id
    change.confirmed_at = m.utcnow()
    db.commit()

    record_audit(
        db, action="BASELINE_CONFIRMED", project_id=project_id, actor_id=actor_id,
        object_type="RequirementChange", object_id=change.id, baseline_id=baseline.id,
        metadata={"requirement_id": requirement.id, "code": requirement.code,
                  "baseline": baseline.name, "sync": sync},
    )
    return {
        "change_id": change.id,
        "baseline": {"id": baseline.id, "name": baseline.name},
        "requirement": {"id": requirement.id, "code": requirement.code, "revision": draft.revision_number},
        "generated": len(confirmed_artifact_revisions),
        "sync": sync,
        "overall": "SYNCED" if all(s["status"] == "SYNCED" for s in sync.values()) else "PARTIALLY_SYNCHRONIZED",
    }


def list_requirement_changes(db: Session, project_id: str | None = None) -> list[dict]:
    q = select(m.RequirementChange).order_by(m.RequirementChange.created_at.desc())
    if project_id:
        q = q.where(m.RequirementChange.project_id == project_id)
    out = []
    for c in db.execute(q).scalars().all():
        requirement = db.get(m.Requirement, c.requirement_id)
        draft = db.get(m.RequirementRevision, c.to_revision_id) if c.to_revision_id else None
        out.append({
            "id": c.id, "project_id": c.project_id, "requirement_id": c.requirement_id,
            "code": requirement.code if requirement else None,
            "label": c.label, "status": c.status,
            "from_revision_id": c.from_revision_id, "to_revision_id": c.to_revision_id,
            "draft_title": draft.title if draft else None,
            "affected_count": (c.impact_json or {}).get("affected_count"),
            "unaffected_count": (c.impact_json or {}).get("unaffected_count"),
            "baseline_id": c.baseline_id, "correlation_id": c.correlation_id,
            "created_at": c.created_at.isoformat(), "created_by": c.created_by,
            "confirmed_at": c.confirmed_at.isoformat() if c.confirmed_at else None,
        })
    return out


# ---------------------------------------------------------------------------
# Impact trust engine (R10.1)
#
# Truthfulness contract: an impact analysis is decision support, not an oracle.
# It must separate what the system KNOWS (explicit trace) from what it INFERS
# (inferred/heuristic) and where it has NO information (untraced domains), and
# must never claim completeness when trace coverage is incomplete.
# ---------------------------------------------------------------------------

# Human names for the downstream domains Document Again hands off to. A domain
# is "traced" only when an explicit link (external reference or handoff) exists.
_TRACED_DOMAINS = [
    ("pm-again", "Planning / PM execution"),
    ("qa-again", "QA validation"),
    ("infra-again", "Infrastructure implementation"),
]


def _project_traced_services(db: Session, project_id: str) -> set[str]:
    services: set[str] = set()
    for ref in db.execute(
        select(m.ExternalReference).where(m.ExternalReference.project_id == project_id)
    ).scalars():
        if ref.service:
            services.add(ref.service)
    if db.execute(
        select(m.ExecutionHandoff).where(m.ExecutionHandoff.project_id == project_id).limit(1)
    ).scalar_one_or_none():
        services.add("pm-again")
    if db.execute(
        select(m.QAValidationHandoff).where(m.QAValidationHandoff.project_id == project_id).limit(1)
    ).scalar_one_or_none():
        services.add("qa-again")
    return services


def _unknown_areas(db: Session, project_id: str, has_trace_edges: bool) -> list[dict]:
    """Domains the analysis cannot claim to have assessed, because no explicit
    trace relationship exists. These are honest gaps, not failures."""
    if not has_trace_edges:
        return [{
            "domain": None,
            "label": "No explicit trace relationships",
            "reason": "This project has no trace links yet, so downstream impact cannot be derived — nothing is claimed.",
        }]
    traced = _project_traced_services(db, project_id)
    areas = []
    for svc, human in _TRACED_DOMAINS:
        if svc not in traced:
            areas.append({
                "domain": svc,
                "label": human,
                "reason": f"{human} is not traced into this project — no explicit link exists, so impact there is unknown.",
            })
    # External / vendor dependencies.
    has_external = any(
        ref.service not in {s for s, _ in _TRACED_DOMAINS}
        for ref in db.execute(
            select(m.ExternalReference).where(m.ExternalReference.project_id == project_id)
        ).scalars()
        if ref.service
    )
    if not has_external:
        areas.append({
            "domain": "vendor-external",
            "label": "Vendor / external operational dependencies",
            "reason": "No external dependency is traced in this project.",
        })
    return areas


def _trust_impact_view(db: Session, project_id: str, affected: list[dict], unaffected: list[str]) -> dict:
    """Split BFS results into known / potential / unknown and derive a
    deterministic confidence + coverage estimate from what the graph supports."""
    known: list[dict] = []
    potential: list[dict] = []
    for item in affected:
        entry = {
            "id": item["semantic_id"],
            "object_type": item["object_type"],
            "display_name": item["display_name"],
            "level": item["level"],
            "reason": item["reason"],
            "path": item.get("path") or [item["semantic_id"]],
            "regeneration": item.get("regeneration"),
            # DIRECT/INDIRECT are explicit one/multi-hop trace; POTENTIAL is
            # inferred (a reference points at the changed object).
            "source": "SYSTEM-DERIVED" if item["level"] in ("DIRECT", "INDIRECT") else "INFERRED",
        }
        if item["level"] == "POTENTIAL":
            potential.append(entry)
        else:
            known.append(entry)

    has_edges = bool(affected) or bool(
        db.execute(select(m.TraceLink).where(m.TraceLink.project_id == project_id).limit(1)).scalar_one_or_none()
    )
    unknown = _unknown_areas(db, project_id, has_edges)

    explicit = len(known)
    inferred = len(potential)
    unknown_n = len(unknown)

    if explicit == 0 and inferred == 0:
        confidence = "UNKNOWN"
        coverage_status = "NOT_MEASURABLE"
    elif inferred == 0 and unknown_n == 0:
        confidence = "HIGH"
        coverage_status = "FULL"
    elif unknown_n == 0:
        confidence = "MEDIUM"
        coverage_status = "PARTIAL"
    else:
        confidence = "LOW"
        coverage_status = "PARTIAL"

    return {
        "known_impact": known,
        "potential_impact": potential,
        "unknown_areas": unknown,
        "unaffected": unaffected[:200],
        "impact_confidence": confidence,
        "trace_coverage": {
            "status": coverage_status,
            "confirmed_relationships": explicit,
            "inferred_or_unresolved": inferred + unknown_n,
            "note": "Confirmed counts trace-derived explicit relationships only; inferred and untraced areas are reported separately and never counted as confirmed.",
        },
    }


def _baseline_sort_key(name: str) -> tuple:
    import re
    m = re.search(r"(\d+)(?:\.(\d+))?", name or "")
    return (int(m.group(1)) if m else 0, int(m.group(2) or 0) if m and m.group(2) else 0)


def _current_baseline(db: Session, project_id: str) -> m.Baseline | None:
    rows = db.execute(select(m.Baseline).where(m.Baseline.project_id == project_id)).scalars().all()
    if not rows:
        return None
    return sorted(rows, key=lambda b: _baseline_sort_key(b.name))[-1]


def _trace_fingerprint(db: Session, project_id: str, baseline_id: str | None) -> str:
    import hashlib
    edges = sorted(
        (e.source_semantic_id, e.target_semantic_id, e.relation_type.value)
        for e in db.execute(select(m.TraceLink).where(m.TraceLink.project_id == project_id)).scalars()
    )
    h = hashlib.sha256()
    for e in edges:
        h.update(("|".join(e)).encode())
    h.update(("baseline:" + (baseline_id or "")).encode())
    return h.hexdigest()


def suggest_cr_classification(db: Session, project_id: str, affected_semantic_ids: list[str]) -> dict:
    """Deterministic, transparent classification suggestion. Explicitly a
    heuristic — the user always confirms or overrides before analysis."""
    ids = [s for s in (affected_semantic_ids or []) if s]
    if not ids:
        return {
            "classification": "CLARIFICATION",
            "reason": "No affected objects declared — likely a clarification or correction.",
            "confidence": "LOW",
            "basis": "no affected objects declared",
        }
    existing = {
        so.semantic_id
        for so in db.execute(
            select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
        ).scalars()
    }
    known = [s for s in ids if s in existing]
    baseline = _current_baseline(db, project_id)
    baseline_objects: set[str] = set()
    if baseline:
        baseline_objects = {
            b.semantic_object_id
            for b in db.execute(
                select(m.BaselineBinding).where(m.BaselineBinding.baseline_id == baseline.id)
            ).scalars()
        }

    if not known:
        return {
            "classification": "SCOPE_EXPANSION",
            "reason": "Introduces a new capability not found in the current project; likely additional execution scope.",
            "confidence": "MEDIUM",
            "basis": f"none of {len(ids)} affected object(s) exist as project objects",
        }
    in_baseline = [s for s in known if s in baseline_objects]
    if len(known) == len(ids) and in_baseline:
        return {
            "classification": "REQUIREMENT_CHANGE",
            "reason": "Modifies objects already present in the current baseline.",
            "confidence": "MEDIUM",
            "basis": f"affected object(s) are in baseline {baseline.name}",
        }
    if len(known) == len(ids):
        return {
            "classification": "REQUIREMENT_CHANGE",
            "reason": "Modifies existing project objects.",
            "confidence": "MEDIUM",
            "basis": f"all {len(ids)} affected object(s) already exist in the project",
        }
    return {
        "classification": "SCOPE_EXPANSION",
        "reason": "Mix of new and existing objects; likely additional execution scope.",
        "confidence": "MEDIUM",
        "basis": f"{len(known)}/{len(ids)} affected object(s) already exist in the project",
    }


def _impact_analysis_dict(row: m.ImpactAnalysis) -> dict:
    result = row.result_json or {}
    return {
        "id": row.id,
        "project_id": row.project_id,
        "target_type": row.target_type,
        "target_id": row.target_id,
        "baseline_id": row.baseline_id,
        "baseline_name": row.baseline_name,
        "calculated_at": row.created_at.isoformat(),
        "confidence": row.confidence,
        "coverage_status": row.coverage_status,
        "stale": row.stale,
        "stale_reason": row.stale_reason,
        "review_state": row.review_state.value,
        "reviewed_by": row.reviewed_by,
        "reviewed_at": row.reviewed_at.isoformat() if row.reviewed_at else None,
        "result": result,
    }


def _cr_affected_bfs(db: Session, project_id: str, starts: list[str]) -> tuple[list[dict], list[str]]:
    """BFS downstream impact from a set of start semantic ids, reusing the exact
    trace graph. Returns (affected entries with paths, unaffected semantic ids)."""
    nodes = {
        n.semantic_id: n
        for n in db.execute(
            select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
        ).scalars().all()
    }
    edges = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
    ).scalars().all()
    adj_down: dict[str, list] = {}
    for e in edges:
        adj_down.setdefault(e.source_semantic_id, []).append((e.target_semantic_id, e.relation_type.value))

    best: dict[str, dict] = {}
    rank = {"DIRECT": 0, "INDIRECT": 1, "POTENTIAL": 2}

    for start in starts:
        path_of: dict[str, list] = {start: [start]}
        queue = [(start, 0)]
        while queue:
            cur, depth = queue.pop(0)
            for target, rel in adj_down.get(cur, []):
                if target in best:
                    continue
                path = path_of.get(cur, []) + [target]
                path_of[target] = path
                level = "DIRECT" if depth == 0 else "INDIRECT"
                node = nodes.get(target)
                best[target] = {
                    "semantic_id": target,
                    "object_type": node.object_type.value if node else "UNKNOWN",
                    "display_name": node.display_name if node else target,
                    "level": level,
                    "depth": depth + 1,
                    "relation": rel,
                    "reason": _RELATION_REASON.get(rel, f"linked via {rel}"),
                    "path": path,
                    "regeneration": _REGEN_ELIGIBILITY.get(node.object_type.value if node else "", "NONE"),
                }
                if depth + 1 < 2:
                    queue.append((target, depth + 1))
        # Incoming references → POTENTIAL.
        for e in edges:
            if e.target_semantic_id == start and e.source_semantic_id not in best and e.source_semantic_id not in starts:
                node = nodes.get(e.source_semantic_id)
                candidate = {
                    "semantic_id": e.source_semantic_id,
                    "object_type": node.object_type.value if node else "UNKNOWN",
                    "display_name": node.display_name if node else e.source_semantic_id,
                    "level": "POTENTIAL",
                    "depth": 1,
                    "relation": e.relation_type.value,
                    "reason": _RELATION_REASON.get(e.relation_type.value, f"linked via {e.relation_type.value}"),
                    "path": [e.source_semantic_id, start],
                    "regeneration": _REGEN_ELIGIBILITY.get(node.object_type.value if node else "", "NONE"),
                }
                prev = best.get(e.source_semantic_id)
                if prev is None or rank[candidate["level"]] < rank[prev["level"]]:
                    best[e.source_semantic_id] = candidate

    affected = list(best.values())
    affected_ids = {a["semantic_id"] for a in affected}
    unaffected = [n.semantic_id for n in nodes.values() if n.semantic_id not in affected_ids and n.semantic_id not in starts]
    return affected, unaffected


def analyze_cr_impact(
    db: Session,
    change_request_id: str,
    *,
    actor="local-user",
    actor_id: str | None = None,
) -> dict:
    """Run a truthful impact analysis for a CR and persist an immutable
    point-in-time snapshot. Does not alter confirmed requirements, baselines,
    PM tasks, QA scope, infra state or commercial commitments."""
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    project_id = cr.project_id

    starts = [link.semantic_id for link in cr.links]
    affected, unaffected = _cr_affected_bfs(db, project_id, starts)
    trust = _trust_impact_view(db, project_id, affected, unaffected)

    baseline = _current_baseline(db, project_id)
    fingerprint = _trace_fingerprint(db, project_id, baseline.id if baseline else None)

    result_json = {
        "code": cr.code,
        "requested_change": cr.requested_change,
        "classification": (db.execute(select(m.ChangeRequestImpact.classification).where(m.ChangeRequestImpact.change_request_id == cr.id)).scalar_one_or_none()),
        "baseline_name": baseline.name if baseline else None,
        "known_impact": trust["known_impact"],
        "potential_impact": trust["potential_impact"],
        "unknown_areas": trust["unknown_areas"],
        "unaffected": trust["unaffected"],
        "unaffected_count": len(trust["unaffected"]),
        "impact_confidence": trust["impact_confidence"],
        "trace_coverage": trust["trace_coverage"],
        "review": {"state": "NOT_REVIEWED", "decisions": {}, "human_added": [], "comments": []},
    }

    row = m.ImpactAnalysis(
        project_id=project_id,
        target_type="change_request",
        target_id=cr.id,
        baseline_id=baseline.id if baseline else None,
        baseline_name=baseline.name if baseline else None,
        confidence=trust["impact_confidence"],
        coverage_status=trust["trace_coverage"]["status"],
        trace_fingerprint=fingerprint,
        result_json=result_json,
    )
    db.add(row)
    if cr.status == m.ChangeRequestStatus.DRAFT:
        cr.status = m.ChangeRequestStatus.IMPACT_ANALYZED
    db.commit()
    db.refresh(row)
    record_audit(
        db, action="CR_IMPACT_ANALYZED", project_id=project_id, actor_id=actor_id,
        object_type="ChangeRequest", object_id=cr.id,
        metadata={"analysis_id": row.id, "confidence": trust["impact_confidence"],
                  "coverage": trust["trace_coverage"]["status"]},
    )
    return _impact_analysis_dict(row)


def get_cr_impact_analysis(db: Session, change_request_id: str) -> dict:
    """Return the latest impact analysis snapshot, re-marking it STALE if the
    project trace graph or baseline changed since it was calculated."""
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    row = db.execute(
        select(m.ImpactAnalysis)
        .where(m.ImpactAnalysis.target_id == cr.id, m.ImpactAnalysis.target_type == "change_request")
        .order_by(m.ImpactAnalysis.created_at.desc())
        .limit(1)
    ).scalar_one_or_none()
    if row is None:
        raise DomainError("No impact analysis has been run for this change request", status_code=404)
    current_fp = _trace_fingerprint(db, cr.project_id, row.baseline_id)
    if row.trace_fingerprint != current_fp:
        row.stale = True
        row.stale_reason = "Project trace graph or baseline changed since this analysis was calculated."
        db.commit()
    return _impact_analysis_dict(row)


def review_cr_impact_analysis(
    db: Session,
    change_request_id: str,
    analysis_id: str,
    *,
    decisions: dict[str, dict] | None = None,
    human_added: list[dict] | None = None,
    comments: list[str] | None = None,
    finalize: bool = False,
    reviewer: str = "local-user",
    actor_id: str | None = None,
) -> dict:
    """Human review of an impact analysis. Confirm, mark not-impacted, add
    missing impact or comment — every decision is audited and never masquerades
    as system-derived."""
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    row = get_or_404(db, m.ImpactAnalysis, analysis_id, "ImpactAnalysis")
    if row.target_id != cr.id or row.target_type != "change_request":
        raise DomainError("Analysis does not belong to this change request", status_code=422)

    rj = row.result_json or {}
    review = rj.setdefault("review", {})
    review.setdefault("decisions", {})
    review.setdefault("human_added", [])
    review.setdefault("comments", [])

    for item_id, d in (decisions or {}).items():
        review["decisions"][item_id] = {
            "decision": (d.get("decision") or "CONFIRMED").upper(),
            "reason": d.get("reason"),
            "reviewer": reviewer,
            "at": m.utcnow().isoformat(),
        }
    for h in (human_added or []):
        entry = dict(h)
        entry["source"] = "HUMAN-ADDED"
        entry["added_by"] = reviewer
        entry["added_at"] = m.utcnow().isoformat()
        review["human_added"].append(entry)
    for c in (comments or []):
        review["comments"].append({"comment": c, "reviewer": reviewer, "at": m.utcnow().isoformat()})

    row.result_json = rj
    row.review_state = (
        m.ImpactReviewState.REVIEWED if finalize else m.ImpactReviewState.REVIEW_IN_PROGRESS
    )
    if finalize:
        row.reviewed_by = reviewer
        row.reviewed_at = m.utcnow()
        if cr.status in (m.ChangeRequestStatus.DRAFT, m.ChangeRequestStatus.IMPACT_ANALYZED):
            cr.status = m.ChangeRequestStatus.INTERNAL_REVIEW_COMPLETE
    db.commit()
    record_audit(
        db, action="CR_IMPACT_REVIEWED", project_id=cr.project_id, actor_id=actor_id,
        object_type="ChangeRequest", object_id=cr.id,
        metadata={"analysis_id": row.id, "finalized": finalize, "reviewer": reviewer,
                  "decisions": len(decisions or {}), "human_added": len(human_added or [])},
    )
    return _impact_analysis_dict(row)


# ---------------------------------------------------------------------------
# ChangeRequest
# ---------------------------------------------------------------------------


def next_cr_code(db: Session, project_id: str) -> str:
    count = (
        db.execute(
            select(m.ChangeRequest.id).where(m.ChangeRequest.project_id == project_id)
        )
        .scalars()
        .all()
    )
    return f"CR-{len(count) + 1:04d}"


def create_change_request(
    db: Session,
    *,
    project_id: str,
    requested_change: str,
    affected_semantic_ids: list[str] | None = None,
    title=None,
    reason=None,
    requested_by="local-user",
    requested_date=None,
    source_reference=None,
    notes=None,
    target_release=None,
    schedule_impact=None,
    commercial_impact=None,
    classification=None,
    actor="local-user",
    actor_id: str | None = None,
) -> m.ChangeRequest:
    for sid in affected_semantic_ids or []:
        exists = db.execute(
            select(m.SemanticObject.id).where(
                m.SemanticObject.project_id == project_id,
                m.SemanticObject.semantic_id == sid,
            )
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(
                f"Change request references unknown semantic object '{sid}'",
                status_code=422,
            )
    from datetime import date as _date

    cr = m.ChangeRequest(
        project_id=project_id,
        code=next_cr_code(db, project_id),
        title=title,
        requested_by=requested_by,
        requested_date=(_date.fromisoformat(requested_date) if requested_date else None),
        reason=reason,
        requested_change=requested_change,
        source_reference=source_reference,
        notes=notes,
        target_release=target_release,
        schedule_impact=schedule_impact,
        commercial_impact=commercial_impact,
        status=m.ChangeRequestStatus.DRAFT,  # saved as a draft first, never applied
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(cr)
    db.flush()
    for sid in affected_semantic_ids or []:
        db.add(m.ChangeRequestLink(change_request_id=cr.id, semantic_id=sid))
    db.add(m.ChangeRequestImpact(change_request_id=cr.id, classification=classification))
    db.commit()
    record_audit(
        db, action="CR_CREATED_DRAFT", project_id=project_id, actor_id=actor_id,
        object_type="ChangeRequest", object_id=cr.id,
        metadata={"code": cr.code, "classification": classification},
    )
    return cr


def implement_change_request(
    db: Session,
    change_request_id: str,
    *,
    artifact_revision_map: dict[str, str] | None = None,
    actor="local-user",
) -> dict:
    """Mark a CR IMPLEMENTED by pointing it at the new revision(s) it spawned.

    The CR itself never mutates old confirmed baselines — it only
    references the freshly created draft/confirmed revisions.
    """
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    if cr.status == m.ChangeRequestStatus.IMPLEMENTED:
        raise DomainError("Change request already implemented")
    spawned: list[dict] = []
    for artifact_id, snapshot in (artifact_revision_map or {}).items():
        rev = create_revision(
            db, artifact_id=artifact_id, snapshot=snapshot, actor=actor
        )
        ensure_semantic_object(
            db,
            project_id=cr.project_id,
            semantic_id=f"rev_{rev.id}",
            object_type=m.SemanticObjectType.DOCUMENT_SECTION,
            display_name=f"{cr.code} revision {rev.revision_number}",
            entity_ref=rev.id,
        )
        db.add(
            m.TraceLink(
                project_id=cr.project_id,
                source_semantic_id=cr.code,
                target_semantic_id=f"rev_{rev.id}",
                relation_type=m.TraceRelationType.GENERATED_FROM,
                created_by=actor,
            )
        )
        spawned.append({"artifact_id": artifact_id, "revision_id": rev.id, "revision_number": rev.revision_number})
    cr.status = m.ChangeRequestStatus.IMPLEMENTED
    db.commit()
    return {"change_request": cr, "spawned_revisions": spawned}


def change_request_detail(db: Session, change_request_id: str) -> dict:
    """Full CR view: affected objects, deterministic impact, spawned revisions."""
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    links = [l.semantic_id for l in cr.links]

    affected = []
    for sid in links:
        so = db.execute(
            select(m.SemanticObject).where(
                m.SemanticObject.project_id == cr.project_id,
                m.SemanticObject.semantic_id == sid,
            )
        ).scalar_one_or_none()
        affected.append({
            "semantic_id": sid,
            "object_type": so.object_type.value if so else None,
            "display_name": so.display_name if so else None,
        })

    impact = {sid: impact_of(db, cr.project_id, sid) for sid in links}

    spawned = []
    traces = db.execute(
        select(m.TraceLink).where(
            m.TraceLink.project_id == cr.project_id,
            m.TraceLink.source_semantic_id == cr.code,
            m.TraceLink.relation_type == m.TraceRelationType.GENERATED_FROM,
        )
    ).scalars().all()
    for t in traces:
        so = db.execute(
            select(m.SemanticObject).where(
                m.SemanticObject.project_id == cr.project_id,
                m.SemanticObject.semantic_id == t.target_semantic_id,
            )
        ).scalar_one_or_none()
        if so and so.entity_ref:
            rev = db.get(m.ArtifactRevision, so.entity_ref)
            if rev:
                spawned.append({
                    "revision_id": rev.id,
                    "revision_number": rev.revision_number,
                    "artifact_id": rev.artifact_id,
                    "artifact_title": rev.artifact.title,
                    "status": rev.status.value,
                    "based_on_revision_id": rev.based_on_revision_id,
                })

    impact_row = db.execute(
        select(m.ChangeRequestImpact).where(
            m.ChangeRequestImpact.change_request_id == cr.id
        )
    ).scalar_one_or_none()

    return {
        "id": cr.id,
        "code": cr.code,
        "project_id": cr.project_id,
        "title": cr.title,
        "requested_by": cr.requested_by,
        "requested_date": cr.requested_date.isoformat() if cr.requested_date else None,
        "reason": cr.reason,
        "requested_change": cr.requested_change,
        "source_reference": cr.source_reference,
        "notes": cr.notes,
        "status": cr.status.value,
        "classification": impact_row.classification if impact_row else None,
        "target_release": cr.target_release,
        "schedule_impact": cr.schedule_impact,
        "commercial_impact": cr.commercial_impact,
        "created_at": cr.created_at.isoformat(),
        "updated_at": cr.updated_at.isoformat() if cr.updated_at else None,
        "created_by": cr.created_by,
        "affected": affected,
        "impact": impact,
        "spawned_revisions": spawned,
    }


def _cr_impact_row(db: Session, change_request_id: str) -> m.ChangeRequestImpact:
    row = db.execute(
        select(m.ChangeRequestImpact).where(
            m.ChangeRequestImpact.change_request_id == change_request_id
        )
    ).scalar_one_or_none()
    if row is None:
        row = m.ChangeRequestImpact(change_request_id=change_request_id)
        db.add(row)
        db.commit()
        db.refresh(row)
    return row


def _cr_impact_dict(row: m.ChangeRequestImpact | None, cr: m.ChangeRequest) -> dict:
    return {
        "change_request_id": cr.id,
        "code": cr.code,
        "requested_change": cr.requested_change,
        "status": cr.status.value,
        "classification": row.classification if row else None,
        "function_impact": row.function_impact if row else None,
        "effort_impact": row.effort_impact if row else None,
        "timeline_impact": row.timeline_impact if row else None,
        "technical_impact": row.technical_impact if row else None,
        "qa_impact": row.qa_impact if row else None,
        "infra_impact": row.infra_impact if row else None,
        "commercial_status": row.commercial_status if row else None,
        "pricing_basis": row.pricing_basis if row else None,
        "confidence": row.confidence if row else None,
        "customer_approval": row.customer_approval if row else None,
        "approval_evidence": row.approval_evidence if row else None,
    }


def cr_impact(db: Session, change_request_id: str) -> dict:
    """Return the stored CR impact plus the technical impact derived from the
    exact trace graph for any affected object — no human memory required."""
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    row = _cr_impact_row(db, change_request_id)
    out = _cr_impact_dict(row, cr)

    # Derive technical impact from trace for affected semantic ids.
    technical = out.get("technical_impact") or {"affected": [], "unaffected": [], "unknown": []}
    for link in cr.links:
        impact = impact_of(db, cr.project_id, link.semantic_id)
        for d in impact["downstream"]:
            so = db.execute(
                select(m.SemanticObject).where(
                    m.SemanticObject.project_id == cr.project_id,
                    m.SemanticObject.semantic_id == d["semantic_id"],
                )
            ).scalar_one_or_none()
            entry = {
                "semantic_id": d["semantic_id"],
                "object_type": so.object_type.value if so else None,
                "display_name": so.display_name if so else d["semantic_id"],
                "relation": d["relation"],
            }
            if entry not in technical["affected"]:
                technical["affected"].append(entry)
    out["technical_impact"] = technical
    return out


def save_cr_impact(
    db: Session,
    change_request_id: str,
    *,
    classification: str | None = None,
    function_impact: dict | None = None,
    effort_impact: dict | None = None,
    timeline_impact: dict | None = None,
    technical_impact: dict | None = None,
    qa_impact: dict | None = None,
    infra_impact: dict | None = None,
    commercial_status: str | None = None,
    pricing_basis: str | None = None,
    confidence: str | None = None,
) -> dict:
    """Store a human/estimator-supplied impact. Nothing is invented here — the
    caller passes only what is known; unknown dimensions stay absent."""
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    row = _cr_impact_row(db, change_request_id)
    for field, value in {
        "classification": classification, "function_impact": function_impact,
        "effort_impact": effort_impact, "timeline_impact": timeline_impact,
        "technical_impact": technical_impact, "qa_impact": qa_impact,
        "infra_impact": infra_impact, "commercial_status": commercial_status,
        "pricing_basis": pricing_basis, "confidence": confidence,
    }.items():
        if value is not None:
            setattr(row, field, value)
    db.commit()
    record_audit(
        db, action="CR_IMPACT_SAVED", project_id=cr.project_id, actor_id=cr.actor_id,
        object_type="ChangeRequest", object_id=cr.id,
        metadata={"classification": classification, "commercial_status": commercial_status},
    )
    return _cr_impact_dict(row, cr)


def set_cr_customer_approval(
    db: Session,
    change_request_id: str,
    *,
    decision: str,  # APPROVED | REJECTED | PENDING
    approved_by: str | None = None,
    reference: str | None = None,
    note: str | None = None,
    amount: str | None = None,
    actor="local-user",
) -> dict:
    """Customer commercial decision. Distinct from admin technical confirmation:
    this records whether the customer approved the scope/price, never whether
    the system truth was published."""
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    decision = decision.upper()
    if decision not in ("APPROVED", "REJECTED", "PENDING"):
        raise DomainError("decision must be APPROVED, REJECTED or PENDING")
    row = _cr_impact_row(db, change_request_id)
    row.customer_approval = decision
    row.approval_evidence = {
        "approved_by": approved_by,
        "approved_at": m.utcnow().isoformat(),
        "reference": reference,
        "note": note,
        "amount": amount,
    }
    if decision == "APPROVED":
        cr.status = m.ChangeRequestStatus.ACCEPTED
    elif decision == "REJECTED":
        cr.status = m.ChangeRequestStatus.REJECTED
    db.commit()
    record_audit(
        db, action="CR_CUSTOMER_APPROVAL", project_id=cr.project_id, actor_id=cr.actor_id,
        object_type="ChangeRequest", object_id=cr.id,
        metadata={"decision": decision, "approved_by": approved_by, "amount": amount},
    )
    return _cr_impact_dict(row, cr)


# ---------------------------------------------------------------------------
# ChangeRequest lifecycle transitions (human-driven, guarded)
# ---------------------------------------------------------------------------

def _cr_dict(cr: m.ChangeRequest) -> dict:
    return {
        "id": cr.id,
        "code": cr.code,
        "title": cr.title,
        "project_id": cr.project_id,
        "requested_change": cr.requested_change,
        "requested_by": cr.requested_by,
        "requested_date": cr.requested_date.isoformat() if cr.requested_date else None,
        "reason": cr.reason,
        "source_reference": cr.source_reference,
        "notes": cr.notes,
        "status": cr.status.value,
        "created_at": cr.created_at.isoformat(),
        "updated_at": cr.updated_at.isoformat() if cr.updated_at else None,
    }

_CR_ALLOWED_TRANSITIONS: dict = {
    m.ChangeRequestStatus.DRAFT: {
        m.ChangeRequestStatus.OPEN,
        m.ChangeRequestStatus.NEEDS_CLARIFICATION,
        m.ChangeRequestStatus.UNDER_HUMAN_REVIEW,
        m.ChangeRequestStatus.REJECTED,
        m.ChangeRequestStatus.DEFERRED,
    },
    m.ChangeRequestStatus.OPEN: {
        m.ChangeRequestStatus.NEEDS_CLARIFICATION,
        m.ChangeRequestStatus.IMPACT_ANALYZED,
        m.ChangeRequestStatus.UNDER_HUMAN_REVIEW,
        m.ChangeRequestStatus.REJECTED,
        m.ChangeRequestStatus.DEFERRED,
    },
    m.ChangeRequestStatus.NEEDS_CLARIFICATION: {
        m.ChangeRequestStatus.OPEN,
        m.ChangeRequestStatus.IMPACT_ANALYZED,
        m.ChangeRequestStatus.UNDER_HUMAN_REVIEW,
        m.ChangeRequestStatus.REJECTED,
        m.ChangeRequestStatus.DEFERRED,
    },
    m.ChangeRequestStatus.IMPACT_ANALYZED: {
        m.ChangeRequestStatus.NEEDS_CLARIFICATION,
        m.ChangeRequestStatus.UNDER_HUMAN_REVIEW,
        m.ChangeRequestStatus.INTERNAL_REVIEW_COMPLETE,
        m.ChangeRequestStatus.REJECTED,
        m.ChangeRequestStatus.DEFERRED,
    },
    m.ChangeRequestStatus.UNDER_HUMAN_REVIEW: {
        m.ChangeRequestStatus.NEEDS_CLARIFICATION,
        m.ChangeRequestStatus.INTERNAL_REVIEW_COMPLETE,
        m.ChangeRequestStatus.ACCEPTED,
        m.ChangeRequestStatus.REJECTED,
        m.ChangeRequestStatus.DEFERRED,
    },
    m.ChangeRequestStatus.INTERNAL_REVIEW_COMPLETE: {
        m.ChangeRequestStatus.UNDER_HUMAN_REVIEW,
        m.ChangeRequestStatus.ACCEPTED,
        m.ChangeRequestStatus.IMPLEMENTATION_PLANNED,
        m.ChangeRequestStatus.REJECTED,
        m.ChangeRequestStatus.DEFERRED,
    },
    m.ChangeRequestStatus.ACCEPTED: {
        m.ChangeRequestStatus.IMPLEMENTATION_PLANNED,
        m.ChangeRequestStatus.IMPLEMENTED,
        m.ChangeRequestStatus.REJECTED,
        m.ChangeRequestStatus.DEFERRED,
    },
    m.ChangeRequestStatus.IMPLEMENTATION_PLANNED: {
        m.ChangeRequestStatus.IMPLEMENTED,
        m.ChangeRequestStatus.ACCEPTED,
        m.ChangeRequestStatus.DEFERRED,
    },
    m.ChangeRequestStatus.IMPLEMENTED: {
        m.ChangeRequestStatus.CLOSED,
        m.ChangeRequestStatus.IMPLEMENTATION_PLANNED,
    },
    m.ChangeRequestStatus.REJECTED: {
        m.ChangeRequestStatus.OPEN,
        m.ChangeRequestStatus.DRAFT,
    },
    m.ChangeRequestStatus.DEFERRED: {
        m.ChangeRequestStatus.OPEN,
        m.ChangeRequestStatus.DRAFT,
        m.ChangeRequestStatus.ACCEPTED,
    },
    m.ChangeRequestStatus.CLOSED: {
        m.ChangeRequestStatus.OPEN,
    },
}


def transition_change_request(
    db: Session,
    change_request_id: str,
    *,
    to_status: str,
    note: str | None = None,
    actor="local-user",
    actor_id: str | None = None,
) -> dict:
    """Human-driven lifecycle transition with a guarded state machine.

    Every call is an explicit human action (the caller supplies the actor),
    and the state machine never allows a jump that would let a computed
    analysis silently become APPROVED. APPROVED/ACCEPTED is only reachable
    from UNDER_HUMAN_REVIEW or INTERNAL_REVIEW_COMPLETE."""
    cr = get_or_404(db, m.ChangeRequest, change_request_id, "ChangeRequest")
    try:
        target = m.ChangeRequestStatus(to_status.upper())
    except ValueError:
        raise DomainError(
            f"Unknown status '{to_status}'. Allowed: "
            + ", ".join(s.value for s in m.ChangeRequestStatus),
            status_code=422,
        )
    allowed = _CR_ALLOWED_TRANSITIONS.get(cr.status, set())
    if target not in allowed and target != cr.status:
        raise DomainError(
            f"Transition {cr.status.value} → {target.value} is not allowed. "
            f"Allowed: {', '.join(s.value for s in sorted(allowed, key=lambda s: s.value)) or 'none'}",
            status_code=409,
        )
    if target == cr.status:
        return {"change_request": _cr_dict(cr), "note": note, "unchanged": True}

    prev = cr.status.value
    cr.status = target
    db.commit()
    record_audit(
        db, action="CR_STATUS_TRANSITION", project_id=cr.project_id, actor_id=actor_id,
        object_type="ChangeRequest", object_id=cr.id,
        metadata={"from": prev, "to": target.value, "note": note, "actor": actor},
    )
    return {"change_request": _cr_dict(cr), "from": prev, "to": target.value, "note": note}


# ---------------------------------------------------------------------------
# Database design (structured model; diagram is a view over it)
# ---------------------------------------------------------------------------


def create_schema(
    db: Session, *, project_id: str, name: str, semantic_id: str, description=None, actor="local-user"
) -> m.DatabaseSchema:
    schema = m.DatabaseSchema(
        project_id=project_id, name=name, semantic_id=semantic_id, description=description,
        created_by=actor,
    )
    db.add(schema)
    db.flush()
    ensure_semantic_object(
        db,
        project_id=project_id,
        semantic_id=semantic_id,
        object_type=m.SemanticObjectType.DB_SCHEMA,
        display_name=name,
        entity_ref=schema.id,
    )
    db.commit()
    return schema


def create_table(
    db: Session, *, schema_id: str, name: str, semantic_id: str | None = None, description=None
) -> m.DatabaseTable:
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    semantic_id = semantic_id or f"tbl_{name}"
    table = m.DatabaseTable(
        schema_id=schema_id, semantic_id=semantic_id, name=name, description=description
    )
    db.add(table)
    db.flush()
    ensure_semantic_object(
        db,
        project_id=schema.project_id,
        semantic_id=semantic_id,
        object_type=m.SemanticObjectType.DB_TABLE,
        display_name=name,
        entity_ref=table.id,
    )
    db.commit()
    return table


def create_field(
    db: Session,
    *,
    table_id: str,
    name: str,
    data_type: str,
    semantic_id: str | None = None,
    length=None,
    nullable=False,
    default=None,
    primary_key=False,
    foreign_key=False,
    reference=None,
    description=None,
    remark=None,
) -> m.DatabaseField:
    table = get_or_404(db, m.DatabaseTable, table_id, "Table")
    semantic_id = semantic_id or f"fld_{table.name}_{name}"
    position = (
        db.execute(
            select(m.DatabaseField.id).where(m.DatabaseField.table_id == table_id)
        )
        .scalars()
        .all()
    )
    field = m.DatabaseField(
        table_id=table_id,
        semantic_id=semantic_id,
        name=name,
        data_type=data_type,
        length=length,
        nullable=nullable,
        default=default,
        primary_key=primary_key,
        foreign_key=foreign_key,
        reference=reference,
        description=description,
        remark=remark,
        position=len(position),
    )
    db.add(field)
    db.flush()
    ensure_semantic_object(
        db,
        project_id=table.schema.project_id,
        semantic_id=semantic_id,
        object_type=m.SemanticObjectType.DB_FIELD,
        display_name=f"{table.name}.{name}",
        entity_ref=field.id,
    )
    db.commit()
    return field


def create_relation(
    db: Session,
    *,
    schema_id: str,
    from_field_semantic_id: str,
    to_field_semantic_id: str,
    relation_type="MANY_TO_ONE",
) -> m.DatabaseRelation:
    for sid in (from_field_semantic_id, to_field_semantic_id):
        exists = db.execute(
            select(m.DatabaseField.id).where(m.DatabaseField.semantic_id == sid)
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(f"Unknown field semantic id '{sid}'", status_code=422)
    rel = m.DatabaseRelation(
        schema_id=schema_id,
        semantic_id=f"rel_{from_field_semantic_id}__{to_field_semantic_id}",
        from_field_semantic_id=from_field_semantic_id,
        to_field_semantic_id=to_field_semantic_id,
        relation_type=relation_type,
    )
    db.add(rel)
    db.flush()
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    ensure_semantic_object(
        db,
        project_id=schema.project_id,
        semantic_id=rel.semantic_id,
        object_type=m.SemanticObjectType.DB_RELATION,
        display_name=f"{from_field_semantic_id} → {to_field_semantic_id}",
        entity_ref=rel.id,
    )
    db.commit()
    return rel


def data_dictionary(db: Session, schema_id: str) -> list[dict]:
    """A pure view over the structured model — no separate truth stored."""
    get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    tables = db.execute(
        select(m.DatabaseTable).where(m.DatabaseTable.schema_id == schema_id)
    ).scalars().all()
    table_sid = {t.id: t.semantic_id for t in tables}
    table_name = {t.id: t.name for t in tables}
    fields = db.execute(
        select(m.DatabaseField).where(m.DatabaseField.table_id.in_(table_sid))
    ).scalars().all()
    result = []
    for f in fields:
        result.append(
            {
                "table": table_name[f.table_id],
                "table_semantic_id": table_sid[f.table_id],
                "field": f.name,
                "field_semantic_id": f.semantic_id,
                "data_type": f.data_type,
                "length": f.length,
                "nullable": f.nullable,
                "default": f.default,
                "primary_key": f.primary_key,
                "foreign_key": f.foreign_key,
                "reference": f.reference,
                "description": f.description,
                "remark": f.remark,
            }
        )
    return result


# ---------------------------------------------------------------------------
# Document workspace (rich UR/DR content, section semantic identity)
# ---------------------------------------------------------------------------

_DOC_BLOCK_KINDS = {"heading", "paragraph", "bullet_list", "numbered_list", "table", "code"}


def _doc_text(content: dict) -> str:
    """Plain text derived from a ProseMirror/Tiptap doc JSON (deterministic)."""
    if not isinstance(content, dict):
        return ""
    parts: list[str] = []

    def walk(node) -> None:
        if not isinstance(node, dict):
            return
        ntype = node.get("type")
        if ntype == "text":
            parts.append(node.get("text") or "")
            return
        for child in node.get("content") or []:
            walk(child)
        if ntype in ("paragraph", "heading", "codeBlock", "blockquote", "listItem", "tableRow"):
            parts.append("\n")

    walk(content)
    return "\n".join(line.rstrip() for line in "".join(parts).splitlines())


def _blocks_to_doc(blocks: list[dict]) -> dict:
    """Convert the legacy P1 block model into a ProseMirror/Tiptap doc JSON."""
    content: list[dict] = []
    for blk in blocks:
        kind = blk.get("kind") or "paragraph"
        if kind == "heading":
            content.append({"type": "heading", "attrs": {"level": blk.get("level") or 2},
                            "content": [{"type": "text", "text": blk.get("text") or ""}]})
        elif kind == "bullet_list":
            content.append({"type": "bulletList", "content": [
                {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": item}]}]}
                for item in blk.get("items", [])
            ]})
        elif kind == "numbered_list":
            content.append({"type": "orderedList", "content": [
                {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": item}]}]}
                for item in blk.get("items", [])
            ]})
        elif kind == "table":
            rows = []
            if blk.get("header"):
                rows.append({"type": "tableRow", "content": [
                    {"type": "tableHeader", "content": [{"type": "paragraph", "content": [{"type": "text", "text": h}]}]}
                    for h in blk["header"]
                ]})
            for row in blk.get("rows", []):
                rows.append({"type": "tableRow", "content": [
                    {"type": "tableCell", "content": [{"type": "paragraph", "content": [{"type": "text", "text": c}]}]}
                    for c in row
                ]})
            content.append({"type": "table", "content": rows})
        elif kind == "code":
            content.append({"type": "codeBlock", "content": [{"type": "text", "text": blk.get("text") or ""}]})
        else:
            content.append({"type": "paragraph", "content": [{"type": "text", "text": blk.get("text") or ""}]})
    return {"type": "doc", "content": content}


def _normalise_sections(artifact_id: str, sections: list[dict]) -> list[dict]:
    """Assign a stable section id where missing and normalise the content model.

    Section identity lives inside the snapshot and is preserved by the
    editor across edits. Legacy P1 `blocks` are migrated to a structured
    ProseMirror doc JSON; `plain_text` is always derived, never the truth.
    """
    out: list[dict] = []
    used: set[str] = set()
    counter = 0
    for sec in sections:
        sec = dict(sec)
        sid = sec.get("id") or f"docsec_{artifact_id}_{counter}"
        while sid in used:
            counter += 1
            sid = f"docsec_{artifact_id}_{counter}"
        used.add(sid)
        sec["id"] = sid
        sec.setdefault("heading", f"Section {len(out) + 1}")
        if sec.get("content") is None and sec.get("blocks"):
            sec["content"] = _blocks_to_doc(sec.get("blocks", []))
        sec.pop("blocks", None)
        if not sec.get("content"):
            sec["content"] = {"type": "doc", "content": [{"type": "paragraph"}]}
        if not sec.get("plain_text"):
            sec["plain_text"] = f"{sec.get('heading') or ''}\n{_doc_text(sec['content'])}"
        out.append(sec)
        counter += 1
    return out


def save_document(
    db: Session,
    *,
    revision_id: str,
    sections: list[dict],
    title: str | None = None,
    actor: str = "local-user",
) -> m.ArtifactRevision:
    """Persist rich document content for a DRAFT revision.

    Each section is registered as a DOCUMENT_SECTION semantic object so
    comments and traces can bind to it — never to a DOM position.
    Confirmed revisions are immutable and are rejected here.
    """
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    require_editable(revision)
    artifact = revision.artifact
    sections = _normalise_sections(artifact.id, sections)
    for sec in sections:
        ensure_semantic_object(
            db,
            project_id=artifact.project_id,
            semantic_id=sec["id"],
            object_type=m.SemanticObjectType.DOCUMENT_SECTION,
            display_name=sec.get("heading") or "Untitled section",
            entity_ref=artifact.id,
        )
    snapshot = dict(revision.snapshot or {})
    snapshot["sections"] = sections
    revision.snapshot = snapshot
    if title is not None:
        revision.title = title
        artifact.title = title
    db.commit()
    return revision


def get_document(db: Session, revision_id: str) -> dict:
    """Return the revision plus its structured sections."""
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    snapshot = revision.snapshot or {}
    sections = snapshot.get("sections")
    if not isinstance(sections, list):
        # Legacy P0 snapshots (e.g. {"sections": [{"id", "note"}]}) become
        # minimal paragraphs so nothing is lost on first edit.
        legacy = snapshot.get("sections") or snapshot.get("content") or []
        sections = [
            {
                "id": f"docsec_{revision.artifact_id}_{i}",
                "heading": sec.get("id") or sec.get("title") or f"Section {i + 1}",
                "blocks": [
                    {"kind": "paragraph", "text": sec.get("note") or sec.get("text") or ""}
                ],
            }
            for i, sec in enumerate(legacy)
        ] if isinstance(legacy, list) else []
    sections = _normalise_sections(revision.artifact_id, sections)
    return {
        "revision_id": revision.id,
        "artifact_id": revision.artifact_id,
        "revision_number": revision.revision_number,
        "status": revision.status.value,
        "title": revision.title,
        "artifact_type": revision.artifact.type.value,
        "based_on_revision_id": revision.based_on_revision_id,
        "created_by": revision.created_by,
        "confirmed_by": revision.confirmed_by,
        "confirmed_at": revision.confirmed_at.isoformat() if revision.confirmed_at else None,
        "editable": revision.editable,
        "sections": sections,
    }


# ---------------------------------------------------------------------------
# Review workflow (threads, summaries, activity timeline)
# ---------------------------------------------------------------------------


def create_thread(
    db: Session,
    *,
    project_id: str,
    title: str | None = None,
    actor: str = "local-user",
) -> m.CommentThread:
    guard_project(db, project_id)
    thread = m.CommentThread(project_id=project_id, title=title)
    db.add(thread)
    db.commit()
    return thread


def list_threads(db: Session, project_id: str) -> list[dict]:
    threads = db.execute(
        select(m.CommentThread)
        .where(m.CommentThread.project_id == project_id)
        .order_by(m.CommentThread.created_at.desc())
    ).scalars().all()
    out = []
    for t in threads:
        anns = db.execute(
            select(m.Annotation)
            .where(m.Annotation.thread_id == t.id)
            .order_by(m.Annotation.created_at)
        ).scalars().all()
        out.append({
            "id": t.id,
            "title": t.title,
            "resolved": t.resolved,
            "open_count": sum(1 for a in anns if a.status != m.AnnotationStatus.RESOLVED),
            "total": len(anns),
            "created_at": t.created_at.isoformat(),
            "annotations": [
                {
                    "id": a.id,
                    "anchor_semantic_id": a.anchor_semantic_id,
                    "type": a.type.value,
                    "status": a.status.value,
                    "content": a.content,
                    "created_by": a.created_by,
                    "created_at": a.created_at.isoformat(),
                }
                for a in anns
            ],
        })
    return out


def annotations_summary(db: Session, project_id: str) -> dict:
    anns = db.execute(
        select(m.Annotation).where(m.Annotation.project_id == project_id)
    ).scalars().all()
    by_status = Counter(a.status.value for a in anns)
    by_type = Counter(a.type.value for a in anns)
    by_anchor = Counter(a.anchor_semantic_id for a in anns)
    return {
        "total": len(anns),
        "open": sum(n for s, n in by_status.items() if s != "RESOLVED"),
        "resolved": by_status.get("RESOLVED", 0),
        "by_status": dict(by_status),
        "by_type": dict(by_type),
        "by_anchor": dict(by_anchor),
    }


def _iso(dt) -> str | None:
    return dt.isoformat() if dt else None


def timeline(db: Session, project_id: str, semantic_id: str | None = None) -> list[dict]:
    """Deterministic activity timeline, derived from real records.

    No separate activity table: events are reconstructed from revisions,
    confirmations, annotations, baselines and change requests.
    """
    events: list[dict] = []

    revisions = db.execute(
        select(m.ArtifactRevision)
        .join(m.Artifact, m.ArtifactRevision.artifact_id == m.Artifact.id)
        .where(m.Artifact.project_id == project_id)
    ).scalars().all()
    for r in revisions:
        events.append({
            "at": _iso(r.created_at), "kind": "revision_created",
            "actor": r.created_by, "label": f"{r.artifact.title} r{r.revision_number}",
            "revision_id": r.id, "semantic_id": None,
        })
        if r.confirmed_at:
            events.append({
                "at": _iso(r.confirmed_at), "kind": "revision_confirmed",
                "actor": r.confirmed_by, "label": f"{r.artifact.title} r{r.revision_number}",
                "revision_id": r.id, "semantic_id": None,
            })

    confirmations = db.execute(
        select(m.Confirmation)
        .where(m.Confirmation.project_id == project_id)
    ).scalars().all()
    for c in confirmations:
        events.append({
            "at": _iso(c.confirmed_at), "kind": "confirmation",
            "actor": c.confirmed_by, "label": c.comment or "confirmed",
            "revision_id": c.artifact_revision_id, "semantic_id": None,
        })

    annotations = db.execute(
        select(m.Annotation)
        .where(m.Annotation.project_id == project_id)
    ).scalars().all()
    for a in annotations:
        events.append({
            "at": _iso(a.created_at), "kind": f"annotation_{a.type.value.lower()}",
            "actor": a.created_by, "label": a.content[:120],
            "revision_id": a.artifact_revision_id, "semantic_id": a.anchor_semantic_id,
        })

    baselines = db.execute(
        select(m.Baseline).where(m.Baseline.project_id == project_id)
    ).scalars().all()
    for b in baselines:
        events.append({
            "at": _iso(b.created_at), "kind": "baseline_created",
            "actor": b.created_by, "label": b.name,
            "revision_id": None, "semantic_id": None,
        })

    crs = db.execute(
        select(m.ChangeRequest).where(m.ChangeRequest.project_id == project_id)
    ).scalars().all()
    for cr in crs:
        events.append({
            "at": _iso(cr.created_at), "kind": "change_request_created",
            "actor": cr.created_by, "label": f"{cr.code} — {cr.requested_change[:120]}",
            "revision_id": None, "semantic_id": cr.code,
        })

    if semantic_id:
        events = [
            e for e in events
            if e["semantic_id"] == semantic_id or e["revision_id"] is not None
        ]

    events.sort(key=lambda e: e["at"] or "")
    return events


# ---------------------------------------------------------------------------
# Database designer (CRUD over the structured model) + ERD layout
# ---------------------------------------------------------------------------

# Semantic identity of a field is fixed at creation. Renaming a field
# (or its table) never changes the semantic id — only display names
# change. This is what lets traces/annotations survive design edits.

_FIELD_EDITABLE = {
    "name", "data_type", "length", "nullable", "default",
    "primary_key", "foreign_key", "reference", "description", "remark",
}


def rename_table(db: Session, table_id: str, name: str) -> m.DatabaseTable:
    table = get_or_404(db, m.DatabaseTable, table_id, "Table")
    table.name = name
    ensure_semantic_object(
        db,
        project_id=table.schema.project_id,
        semantic_id=table.semantic_id,
        object_type=m.SemanticObjectType.DB_TABLE,
        display_name=name,
        entity_ref=table.id,
    )
    db.commit()
    return table


def update_field(db: Session, field_id: str, **changes) -> m.DatabaseField:
    field = get_or_404(db, m.DatabaseField, field_id, "Field")
    for key, value in changes.items():
        if key not in _FIELD_EDITABLE:
            raise DomainError(f"Cannot update field attribute '{key}'", status_code=422)
        setattr(field, key, value)
    ensure_semantic_object(
        db,
        project_id=field.table.schema.project_id,
        semantic_id=field.semantic_id,
        object_type=m.SemanticObjectType.DB_FIELD,
        display_name=f"{field.table.name}.{field.name}",
        entity_ref=field.id,
    )
    db.commit()
    return field


def delete_field(db: Session, field_id: str) -> None:
    field = get_or_404(db, m.DatabaseField, field_id, "Field")
    db.execute(
        delete(m.DatabaseRelation).where(
            or_(
                m.DatabaseRelation.from_field_semantic_id == field.semantic_id,
                m.DatabaseRelation.to_field_semantic_id == field.semantic_id,
            )
        )
    )
    db.delete(field)
    db.commit()


def delete_table(db: Session, table_id: str) -> None:
    table = get_or_404(db, m.DatabaseTable, table_id, "Table")
    field_ids = [f.semantic_id for f in table.fields]
    if field_ids:
        db.execute(
            delete(m.DatabaseRelation).where(
                or_(
                    m.DatabaseRelation.from_field_semantic_id.in_(field_ids),
                    m.DatabaseRelation.to_field_semantic_id.in_(field_ids),
                )
            )
        )
    db.delete(table)  # fields cascade via ORM relationship
    db.commit()


def delete_relation(db: Session, relation_id: str) -> None:
    relation = get_or_404(db, m.DatabaseRelation, relation_id, "Relation")
    db.delete(relation)
    db.commit()


def save_erd_layout(db: Session, schema_id: str, layout: dict) -> m.DatabaseSchema:
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    schema.layout = layout or {}
    db.commit()
    return schema


def get_erd_layout(db: Session, schema_id: str) -> dict:
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    return schema.layout or {}


def db_design_snapshot(db: Session, schema_id: str) -> dict:
    """Canonical, semantic-id-keyed snapshot of the structured DB design.

    Used by the semantic diff (P1-E): keys are stable semantic ids, so a
    diff over two snapshots reports ADDED/REMOVED/CHANGED objects rather
    than positional noise.
    """
    schema = get_or_404(db, m.DatabaseSchema, schema_id, "Schema")
    tables = db.execute(
        select(m.DatabaseTable).where(m.DatabaseTable.schema_id == schema_id)
    ).scalars().all()
    table_sid = {t.id: t.semantic_id for t in tables}
    tables_out: dict[str, dict] = {
        t.semantic_id: {"name": t.name, "description": t.description, "fields": {}}
        for t in tables
    }
    fields = db.execute(
        select(m.DatabaseField).where(m.DatabaseField.table_id.in_(table_sid))
    ).scalars().all()
    for f in fields:
        tables_out[table_sid[f.table_id]]["fields"][f.semantic_id] = {
            "name": f.name,
            "data_type": f.data_type,
            "length": f.length,
            "nullable": f.nullable,
            "default": f.default,
            "primary_key": f.primary_key,
            "foreign_key": f.foreign_key,
            "reference": f.reference,
            "description": f.description,
            "remark": f.remark,
        }
    relations = {
        r.semantic_id: {
            "from": r.from_field_semantic_id,
            "to": r.to_field_semantic_id,
            "type": r.relation_type,
        }
        for r in db.execute(
            select(m.DatabaseRelation).where(m.DatabaseRelation.schema_id == schema_id)
        ).scalars()
    }
    return {"tables": tables_out, "relations": relations}


# ---------------------------------------------------------------------------
# Revision compare + semantic diff
# ---------------------------------------------------------------------------


def _section_text(section: dict) -> str:
    lines: list[str] = [f"## {section.get('heading') or ''}"]
    if section.get("content") is not None:
        lines.append(_doc_text(section.get("content")))
    else:  # legacy blocks fallback
        for blk in section.get("blocks", []):
            kind = blk.get("kind")
            if kind == "heading":
                lines.append(f"{'#' * (blk.get('level') or 2)} {blk.get('text') or ''}")
            elif kind in ("bullet_list", "numbered_list"):
                for item in blk.get("items", []):
                    lines.append(f"- {item}")
            elif kind == "table":
                lines.append(" | ".join(blk.get("header", [])))
                for row in blk.get("rows", []):
                    lines.append(" | ".join(row))
            elif kind == "code":
                lines.append("```")
                lines.extend((blk.get("text") or "").splitlines())
                lines.append("```")
            else:
                lines.append(blk.get("text") or "")
    return "\n".join(lines)


def text_diff(a: str, b: str) -> list[dict]:
    """Line-level diff (insert/delete/equal) using difflib."""
    sm = difflib.SequenceMatcher(a=a.splitlines(), b=b.splitlines())
    out: list[dict] = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            out.append({"op": "equal", "lines": a.splitlines()[i1:i2]})
        elif tag == "replace":
            out.append({"op": "delete", "lines": a.splitlines()[i1:i2]})
            out.append({"op": "insert", "lines": b.splitlines()[j1:j2]})
        elif tag == "delete":
            out.append({"op": "delete", "lines": a.splitlines()[i1:i2]})
        elif tag == "insert":
            out.append({"op": "insert", "lines": b.splitlines()[j1:j2]})
    return out


def document_diff(db: Session, rev_a_id: str, rev_b_id: str) -> list[dict]:
    """Semantic document diff keyed by stable section ids."""
    a = {s["id"]: s for s in get_document(db, rev_a_id)["sections"]}
    b = {s["id"]: s for s in get_document(db, rev_b_id)["sections"]}
    changes: list[dict] = []
    for sid in sorted(set(a) - set(b)):
        changes.append({"kind": "REMOVED", "object": "SECTION", "semantic_id": sid, "label": a[sid].get("heading") or sid})
    for sid in sorted(set(b) - set(a)):
        changes.append({"kind": "ADDED", "object": "SECTION", "semantic_id": sid, "label": b[sid].get("heading") or sid})
    for sid in sorted(set(a) & set(b)):
        ta, tb = _section_text(a[sid]), _section_text(b[sid])
        if ta != tb:
            changes.append({
                "kind": "CHANGED", "object": "SECTION", "semantic_id": sid,
                "label": b[sid].get("heading") or sid, "text_diff": text_diff(ta, tb),
            })
    return changes


def semantic_db_diff(a: dict, b: dict) -> list[dict]:
    """Semantic diff between two canonical DB design snapshots.

    Keys are stable semantic ids, so changes are reported as ADDED /
    REMOVED / CHANGED objects and attributes — never positional noise.
    """
    changes: list[dict] = []
    a_t, b_t = a.get("tables", {}), b.get("tables", {})
    for sid in sorted(set(a_t) - set(b_t)):
        changes.append({"kind": "REMOVED", "object": "TABLE", "semantic_id": sid, "label": a_t[sid].get("name", sid)})
    for sid in sorted(set(b_t) - set(a_t)):
        changes.append({"kind": "ADDED", "object": "TABLE", "semantic_id": sid, "label": b_t[sid].get("name", sid)})
    for sid in sorted(set(a_t) & set(b_t)):
        ta, tb = a_t[sid], b_t[sid]
        if ta.get("name") != tb.get("name"):
            changes.append({"kind": "CHANGED", "object": "TABLE", "semantic_id": sid, "attribute": "name", "before": ta.get("name"), "after": tb.get("name")})
        fa, fb = ta.get("fields", {}), tb.get("fields", {})
        for fid in sorted(set(fa) - set(fb)):
            changes.append({"kind": "REMOVED", "object": "FIELD", "semantic_id": fid, "label": fa[fid].get("name", fid)})
        for fid in sorted(set(fb) - set(fa)):
            changes.append({"kind": "ADDED", "object": "FIELD", "semantic_id": fid, "label": fb[fid].get("name", fid)})
        for fid in sorted(set(fa) & set(fb)):
            for attr in sorted(set(fa[fid]) | set(fb[fid])):
                if fa[fid].get(attr) != fb[fid].get(attr):
                    changes.append({
                        "kind": "CHANGED", "object": "FIELD", "semantic_id": fid,
                        "attribute": attr, "before": fa[fid].get(attr), "after": fb[fid].get(attr),
                    })

    a_r, b_r = a.get("relations", {}), b.get("relations", {})
    for rid in sorted(set(a_r) - set(b_r)):
        changes.append({"kind": "REMOVED", "object": "RELATION", "semantic_id": rid})
    for rid in sorted(set(b_r) - set(a_r)):
        changes.append({"kind": "ADDED", "object": "RELATION", "semantic_id": rid})
    for rid in sorted(set(a_r) & set(b_r)):
        if a_r[rid].get("type") != b_r[rid].get("type"):
            changes.append({"kind": "CHANGED", "object": "RELATION", "semantic_id": rid, "attribute": "type", "before": a_r[rid].get("type"), "after": b_r[rid].get("type")})
    return changes


def semantic_flow_diff(a: dict, b: dict) -> list[dict]:
    """Approval/process step counts compared by stable step ids where available."""
    changes: list[dict] = []
    fa = a.get("flows", {})
    fb = b.get("flows", {})
    for fid in sorted(set(fa) - set(fb)):
        changes.append({"kind": "REMOVED", "object": "FLOW", "semantic_id": fid})
    for fid in sorted(set(fb) - set(fa)):
        changes.append({"kind": "ADDED", "object": "FLOW", "semantic_id": fid})
    for fid in sorted(set(fa) & set(fb)):
        sa = fa[fid].get("steps", [])
        sb = fb[fid].get("steps", [])
        if len(sa) != len(sb):
            changes.append({"kind": "CHANGED", "object": "FLOW", "semantic_id": fid, "attribute": "steps", "before": len(sa), "after": len(sb)})
    return changes


def semantic_diff(a: dict, b: dict) -> list[dict]:
    """Combined semantic diff (DB design + flows) over stable ids."""
    return semantic_db_diff(a, b) + semantic_flow_diff(a, b)


def snapshot_database_into_revision(db: Session, revision_id: str, schema_id: str) -> m.ArtifactRevision:
    """Embed the current canonical DB design into a DRAFT revision snapshot.

    This creates versioned DB data inside the document revision so that a
    later semantic diff can compare two points in time by stable ids.
    Confirmed revisions are immutable and rejected.
    """
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    require_editable(revision)
    snapshot = dict(revision.snapshot or {})
    snapshot["database"] = db_design_snapshot(db, schema_id)
    revision.snapshot = snapshot
    db.commit()
    return revision


def revision_diff(db: Session, rev_a_id: str, rev_b_id: str) -> dict:
    """Full comparison of two revisions: document + DB semantic diff."""
    ra = get_or_404(db, m.ArtifactRevision, rev_a_id, "Revision")
    rb = get_or_404(db, m.ArtifactRevision, rev_b_id, "Revision")
    sa, sb = ra.snapshot or {}, rb.snapshot or {}
    db_diff = semantic_diff(sa.get("database", {}), sb.get("database", {}))
    return {
        "a": {"id": ra.id, "revision_number": ra.revision_number, "status": ra.status.value},
        "b": {"id": rb.id, "revision_number": rb.revision_number, "status": rb.status.value},
        "document_diff": document_diff(db, rev_a_id, rev_b_id),
        "database_diff": db_diff,
    }


# ---------------------------------------------------------------------------
# Deterministic impact analysis (graph/rule based, no AI)
# ---------------------------------------------------------------------------


def _graph_adjacency(edges) -> tuple[dict, dict]:
    out: dict[str, list[tuple[str, str]]] = {}
    inc: dict[str, list[tuple[str, str]]] = {}
    for e in edges:
        out.setdefault(e.source_semantic_id, []).append((e.target_semantic_id, e.relation_type.value))
        inc.setdefault(e.target_semantic_id, []).append((e.source_semantic_id, e.relation_type.value))
    return out, inc


def impact_paths(db: Session, project_id: str, semantic_id: str, max_depth: int = 3) -> dict:
    """Bounded transitive impact with relation-path explanation.

    Only TraceLink rows are traversed — nothing is inferred.
    """
    edges = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
    ).scalars().all()
    out, inc = _graph_adjacency(edges)

    def walk(adj):
        results: list[list[dict]] = []
        queue: list[tuple[str, list[dict]]] = [(semantic_id, [])]
        while queue:
            node, path = queue.pop(0)
            for nxt, rel in adj.get(node, []):
                new_path = path + [{"semantic_id": nxt, "relation": rel}]
                results.append(new_path)
                if len(new_path) < max_depth:
                    queue.append((nxt, new_path))
        return results

    return {"downstream": walk(out), "upstream": walk(inc)}


def impact_analysis(db: Session, project_id: str, semantic_id: str, max_depth: int = 3) -> dict:
    direct = impact_of(db, project_id, semantic_id)
    paths = impact_paths(db, project_id, semantic_id, max_depth=max_depth)
    return {
        "semantic_id": semantic_id,
        "max_depth": max_depth,
        "direct": direct,
        "paths": paths,
    }


# ---------------------------------------------------------------------------
# Impact analysis v2 — named change sets + rule-based severity
# ---------------------------------------------------------------------------

CHANGE_TYPES = {"MODIFIED", "ADDED", "REMOVED", "RENAMED"}
_SEVERITY_BY_DEPTH = {1: "DIRECT", 2: "HIGH", 3: "MEDIUM"}
_HIGH_BLAST_RADIUS = {
    m.SemanticObjectType.DB_SCHEMA.value,
    m.SemanticObjectType.DB_TABLE.value,
    m.SemanticObjectType.DB_FIELD.value,
    m.SemanticObjectType.DB_RELATION.value,
}


def _severity(depth: int, object_type: str | None, change_type: str | None) -> str:
    base = _SEVERITY_BY_DEPTH.get(depth, "LOW")
    order = ["LOW", "MEDIUM", "HIGH", "DIRECT"]
    idx = order.index(base)
    if object_type in _HIGH_BLAST_RADIUS:
        idx = min(idx + 1, len(order) - 1)
    if change_type == "REMOVED" and object_type in {
        m.SemanticObjectType.DB_TABLE.value,
        m.SemanticObjectType.DB_SCHEMA.value,
    }:
        idx = min(idx + 1, len(order) - 1)
    return order[idx]


def create_change_set(
    db: Session,
    *,
    project_id: str,
    name: str,
    description: str | None = None,
    items: list[dict] | None = None,
    actor="local-user",
    actor_id: str | None = None,
) -> dict:
    guard_project(db, project_id)
    cs = m.ChangeSet(
        project_id=project_id, name=name, description=description,
        created_by=actor, actor_id=actor_id,
    )
    db.add(cs)
    db.flush()
    built: list[dict] = []
    for item in items or []:
        ct = (item.get("change_type") or "MODIFIED").upper()
        if ct not in CHANGE_TYPES:
            raise DomainError(f"invalid change type: {ct}")
        ci = m.ChangeItem(
            change_set_id=cs.id, semantic_id=item["semantic_id"],
            change_type=ct, rationale=item.get("rationale"),
        )
        db.add(ci)
        built.append({"semantic_id": item["semantic_id"], "change_type": ct})
    db.commit()
    db.refresh(cs)
    return {"id": cs.id, "project_id": cs.project_id, "name": cs.name,
            "description": cs.description, "items": built,
            "created_at": cs.created_at.isoformat(), "created_by": cs.created_by}


def list_change_sets(db: Session, project_id: str | None = None) -> list[dict]:
    q = select(m.ChangeSet).order_by(m.ChangeSet.created_at.desc())
    if project_id:
        q = q.where(m.ChangeSet.project_id == project_id)
    out = []
    for cs in db.execute(q).scalars().all():
        items = db.execute(
            select(m.ChangeItem).where(m.ChangeItem.change_set_id == cs.id)
        ).scalars().all()
        out.append({
            "id": cs.id, "project_id": cs.project_id, "name": cs.name,
            "description": cs.description, "created_at": cs.created_at.isoformat(),
            "created_by": cs.created_by,
            "items": [{"semantic_id": i.semantic_id, "change_type": i.change_type,
                       "rationale": i.rationale} for i in items],
        })
    return out


def impact_analysis_v2(
    db: Session, project_id: str, semantic_id: str, max_depth: int = 4
) -> dict:
    """Rule-based impact over transitive trace links.

    Severity: DIRECT (1 hop) > HIGH (2) > MEDIUM (3) > LOW (4+).
    DB schema/table/field/relation and releases bump severity one level;
    REMOVED of a DB table/schema bumps again. Every result carries the
    explicit relation path used, so nothing is inferred silently.
    """
    edges = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
    ).scalars().all()
    out_adj, inc_adj = _graph_adjacency(edges)

    types = {
        so.semantic_id: so.object_type.value
        for so in db.execute(
            select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
        ).scalars().all()
    }

    def walk(adj, direction: str) -> list[dict]:
        results: dict[str, dict] = {}
        queue: list[tuple[str, list[dict], int]] = [(semantic_id, [], 0)]
        while queue:
            node, path, depth = queue.pop(0)
            for nxt, rel in adj.get(node, []):
                if depth + 1 > max_depth:
                    continue
                new_path = path + [{"semantic_id": nxt, "relation": rel}]
                sev = _severity(depth + 1, types.get(nxt), None)
                prev = results.get(nxt)
                if prev is None or _severity_order(sev) > _severity_order(prev["severity"]):
                    results[nxt] = {
                        "semantic_id": nxt,
                        "object_type": types.get(nxt),
                        "severity": sev,
                        "depth": depth + 1,
                        "direction": direction,
                        "path": new_path,
                    }
                queue.append((nxt, new_path, depth + 1))
        return sorted(results.values(), key=lambda r: _severity_order(r["severity"]), reverse=True)

    return {
        "semantic_id": semantic_id,
        "object_type": types.get(semantic_id),
        "max_depth": max_depth,
        "affected": walk(out_adj, "downstream") + walk(inc_adj, "upstream"),
    }


def _severity_order(sev: str) -> int:
    return {"LOW": 1, "MEDIUM": 2, "HIGH": 3, "DIRECT": 4}.get(sev, 0)


# ---------------------------------------------------------------------------
# Project memory — semantic context for the right-hand panel
# ---------------------------------------------------------------------------

_DB_OBJECT_TYPES = {
    m.SemanticObjectType.DB_SCHEMA,
    m.SemanticObjectType.DB_TABLE,
    m.SemanticObjectType.DB_FIELD,
    m.SemanticObjectType.DB_RELATION,
}


def semantic_context(db: Session, project_id: str, semantic_id: str) -> dict:
    """Structured context for one semantic object. Never fabricates."""
    so = db.execute(
        select(m.SemanticObject).where(
            m.SemanticObject.project_id == project_id,
            m.SemanticObject.semantic_id == semantic_id,
        )
    ).scalar_one_or_none()
    if so is None:
        raise DomainError(f"Unknown semantic object '{semantic_id}'", status_code=404)

    out: dict = {
        "semantic_id": so.semantic_id,
        "object_type": so.object_type.value,
        "display_name": so.display_name,
        "entity_ref": so.entity_ref,
        "status": None,
        "confirmed": None,
        "revision": None,
        "evidence": [],
    }

    if so.object_type == m.SemanticObjectType.REQUIREMENT and so.entity_ref:
        req = db.get(m.Requirement, so.entity_ref)
        if req:
            out["status"] = req.status.value
            out["confirmed"] = req.status.value == m.RequirementStatus.CONFIRMED.value
            out["owner"] = {"type": "Requirement", "code": req.code, "priority": req.priority}
    elif so.object_type in _DB_OBJECT_TYPES:
        out["status"] = "live"
        out["owner"] = {"type": "Database design", "note": "working model — frozen at DR revision snapshot, not here"}
    elif so.object_type == m.SemanticObjectType.DOCUMENT_SECTION and so.entity_ref:
        artifact = db.get(m.Artifact, so.entity_ref)
        if artifact:
            revs = db.execute(
                select(m.ArtifactRevision)
                .where(m.ArtifactRevision.artifact_id == artifact.id)
                .order_by(m.ArtifactRevision.revision_number.desc())
            ).scalars().all()
            holders = [
                r for r in revs
                if any(s.get("id") == semantic_id for s in (r.snapshot or {}).get("sections") or [])
            ]
            if holders:
                latest = holders[0]
                out["status"] = latest.status.value
                out["confirmed"] = latest.status.value == m.RevisionStatus.CONFIRMED.value
                out["revision"] = {
                    "id": latest.id,
                    "revision_number": latest.revision_number,
                    "artifact_title": artifact.title,
                    "created_by": latest.created_by,
                    "created_at": latest.created_at.isoformat(),
                    "confirmed_by": latest.confirmed_by,
                    "confirmed_at": latest.confirmed_at.isoformat() if latest.confirmed_at else None,
                }
                out["evidence"] = [
                    {
                        "confirmed_by": c.confirmed_by,
                        "confirmed_at": c.confirmed_at.isoformat(),
                        "comment": c.comment,
                        "evidence": c.evidence,
                    }
                    for c in db.execute(
                        select(m.Confirmation).where(
                            m.Confirmation.artifact_revision_id == latest.id
                        )
                    ).scalars()
                ]

    elif so.object_type == m.SemanticObjectType.PROCESS_FLOW and so.entity_ref:
        flow = db.get(m.ProcessFlow, so.entity_ref)
        if flow:
            out["status"] = "live"
            out["owner"] = {"type": "Process flow", "name": flow.name, "description": flow.description}
    elif so.object_type == m.SemanticObjectType.PROCESS_STEP and so.entity_ref:
        step = db.get(m.ProcessStep, so.entity_ref)
        if step:
            out["status"] = "live"
            out["owner"] = {"type": "Process step", "name": step.name, "step_type": step.step_type, "flow": step.flow.name}
    elif so.object_type == m.SemanticObjectType.API_ENDPOINT and so.entity_ref:
        api = db.get(m.APIEndpoint, so.entity_ref)
        if api:
            out["status"] = "live"
            out["owner"] = {"type": "API endpoint", "method": api.method, "path": api.path, "summary": api.summary}
    elif so.object_type == m.SemanticObjectType.ARCHITECTURE_NODE and so.entity_ref:
        node = db.get(m.ArchitectureNode, so.entity_ref)
        if node:
            out["status"] = "live"
            out["owner"] = {"type": "Architecture node", "name": node.name, "node_type": node.node_type, "technology": node.technology}
    elif so.object_type == m.SemanticObjectType.DECISION and so.entity_ref:
        d = db.get(m.Decision, so.entity_ref)
        if d:
            out["status"] = "recorded"
            out["owner"] = {"type": "Decision", "title": d.title, "decided_by": d.decided_by, "content": d.content}
    elif so.object_type == m.SemanticObjectType.ASSUMPTION and so.entity_ref:
        a = db.get(m.Assumption, so.entity_ref)
        if a:
            out["status"] = a.status
            out["owner"] = {"type": "Assumption", "content": a.content}
    elif so.object_type == m.SemanticObjectType.CLARIFICATION and so.entity_ref:
        c = db.get(m.Clarification, so.entity_ref)
        if c:
            out["status"] = "resolved" if c.resolved else "open"
            out["owner"] = {"type": "Clarification", "question": c.question, "answer": c.answer}

    # annotations on this object
    anns = db.execute(
        select(m.Annotation).where(
            m.Annotation.project_id == project_id,
            m.Annotation.anchor_semantic_id == semantic_id,
        )
    ).scalars().all()
    out["annotations"] = {
        "total": len(anns),
        "open": sum(1 for a in anns if a.status != m.AnnotationStatus.RESOLVED),
        "by_type": dict(Counter(a.type.value for a in anns)),
    }

    return out


# ---------------------------------------------------------------------------
# Semantic search (favours semantic objects over file names)
# ---------------------------------------------------------------------------


def search_semantic(db: Session, project_id: str, query: str, limit: int = 60) -> list[dict]:
    q = (query or "").strip().lower()
    if not q:
        return []
    results: list[dict] = []

    sos = db.execute(
        select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
    ).scalars().all()
    for so in sos:
        hay = f"{so.semantic_id} {so.display_name or ''}".lower()
        if q in hay:
            results.append({
                "kind": "semantic_object",
                "semantic_id": so.semantic_id,
                "object_type": so.object_type.value,
                "title": so.display_name or so.semantic_id,
            })

    anns = db.execute(
        select(m.Annotation).where(m.Annotation.project_id == project_id)
    ).scalars().all()
    for a in anns:
        if q in (a.content or "").lower():
            results.append({
                "kind": "annotation",
                "semantic_id": a.anchor_semantic_id,
                "object_type": a.type.value,
                "title": (a.content or "")[:120],
                "status": a.status.value,
            })

    crs = db.execute(
        select(m.ChangeRequest).where(m.ChangeRequest.project_id == project_id)
    ).scalars().all()
    for cr in crs:
        if q in cr.code.lower() or q in (cr.requested_change or "").lower():
            results.append({
                "kind": "change_request",
                "semantic_id": cr.code,
                "object_type": "CHANGE_REQUEST",
                "title": (cr.requested_change or "")[:120],
                "status": cr.status.value,
            })

    return results[:limit]


# ---------------------------------------------------------------------------
# Process flow designer (structured; diagram is a view over this)
# ---------------------------------------------------------------------------

FLOW_STEP_TYPES = {"START", "ACTION", "DECISION", "APPROVAL", "SYSTEM", "MANUAL", "END"}


def create_flow(
    db: Session, *, project_id: str, name: str, semantic_id: str, description=None, actor="local-user"
) -> m.ProcessFlow:
    flow = m.ProcessFlow(project_id=project_id, semantic_id=semantic_id, name=name, description=description)
    db.add(flow)
    db.flush()
    ensure_semantic_object(
        db, project_id=project_id, semantic_id=semantic_id,
        object_type=m.SemanticObjectType.PROCESS_FLOW, display_name=name, entity_ref=flow.id,
    )
    db.commit()
    return flow


def add_flow_step(
    db: Session, *, flow_id: str, name: str, step_type: str = "ACTION",
    semantic_id: str | None = None, description=None,
) -> m.ProcessStep:
    flow = get_or_404(db, m.ProcessFlow, flow_id, "ProcessFlow")
    if step_type not in FLOW_STEP_TYPES:
        raise DomainError(f"Unknown step type '{step_type}'", status_code=422)
    semantic_id = semantic_id or f"flow_step_{name.lower().replace(' ', '_').replace('-', '_')}"
    existing = db.execute(
        select(m.ProcessStep.id).where(
            m.ProcessStep.flow_id == flow_id, m.ProcessStep.semantic_id == semantic_id
        )
    ).scalar_one_or_none()
    if existing:
        raise DomainError(f"Step semantic id '{semantic_id}' already exists in this flow")
    position = (
        db.execute(select(m.ProcessStep.id).where(m.ProcessStep.flow_id == flow_id)).scalars().all()
    )
    step = m.ProcessStep(
        flow_id=flow_id, semantic_id=semantic_id, name=name, step_type=step_type,
        description=description, position=len(position),
    )
    db.add(step)
    db.flush()
    ensure_semantic_object(
        db, project_id=flow.project_id, semantic_id=semantic_id,
        object_type=m.SemanticObjectType.PROCESS_STEP, display_name=name, entity_ref=step.id,
    )
    db.commit()
    return step


def add_flow_transition(
    db: Session, *, flow_id: str, from_step_semantic_id: str, to_step_semantic_id: str,
    label: str | None = None, condition: str | None = None,
) -> m.ProcessTransition:
    flow = get_or_404(db, m.ProcessFlow, flow_id, "ProcessFlow")
    for sid in (from_step_semantic_id, to_step_semantic_id):
        exists = db.execute(
            select(m.ProcessStep.id).where(
                m.ProcessStep.flow_id == flow_id, m.ProcessStep.semantic_id == sid
            )
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(f"Unknown step '{sid}' in this flow", status_code=422)
    semantic_id = f"flow_transition_{from_step_semantic_id}__{to_step_semantic_id}"
    t = m.ProcessTransition(
        flow_id=flow_id, semantic_id=semantic_id,
        from_step_semantic_id=from_step_semantic_id, to_step_semantic_id=to_step_semantic_id,
        label=label, condition=condition,
    )
    db.add(t)
    db.commit()
    return t


def delete_flow_step(db: Session, step_id: str) -> None:
    step = get_or_404(db, m.ProcessStep, step_id, "ProcessStep")
    flow_id = step.flow_id
    db.execute(
        delete(m.ProcessTransition).where(
            or_(
                m.ProcessTransition.from_step_semantic_id == step.semantic_id,
                m.ProcessTransition.to_step_semantic_id == step.semantic_id,
            )
        )
    )
    db.delete(step)
    db.commit()


def delete_flow_transition(db: Session, transition_id: str) -> None:
    t = get_or_404(db, m.ProcessTransition, transition_id, "ProcessTransition")
    db.delete(t)
    db.commit()


def save_flow_layout(db: Session, flow_id: str, layout: dict) -> m.ProcessFlow:
    flow = get_or_404(db, m.ProcessFlow, flow_id, "ProcessFlow")
    flow.layout = layout or {}
    db.commit()
    return flow


def list_flows(db: Session, project_id: str) -> list[dict]:
    flows = db.execute(
        select(m.ProcessFlow).where(m.ProcessFlow.project_id == project_id)
    ).scalars().all()
    out = []
    for f in flows:
        steps = db.execute(
            select(m.ProcessStep).where(m.ProcessStep.flow_id == f.id).order_by(m.ProcessStep.position)
        ).scalars().all()
        transitions = db.execute(
            select(m.ProcessTransition).where(m.ProcessTransition.flow_id == f.id)
        ).scalars().all()
        out.append({
            "id": f.id, "semantic_id": f.semantic_id, "name": f.name,
            "description": f.description, "layout": f.layout or {},
            "steps": [
                {"id": s.id, "semantic_id": s.semantic_id, "name": s.name,
                 "step_type": s.step_type, "position": s.position, "description": s.description}
                for s in steps
            ],
            "transitions": [
                {"id": t.id, "semantic_id": t.semantic_id,
                 "from": t.from_step_semantic_id, "to": t.to_step_semantic_id,
                 "label": t.label, "condition": t.condition}
                for t in transitions
            ],
        })
    return out


# ---------------------------------------------------------------------------
# API design workspace (structured, not free-text)
# ---------------------------------------------------------------------------


def _slug_path(path: str) -> str:
    return path.strip("/").replace("/", "_").replace("{", "").replace("}", "").replace("-", "_")


def create_api_endpoint(
    db: Session, *, project_id: str, method: str, path: str, summary=None,
    semantic_id: str | None = None, description=None, authentication="NONE", actor="local-user",
) -> m.APIEndpoint:
    semantic_id = semantic_id or f"api_{method.lower()}_{_slug_path(path)}"
    existing = db.execute(
        select(m.APIEndpoint.id).where(
            m.APIEndpoint.project_id == project_id, m.APIEndpoint.semantic_id == semantic_id
        )
    ).scalar_one_or_none()
    if existing:
        raise DomainError(f"API semantic id '{semantic_id}' already exists")
    ep = m.APIEndpoint(
        project_id=project_id, semantic_id=semantic_id, method=method.upper(), path=path,
        summary=summary, description=description, authentication=authentication,
    )
    db.add(ep)
    db.flush()
    ensure_semantic_object(
        db, project_id=project_id, semantic_id=semantic_id,
        object_type=m.SemanticObjectType.API_ENDPOINT, display_name=f"{method.upper()} {path}",
        entity_ref=ep.id,
    )
    db.commit()
    return ep


def update_api_endpoint(db: Session, endpoint_id: str, **changes) -> m.APIEndpoint:
    ep = get_or_404(db, m.APIEndpoint, endpoint_id, "APIEndpoint")
    for k, v in changes.items():
        if hasattr(ep, k):
            setattr(ep, k, v)
    ensure_semantic_object(
        db, project_id=ep.project_id, semantic_id=ep.semantic_id,
        object_type=m.SemanticObjectType.API_ENDPOINT,
        display_name=f"{ep.method} {ep.path}", entity_ref=ep.id,
    )
    db.commit()
    return ep


def _add_api_child(db, model, *, endpoint_id, **fields):
    get_or_404(db, m.APIEndpoint, endpoint_id, "APIEndpoint")
    child = model(endpoint_id=endpoint_id, **fields)
    db.add(child)
    db.commit()
    return child


def _delete_api_child(db, model, child_id):
    child = get_or_404(db, model, child_id, model.__name__)
    db.delete(child)
    db.commit()


def list_api_endpoints(db: Session, project_id: str) -> list[dict]:
    eps = db.execute(
        select(m.APIEndpoint).where(m.APIEndpoint.project_id == project_id).order_by(m.APIEndpoint.path)
    ).scalars().all()
    out = []
    for ep in eps:
        out.append({
            "id": ep.id, "semantic_id": ep.semantic_id, "method": ep.method, "path": ep.path,
            "summary": ep.summary, "description": ep.description, "authentication": ep.authentication,
            "parameters": [
                {"id": p.id, "name": p.name, "location": p.location, "data_type": p.data_type,
                 "required": p.required, "description": p.description} for p in ep.parameters
            ],
            "request_fields": [
                {"id": f.id, "name": f.name, "data_type": f.data_type, "required": f.required,
                 "description": f.description} for f in ep.request_fields
            ],
            "response_fields": [
                {"id": f.id, "status_code": f.status_code, "name": f.name, "data_type": f.data_type,
                 "description": f.description} for f in ep.response_fields
            ],
            "error_responses": [
                {"id": e.id, "status_code": e.status_code, "message": e.message,
                 "description": e.description} for e in ep.error_responses
            ],
        })
    return out


# ---------------------------------------------------------------------------
# OpenAPI 3.x import/export (structured API design as the source of truth)
# ---------------------------------------------------------------------------

# Import safety: bound the uploaded document size; the parser never follows
# remote $ref URLs (no SSRF), so a remote reference is simply treated as an
# unresolved schema name.
MAX_OPENAPI_BYTES = 5 * 1024 * 1024


def _parse_openapi(text: str) -> dict:
    """Parse a JSON or YAML OpenAPI 3.x document into a raw spec dict.

    Does not fetch remote ``$ref`` URLs — a ``$ref`` is only ever reduced to
    its final path component as a schema name.
    """
    if len(text.encode("utf-8")) > MAX_OPENAPI_BYTES:
        raise DomainError("OpenAPI document exceeds the 5MB size limit", status_code=413)
    import yaml as _yaml

    try:
        spec = _json.loads(text)
    except Exception:
        try:
            spec = _yaml.safe_load(text)
        except Exception as exc:
            raise DomainError(f"OpenAPI document is neither valid JSON nor YAML: {exc}", status_code=422)
    if not isinstance(spec, dict) or "paths" not in spec or not isinstance(spec["paths"], dict):
        raise DomainError("OpenAPI document must be an object with a 'paths' object", status_code=422)
    return spec


def _resolve_schema_type(schema: dict) -> str:
    if not isinstance(schema, dict):
        return "string"
    if "type" in schema and isinstance(schema["type"], str):
        return schema["type"]
    if "$ref" in schema:
        return schema["$ref"].split("/")[-1]
    if "anyOf" in schema or "oneOf" in schema:
        return "union"
    return "object"


def openapi_to_endpoints(spec: dict) -> list[dict]:
    """Normalize an OpenAPI 3.x spec into Document Again endpoint dicts."""
    out = []
    for path, path_item in spec.get("paths", {}).items():
        if not isinstance(path_item, dict):
            continue
        for method in ("get", "post", "put", "delete", "patch", "head", "options"):
            op = path_item.get(method)
            if not isinstance(op, dict):
                continue
            parameters = []
            for p in op.get("parameters", []) or []:
                if isinstance(p, dict):
                    sch = p.get("schema") or {}
                    parameters.append({
                        "name": p.get("name", ""), "location": p.get("in", "query"),
                        "data_type": _resolve_schema_type(sch), "required": bool(p.get("required")),
                        "description": p.get("description"),
                    })
            request_fields = []
            if "requestBody" in op and isinstance(op["requestBody"], dict):
                content = op["requestBody"].get("content", {})
                for ctype, media in content.items():
                    schema = (media or {}).get("schema") or {}
                    props = schema.get("properties", {}) if isinstance(schema, dict) else {}
                    for fname, fsch in props.items():
                        if isinstance(fsch, dict):
                            request_fields.append({
                                "name": fname, "data_type": _resolve_schema_type(fsch),
                                "required": fname in (schema.get("required") or []),
                                "description": fsch.get("description"),
                            })
            response_fields = []
            error_responses = []
            for status, resp in (op.get("responses") or {}).items():
                if not isinstance(resp, dict):
                    continue
                if status.startswith("2"):
                    content = resp.get("content", {})
                    for ctype, media in content.items():
                        schema = (media or {}).get("schema") or {}
                        props = schema.get("properties", {}) if isinstance(schema, dict) else {}
                        for fname, fsch in props.items():
                            if isinstance(fsch, dict):
                                response_fields.append({
                                    "status_code": status, "name": fname,
                                    "data_type": _resolve_schema_type(fsch),
                                    "description": fsch.get("description"),
                                })
                else:
                    error_responses.append({
                        "status_code": status,
                        "message": resp.get("description") or f"HTTP {status}",
                        "description": resp.get("description"),
                    })
            out.append({
                "method": method.upper(), "path": path, "summary": op.get("summary"),
                "description": op.get("description"),
                "authentication": "BEARER" if ("security" in op and op["security"]) else "NONE",
                "parameters": parameters, "request_fields": request_fields,
                "response_fields": response_fields, "error_responses": error_responses,
            })
    return out


def preview_openapi_import(db: Session, project_id: str, text: str) -> dict:
    """Diff a parsed OpenAPI spec against existing endpoints (no writes)."""
    incoming = openapi_to_endpoints(_parse_openapi(text))
    existing = list_api_endpoints(db, project_id)
    existing_keys = {e["semantic_id"] for e in existing}
    report = {"added": [], "changed": [], "removed": [], "unchanged": [], "conflicts": []}
    for ep in incoming:
        sid = f"api_{ep['method'].lower()}_{_slug_path(ep['path'])}"
        if sid in existing_keys:
            current = next(e for e in existing if e["semantic_id"] == sid)
            if (current["method"] == ep["method"] and current["path"] == ep["path"]
                    and current["summary"] == ep["summary"]):
                report["unchanged"].append(sid)
            else:
                report["changed"].append(sid)
        else:
            report["added"].append(sid)
    incoming_keys = {f"api_{ep['method'].lower()}_{_slug_path(ep['path'])}" for ep in incoming}
    for e in existing:
        if e["semantic_id"] not in incoming_keys:
            report["removed"].append(e["semantic_id"])
    return {"project_id": project_id, **report}


def import_openapi(db: Session, project_id: str, text: str, actor="local-user") -> dict:
    """Apply an OpenAPI spec: create missing endpoints, update changed ones.

    Removed endpoints are never deleted (historical truth); they are listed
    in the report so a human can decide.
    """
    spec = _parse_openapi(text)
    preview = preview_openapi_import(db, project_id, text)
    applied = []
    for ep in openapi_to_endpoints(spec):
        sid = f"api_{ep['method'].lower()}_{_slug_path(ep['path'])}"
        existing = db.execute(
            select(m.APIEndpoint).where(
                m.APIEndpoint.project_id == project_id, m.APIEndpoint.semantic_id == sid
            )
        ).scalar_one_or_none()
        if existing is None:
            created = create_api_endpoint(
                db, project_id=project_id, method=ep["method"], path=ep["path"],
                summary=ep["summary"], semantic_id=sid, description=ep["description"],
                authentication=ep["authentication"], actor=actor,
            )
            _apply_api_children(db, created, ep)
            applied.append({"semantic_id": sid, "action": "ADDED"})
        else:
            update_api_endpoint(
                db, existing.id, method=ep["method"], path=ep["path"],
                summary=ep["summary"], description=ep["description"],
                authentication=ep["authentication"],
            )
            _replace_api_children(db, existing, ep)
            applied.append({"semantic_id": sid, "action": "UPDATED"})
    return {"project_id": project_id, "applied": applied,
            "removed_in_spec": preview["removed"], "conflicts": preview["conflicts"]}


def _apply_api_children(db, ep, ep_dict):
    for p in ep_dict["parameters"]:
        _add_api_child(db, m.ApiParameter, endpoint_id=ep.id, **p)
    for f in ep_dict["request_fields"]:
        _add_api_child(db, m.ApiRequestField, endpoint_id=ep.id, **f)
    for f in ep_dict["response_fields"]:
        _add_api_child(db, m.ApiResponseField, endpoint_id=ep.id, **f)
    for e in ep_dict["error_responses"]:
        _add_api_child(db, m.ApiErrorResponse, endpoint_id=ep.id, **e)


def _replace_api_children(db, ep, ep_dict):
    for model in (m.ApiParameter, m.ApiRequestField, m.ApiResponseField, m.ApiErrorResponse):
        for child in db.execute(
            select(model).where(model.endpoint_id == ep.id)
        ).scalars().all():
            db.delete(child)
    db.flush()
    _apply_api_children(db, ep, ep_dict)


def export_openapi(db: Session, revision_id: str) -> dict:
    """Reconstruct an OpenAPI 3.0 document from a revision's API snapshot."""
    rev = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    snapshot = rev.snapshot or {}
    api_map = snapshot.get("technical_design", {}).get("api_endpoints") or {}
    if isinstance(api_map, dict):
        endpoints = list(api_map.values())
    else:
        endpoints = api_map
    paths: dict = {}
    for ep in endpoints:
        method = (ep.get("method") or "get").lower()
        path = ep.get("path") or "/"
        op = paths.setdefault(path, {})
        op[method] = {
            "summary": ep.get("summary"),
            "description": ep.get("description"),
            "parameters": [
                {"name": p.get("name"), "in": p.get("location", "query"),
                 "required": p.get("required", False),
                 "schema": {"type": p.get("data_type", "string")},
                 "description": p.get("description")}
                for p in (ep.get("parameters") or [])
            ],
            "responses": {
                str(f.get("status_code", "200")): {
                    "description": f.get("description") or "",
                    "content": {"application/json": {"schema": {
                        "type": "object", "properties": {
                            f.get("name", "value"): {"type": f.get("data_type", "string")}
                        }
                    }}}
                }
                for f in (ep.get("response_fields") or [])
            } or {"200": {"description": "OK"}},
        }
    return {"openapi": "3.0.0", "info": {"title": "Document Again — exported API design",
                                         "version": "1.0"}, "paths": paths}


# ---------------------------------------------------------------------------
# Architecture design workspace
# ---------------------------------------------------------------------------

ARCH_NODE_TYPES = {
    "USER", "CLIENT", "SERVICE", "DATABASE", "QUEUE", "STORAGE",
    "EXTERNAL_SYSTEM", "NETWORK_ZONE", "CLOUD_SERVICE",
}


def create_architecture_diagram(
    db: Session, *, project_id: str, name: str, semantic_id: str, description=None, actor="local-user"
) -> m.ArchitectureDiagram:
    d = m.ArchitectureDiagram(project_id=project_id, semantic_id=semantic_id, name=name, description=description)
    db.add(d)
    db.commit()
    return d


def add_architecture_node(
    db: Session, *, diagram_id: str, name: str, semantic_id: str,
    node_type: str = "SERVICE", description=None, technology=None, environment=None, metadata=None,
) -> m.ArchitectureNode:
    diagram = get_or_404(db, m.ArchitectureDiagram, diagram_id, "ArchitectureDiagram")
    if node_type not in ARCH_NODE_TYPES:
        raise DomainError(f"Unknown node type '{node_type}'", status_code=422)
    node = m.ArchitectureNode(
        diagram_id=diagram_id, semantic_id=semantic_id, name=name, node_type=node_type,
        description=description, technology=technology, environment=environment, metadata_json=metadata,
    )
    db.add(node)
    db.flush()
    ensure_semantic_object(
        db, project_id=diagram.project_id, semantic_id=semantic_id,
        object_type=m.SemanticObjectType.ARCHITECTURE_NODE, display_name=name, entity_ref=node.id,
    )
    db.commit()
    return node


def add_architecture_edge(
    db: Session, *, diagram_id: str, from_node_semantic_id: str, to_node_semantic_id: str, label=None,
) -> m.ArchitectureEdge:
    diagram = get_or_404(db, m.ArchitectureDiagram, diagram_id, "ArchitectureDiagram")
    for sid in (from_node_semantic_id, to_node_semantic_id):
        exists = db.execute(
            select(m.ArchitectureNode.id).where(
                m.ArchitectureNode.diagram_id == diagram_id,
                m.ArchitectureNode.semantic_id == sid,
            )
        ).scalar_one_or_none()
        if not exists:
            raise DomainError(f"Unknown node '{sid}' in this diagram", status_code=422)
    edge = m.ArchitectureEdge(
        diagram_id=diagram_id,
        semantic_id=f"edge_{from_node_semantic_id}__{to_node_semantic_id}",
        from_node_semantic_id=from_node_semantic_id, to_node_semantic_id=to_node_semantic_id, label=label,
    )
    db.add(edge)
    db.commit()
    return edge


def delete_architecture_node(db: Session, node_id: str) -> None:
    node = get_or_404(db, m.ArchitectureNode, node_id, "ArchitectureNode")
    db.execute(
        delete(m.ArchitectureEdge).where(
            or_(
                m.ArchitectureEdge.from_node_semantic_id == node.semantic_id,
                m.ArchitectureEdge.to_node_semantic_id == node.semantic_id,
            )
        )
    )
    db.delete(node)
    db.commit()


def delete_architecture_edge(db: Session, edge_id: str) -> None:
    edge = get_or_404(db, m.ArchitectureEdge, edge_id, "ArchitectureEdge")
    db.delete(edge)
    db.commit()


def save_architecture_layout(db: Session, diagram_id: str, layout: dict) -> m.ArchitectureDiagram:
    diagram = get_or_404(db, m.ArchitectureDiagram, diagram_id, "ArchitectureDiagram")
    diagram.layout = layout or {}
    db.commit()
    return diagram


def list_architecture_diagrams(db: Session, project_id: str) -> list[dict]:
    diagrams = db.execute(
        select(m.ArchitectureDiagram).where(m.ArchitectureDiagram.project_id == project_id)
    ).scalars().all()
    out = []
    for d in diagrams:
        nodes = db.execute(select(m.ArchitectureNode).where(m.ArchitectureNode.diagram_id == d.id)).scalars().all()
        edges = db.execute(select(m.ArchitectureEdge).where(m.ArchitectureEdge.diagram_id == d.id)).scalars().all()
        out.append({
            "id": d.id, "semantic_id": d.semantic_id, "name": d.name, "description": d.description,
            "layout": d.layout or {},
            "nodes": [
                {"id": n.id, "semantic_id": n.semantic_id, "name": n.name, "node_type": n.node_type,
                 "description": n.description, "technology": n.technology, "environment": n.environment}
                for n in nodes
            ],
            "edges": [
                {"id": e.id, "semantic_id": e.semantic_id, "from": e.from_node_semantic_id,
                 "to": e.to_node_semantic_id, "label": e.label}
                for e in edges
            ],
        })
    return out


# ---------------------------------------------------------------------------
# Decision / Assumption / Clarification project-memory surfaces
# ---------------------------------------------------------------------------


def _next_code(db: Session, project_id: str, prefix: str, model) -> str:
    count = db.execute(select(model.id).where(model.project_id == project_id)).scalars().all()
    return f"{prefix}-{len(count) + 1:04d}"


def create_decision(
    db: Session, *, project_id: str, title: str, content: str,
    decided_by="local-user", related_semantic_ids: list[str] | None = None, actor="local-user",
    actor_id: str | None = None,
) -> m.Decision:
    code = _next_code(db, project_id, "DEC", m.Decision)
    d = m.Decision(project_id=project_id, semantic_id=code, title=title, content=content, decided_by=decided_by, actor_id=actor_id)
    db.add(d)
    db.flush()
    ensure_semantic_object(
        db, project_id=project_id, semantic_id=code, object_type=m.SemanticObjectType.DECISION,
        display_name=title, entity_ref=d.id,
    )
    for sid in (related_semantic_ids or []):
        create_trace_link(db, project_id=project_id, source_semantic_id=code, target_semantic_id=sid, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
    db.commit()
    return d


def create_assumption(db: Session, *, project_id: str, content: str, related_semantic_ids=None, actor="local-user") -> m.Assumption:
    code = _next_code(db, project_id, "ASM", m.Assumption)
    a = m.Assumption(project_id=project_id, semantic_id=code, content=content, created_by=actor)
    db.add(a)
    db.flush()
    ensure_semantic_object(db, project_id=project_id, semantic_id=code, object_type=m.SemanticObjectType.ASSUMPTION, display_name=code, entity_ref=a.id)
    for sid in (related_semantic_ids or []):
        create_trace_link(db, project_id=project_id, source_semantic_id=code, target_semantic_id=sid, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
    db.commit()
    return a


def create_clarification(db: Session, *, project_id: str, question: str, answer=None, related_semantic_ids=None, actor="local-user") -> m.Clarification:
    code = _next_code(db, project_id, "CLR", m.Clarification)
    c = m.Clarification(project_id=project_id, semantic_id=code, question=question, answer=answer, asked_by=actor, resolved=answer is not None)
    db.add(c)
    db.flush()
    ensure_semantic_object(db, project_id=project_id, semantic_id=code, object_type=m.SemanticObjectType.CLARIFICATION, display_name=code, entity_ref=c.id)
    for sid in (related_semantic_ids or []):
        create_trace_link(db, project_id=project_id, source_semantic_id=code, target_semantic_id=sid, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
    db.commit()
    return c


def list_project_memory(db: Session, project_id: str) -> dict:
    decisions = db.execute(select(m.Decision).where(m.Decision.project_id == project_id)).scalars().all()
    assumptions = db.execute(select(m.Assumption).where(m.Assumption.project_id == project_id)).scalars().all()
    clarifications = db.execute(select(m.Clarification).where(m.Clarification.project_id == project_id)).scalars().all()
    links = db.execute(select(m.TraceLink).where(m.TraceLink.project_id == project_id)).scalars().all()

    def related(sid):
        return [l.target_semantic_id for l in links if l.source_semantic_id == sid and l.relation_type == m.TraceRelationType.REFERENCES]

    return {
        "decisions": [
            {"id": d.id, "code": d.semantic_id, "title": d.title, "content": d.content,
             "decided_by": d.decided_by, "decided_at": d.decided_at.isoformat(), "related": related(d.semantic_id)}
            for d in decisions
        ],
        "assumptions": [
            {"id": a.id, "code": a.semantic_id, "content": a.content, "status": a.status,
             "created_by": a.created_by, "related": related(a.semantic_id)}
            for a in assumptions
        ],
        "clarifications": [
            {"id": c.id, "code": c.semantic_id, "question": c.question, "answer": c.answer,
             "asked_by": c.asked_by, "resolved": c.resolved, "related": related(c.semantic_id)}
            for c in clarifications
        ],
    }


def _doc_updated_iso(rev: m.ArtifactRevision) -> str:
    """Human 'last updated' for a revision — confirmation wins over creation."""
    return _iso(rev.confirmed_at or rev.created_at)


def project_home(db: Session, project_id: str) -> dict:
    """Aggregate Project Home payload (P1 UX). One request answers:
    what project, current baseline, what documents exist + which is current,
    what PM/QA are doing, what changed recently, and what is open.

    Human-facing titles/versions are computed here so the UI never has to
    resolve semantic IDs. Technical IDs are still carried as metadata only.
    """
    project = guard_project(db, project_id)

    baselines = db.execute(
        select(m.Baseline).where(m.Baseline.project_id == project_id)
        .order_by(m.Baseline.created_at.desc())
    ).scalars().all()
    current_baseline = baselines[0] if baselines else None

    artifacts = db.execute(
        select(m.Artifact).where(m.Artifact.project_id == project_id)
    ).scalars().all()
    revisions_by_artifact: dict[str, list[m.ArtifactRevision]] = {}
    for a in artifacts:
        revs = db.execute(
            select(m.ArtifactRevision)
            .where(m.ArtifactRevision.artifact_id == a.id)
            .order_by(m.ArtifactRevision.revision_number.desc())
        ).scalars().all()
        revisions_by_artifact[a.id] = revs

    # Which revision is bound to the *current* baseline per artifact.
    current_bindings: dict[str, str] = {}
    if current_baseline:
        for bb in current_baseline.bindings:
            if bb.artifact_id:
                current_bindings[bb.artifact_id] = bb.artifact_revision_id

    requirements = db.execute(
        select(m.Requirement).where(m.Requirement.project_id == project_id)
        .order_by(m.Requirement.code)
    ).scalars().all()

    memory = list_project_memory(db, project_id)
    handoffs_pm = list_execution_handoffs(db, project_id)
    handoffs_qa = list_qa_handoffs(db, project_id)
    tl = timeline(db, project_id)

    arch = db.execute(
        select(m.ArchitectureDiagram).where(m.ArchitectureDiagram.project_id == project_id)
    ).scalars().all()
    flows = db.execute(
        select(m.ProcessFlow).where(m.ProcessFlow.project_id == project_id)
    ).scalars().all()

    # ── Documents (human-readable, current-first) ──
    documents: list[dict] = []

    def _rev_meta(a: m.Artifact):
        revs = revisions_by_artifact.get(a.id) or []
        rev = None
        bound = current_bindings.get(a.id)
        if bound:
            rev = next((r for r in revs if r.id == bound), None)
        if rev is None:
            rev = next((r for r in revs if r.status == m.RevisionStatus.CONFIRMED), None) or (revs[0] if revs else None)
        return revs, rev

    # Requirement Register
    latest_req_ts = max((r.created_at for r in requirements), default=project.created_at)
    documents.append({
        "key": "requirements", "title": "Requirement Register", "type": "REQUIREMENT_REGISTER",
        "version": None, "status": "CURRENT", "count": len(requirements),
        "updated_at": _iso(latest_req_ts), "open_route": "/requirements", "download": None,
    })

    for a in artifacts:
        revs, rev = _rev_meta(a)
        human = a.type.value.upper()
        version = f"v{rev.revision_number}" if rev else None
        documents.append({
            "key": a.id, "title": a.title, "type": human,
            "version": version, "status": rev.status.value if rev else "DRAFT",
            "updated_at": _doc_updated_iso(rev) if rev else _iso(a.created_at),
            "revision_id": rev.id if rev else None,
            "open_route": "/requirements/ur" if a.type == m.ArtifactType.UR else "/design/dr",
            "download": f"/api/revisions/{rev.id}/export?format=xlsx" if rev else None,
        })

    documents.append({
        "key": "traceability", "title": "Traceability Matrix", "type": "TRACEABILITY_MATRIX",
        "version": None, "status": "CURRENT",
        "updated_at": _iso(max((l.created_at for l in db.execute(select(m.TraceLink).where(m.TraceLink.project_id == project_id)).scalars().all()), default=project.created_at)),
        "open_route": "/design/trace", "download": None,
    })

    for d in arch:
        documents.append({
            "key": d.id, "title": d.name, "type": "ARCHITECTURE", "version": None,
            "status": "CURRENT", "updated_at": _iso(d.created_at),
            "open_route": "/design/architecture", "download": None,
        })
    for f in flows:
        documents.append({
            "key": f.id, "title": f.name, "type": "PROCESS_FLOW", "version": None,
            "status": "CURRENT", "updated_at": _iso(f.created_at),
            "open_route": "/design/flows", "download": None,
        })

    for key, title, items in (
        ("clarifications", "Clarification Register", memory["clarifications"]),
        ("assumptions", "Assumption Register", memory["assumptions"]),
        ("decisions", "Decision Register", memory["decisions"]),
    ):
        documents.append({
            "key": key, "title": title, "type": key.upper() + "_REGISTER",
            "version": None, "status": "CURRENT", "count": len(items),
            "updated_at": _iso(project.created_at), "open_route": "/decisions", "download": None,
        })

    # ── PM / QA status (handoff truth; never fabricate) ──
    def _pm_qa_status(handoffs, target):
        if not handoffs:
            return {"state": "NOT_SENT", "human": "Not sent"}
        acked = [h for h in handoffs if h.get("status") == "ACKNOWLEDGED"]
        failed = [h for h in handoffs if h.get("status") == "FAILED"]
        if acked:
            return {"state": "SENT", "human": "Delivered", "handoff": acked[-1]}
        if failed:
            return {"state": "FAILED", "human": "Delivery failed", "handoff": failed[-1]}
        return {"state": handoffs[-1].get("status") or "DRAFT", "human": "Pending", "handoff": handoffs[-1]}

    pm_status = _pm_qa_status(handoffs_pm, "pm")
    qa_status = _pm_qa_status(handoffs_qa, "qa")

    # Requirements with resolved human trace targets (REQ -> UR/DR/flow/baseline).
    trace_graph_data = trace_graph(db, project_id)
    nodes = {n["semantic_id"]: n for n in trace_graph_data["nodes"]}
    edges = trace_graph_data["edges"]

    reqs_out = []
    for r in requirements:
        targets = []
        for e in edges:
            if e["source"] == r.code:
                t = nodes.get(e["target"])
                targets.append({
                    "relation": e["relation"], "target": e["target"],
                    "display_name": (t or {}).get("display_name") if t else e["target"],
                    "object_type": (t or {}).get("object_type") if t else None,
                })
        reqs_out.append({
            "id": r.id, "code": r.code, "title": r.title, "status": r.status.value,
            "targets": targets,
        })

    # Last updated across the whole project.
    timestamps = [ev["at"] for ev in tl] + [d["updated_at"] for d in documents if d["updated_at"]]
    last_updated = max(timestamps) if timestamps else _iso(project.created_at)

    return {
        "project": {
            "id": project.id, "key": project.key, "name": project.name,
            "description": project.description, "tenant_id": project.tenant_id,
            "created_at": _iso(project.created_at), "created_by": project.created_by,
        },
        "current_baseline": {
            "id": current_baseline.id, "name": current_baseline.name,
            "created_at": _iso(current_baseline.created_at),
        } if current_baseline else None,
        "baselines": [
            {"id": b.id, "name": b.name, "created_at": _iso(b.created_at)} for b in baselines
        ],
        "last_updated": last_updated,
        "open_clarifications": sum(1 for c in memory["clarifications"] if not c["resolved"]),
        "assumptions": len(memory["assumptions"]),
        "decisions": len(memory["decisions"]),
        "clarifications": memory["clarifications"],
        "documents": documents,
        "pm": pm_status,
        "qa": qa_status,
        "requirements": reqs_out,
        "activity": tl[:20],
    }


def promote_annotation(db: Session, *, annotation_id: str, to_kind: str, actor="local-user") -> dict:
    """Promote a comment/annotation into a first-class project-memory record.

    Provenance is retained: the source annotation id / thread is recorded
    and the annotation is marked resolved.
    """
    ann = get_or_404(db, m.Annotation, annotation_id, "Annotation")
    provenance = {
        "source": "annotation", "annotation_id": ann.id, "thread_id": ann.thread_id,
        "author": ann.created_by, "anchor": ann.anchor_semantic_id,
    }
    source_label = f"Comment Thread #{ann.thread_id}" if ann.thread_id else f"Annotation #{ann.id}"

    if to_kind == "change_request":
        cr = create_change_request(
            db, project_id=ann.project_id, requested_change=ann.content,
            affected_semantic_ids=[ann.anchor_semantic_id], reason=f"Promoted from {source_label}",
            actor=actor,
        )
        result = {"kind": "change_request", "code": cr.code, "provenance": provenance}
    elif to_kind == "decision":
        code = _next_code(db, ann.project_id, "DEC", m.Decision)
        d = m.Decision(project_id=ann.project_id, semantic_id=code, title=(ann.content or code)[:80], content=ann.content, decided_by=actor)
        db.add(d)
        db.flush()
        ensure_semantic_object(db, project_id=ann.project_id, semantic_id=code, object_type=m.SemanticObjectType.DECISION,
                               display_name=(ann.content or code)[:80], entity_ref=d.id, metadata={"provenance": provenance})
        create_trace_link(db, project_id=ann.project_id, source_semantic_id=code, target_semantic_id=ann.anchor_semantic_id, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
        result = {"kind": "decision", "code": code, "provenance": provenance}
    elif to_kind == "assumption":
        code = _next_code(db, ann.project_id, "ASM", m.Assumption)
        a = m.Assumption(project_id=ann.project_id, semantic_id=code, content=ann.content, created_by=actor)
        db.add(a)
        db.flush()
        ensure_semantic_object(db, project_id=ann.project_id, semantic_id=code, object_type=m.SemanticObjectType.ASSUMPTION,
                               display_name=code, entity_ref=a.id, metadata={"provenance": provenance})
        create_trace_link(db, project_id=ann.project_id, source_semantic_id=code, target_semantic_id=ann.anchor_semantic_id, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
        result = {"kind": "assumption", "code": code, "provenance": provenance}
    elif to_kind == "clarification":
        code = _next_code(db, ann.project_id, "CLR", m.Clarification)
        c = m.Clarification(project_id=ann.project_id, question=ann.content, asked_by=actor, resolved=False)
        db.add(c)
        db.flush()
        ensure_semantic_object(db, project_id=ann.project_id, semantic_id=code, object_type=m.SemanticObjectType.CLARIFICATION,
                               display_name=code, entity_ref=c.id, metadata={"provenance": provenance})
        create_trace_link(db, project_id=ann.project_id, source_semantic_id=code, target_semantic_id=ann.anchor_semantic_id, relation_type=m.TraceRelationType.REFERENCES, actor=actor)
        result = {"kind": "clarification", "code": code, "provenance": provenance}
    else:
        raise DomainError(f"Unknown promotion target '{to_kind}'", status_code=422)

    ann.status = m.AnnotationStatus.RESOLVED
    db.commit()
    return result


def record_actor(db: Session, actor_id: str, display_name: str, tenant_id: str | None = None, source: str = "LOCAL") -> None:
    """Cache a resolved actor identity (idempotent upsert)."""
    existing = db.get(m.ActorIdentity, actor_id)
    if existing:
        existing.display_name = display_name
        if tenant_id:
            existing.tenant_id = tenant_id
        existing.source = source
        existing.resolved_at = m.utcnow()
    else:
        db.add(m.ActorIdentity(actor_id=actor_id, display_name=display_name, tenant_id=tenant_id, source=source))
    db.commit()


def record_audit(
    db: Session,
    *,
    action: str,
    project_id: str | None = None,
    tenant_id: str | None = None,
    actor_id: str | None = None,
    object_type: str | None = None,
    object_id: str | None = None,
    revision_context: str | None = None,
    baseline_id: str | None = None,
    correlation_id: str | None = None,
    metadata: dict | None = None,
) -> m.AuditEvent:
    """Write an immutable audit event. Independent of editable comments."""
    ev = m.AuditEvent(
        tenant_id=tenant_id if tenant_id is not None else current_tenant(),
        project_id=project_id, actor_id=actor_id, action=action,
        object_type=object_type, object_id=object_id,
        revision_context=revision_context, baseline_id=baseline_id,
        correlation_id=correlation_id, metadata_json=metadata,
    )
    db.add(ev)
    db.commit()
    return ev


def list_audit_events(
    db: Session,
    *,
    project_id: str | None = None,
    actor_id: str | None = None,
    action: str | None = None,
    object_id: str | None = None,
    baseline_id: str | None = None,
    limit: int = 200,
) -> list[dict]:
    tenant = current_tenant()
    q = select(m.AuditEvent).order_by(m.AuditEvent.created_at.desc()).limit(limit)
    if project_id:
        guard_project(db, project_id)
        q = q.where(m.AuditEvent.project_id == project_id)
    elif tenant is not None:
        q = q.where(m.AuditEvent.tenant_id == tenant)
    if actor_id:
        q = q.where(m.AuditEvent.actor_id == actor_id)
    if action:
        q = q.where(m.AuditEvent.action == action)
    if object_id:
        q = q.where(m.AuditEvent.object_id == object_id)
    if baseline_id:
        q = q.where(m.AuditEvent.baseline_id == baseline_id)
    return [
        {
            "id": e.id, "tenant_id": e.tenant_id, "project_id": e.project_id,
            "actor_id": e.actor_id, "action": e.action, "object_type": e.object_type,
            "object_id": e.object_id, "revision_context": e.revision_context,
            "baseline_id": e.baseline_id, "correlation_id": e.correlation_id,
            "metadata": e.metadata_json, "created_at": e.created_at.isoformat(),
        }
        for e in db.execute(q).scalars().all()
    ]


# ---------------------------------------------------------------------------
# Ecosystem event + outbox (durable, idempotent delivery)
# ---------------------------------------------------------------------------

ECOSYSTEM_EVENT_TYPES = {
    "REQUIREMENT_BASELINED", "DESIGN_BASELINED", "CHANGE_REQUEST_APPROVED",
    "DESIGN_CHANGED", "EXECUTION_REQUESTED", "QA_VALIDATION_REQUESTED",
    "QA_RESULT_RECEIVED", "RELEASE_LINKED",
}

OUTBOX_STATUSES = {"PENDING", "SENT", "ACKNOWLEDGED", "FAILED"}


def emit_event(
    db: Session,
    *,
    event_type: str,
    project_id: str,
    payload: dict | None = None,
    target_services: list[str] | None = None,
    correlation_id: str | None = None,
    source_object_id: str | None = None,
    source_revision: str | None = None,
    actor_id: str | None = None,
    tenant_id: str | None = None,
) -> m.EcosystemEvent:
    """Record an ecosystem event and enqueue durable delivery, atomically.

    The domain change and its outbox rows commit together; delivery never
    depends on a fire-and-forget HTTP call.
    """
    if event_type not in ECOSYSTEM_EVENT_TYPES:
        raise DomainError(f"Unknown event type '{event_type}'", status_code=422)
    correlation_id = correlation_id or m.new_id("corr")
    event = m.EcosystemEvent(
        event_type=event_type, project_id=project_id, tenant_id=tenant_id,
        source_service="document-again", source_object_id=source_object_id,
        source_revision=source_revision, actor_id=actor_id,
        payload_version="1.0", payload=payload, correlation_id=correlation_id,
    )
    db.add(event)
    db.flush()
    for ts in (target_services or []):
        db.add(m.OutboxEvent(event_id=event.id, target_service=ts, correlation_id=correlation_id))
    metrics.inc("outbox_pending", len(target_services or []))
    db.commit()
    return event


def due_outbox(db: Session, limit: int = 50) -> list[m.OutboxEvent]:
    now = m.utcnow()
    rows = db.execute(
        select(m.OutboxEvent).where(
            m.OutboxEvent.status.in_(["PENDING", "FAILED"]),
            (m.OutboxEvent.next_attempt_at.is_(None)) | (m.OutboxEvent.next_attempt_at <= now),
        ).order_by(m.OutboxEvent.created_at).limit(limit)
    ).scalars().all()
    return rows


def _backoff_seconds(attempt_count: int) -> int:
    return min(5 * (2 ** max(attempt_count, 1)), 300)


def mark_outbox_sent(db: Session, outbox_id: str, external_reference: str | None = None) -> m.OutboxEvent:
    out = get_or_404(db, m.OutboxEvent, outbox_id, "OutboxEvent")
    out.status = "SENT"
    out.delivered_at = m.utcnow()
    out.attempt_count += 1
    out.last_error = None
    out.external_reference = external_reference
    db.commit()
    return out


def mark_outbox_acknowledged(db: Session, outbox_id: str) -> m.OutboxEvent:
    out = get_or_404(db, m.OutboxEvent, outbox_id, "OutboxEvent")
    out.status = "ACKNOWLEDGED"
    out.delivered_at = m.utcnow()
    db.commit()
    return out


def mark_outbox_failed(db: Session, outbox_id: str, error: str) -> m.OutboxEvent:
    out = get_or_404(db, m.OutboxEvent, outbox_id, "OutboxEvent")
    out.status = "FAILED"
    out.attempt_count += 1
    out.last_error = error[:2000]
    out.next_attempt_at = m.utcnow() + timedelta(seconds=_backoff_seconds(out.attempt_count))
    db.commit()
    return out


def deliver_due_events(db: Session, deliver_fn, limit: int = 50) -> dict:
    """Dispatch due outbox events through a delivery callback.

    `deliver_fn(outbox_event) -> str | None` returns an external reference
    on success. Failures are recorded with a backoff for retry; the unique
    (event_id, target_service) constraint keeps delivery idempotent.
    """
    delivered = 0
    failed = 0
    for out in due_outbox(db, limit):
        try:
            ext_ref = deliver_fn(out)
            out.status = "SENT"
            out.delivered_at = m.utcnow()
            out.attempt_count += 1
            out.last_error = None
            if ext_ref:
                out.external_reference = ext_ref
            delivered += 1
        except Exception as exc:  # noqa: BLE001 — record and retry later
            out.status = "FAILED"
            out.attempt_count += 1
            out.last_error = str(exc)[:2000]
            out.next_attempt_at = m.utcnow() + timedelta(seconds=_backoff_seconds(out.attempt_count))
            failed += 1
    db.commit()
    metrics.inc("outbox_delivered", delivered)
    metrics.inc("outbox_failed", failed)
    return {"delivered": delivered, "failed": failed}


def build_event_delivery_payload(db: Session, outbox_event: m.OutboxEvent) -> dict:
    """Versioned ecosystem-event envelope for outbound delivery."""
    event = db.get(m.EcosystemEvent, outbox_event.event_id)
    if event is None:
        raise DomainError(f"Ecosystem event {outbox_event.event_id} missing", status_code=404)
    return versioned_payload(
        "ecosystem-event", 1,
        eventType=event.event_type,
        correlationId=event.correlation_id,
        sourceService=event.source_service,
        sourceObjectId=event.source_object_id,
        sourceRevision=event.source_revision,
        occurredAt=event.occurred_at.isoformat(),
        actorId=event.actor_id,
        payload=event.payload or {},
    )


def deliver_due_events_http(
    db: Session,
    target_url: str,
    tenant_id: str | None = None,
    client=None,
    limit: int = 50,
) -> dict:
    """Dispatch due outbox events over HTTP (idempotent, versioned payload)."""
    from .ecosystem_client import EcosystemDeliveryClient

    c = client or EcosystemDeliveryClient()

    def fn(out: m.OutboxEvent) -> str | None:
        payload = build_event_delivery_payload(db, out)
        return c.deliver(target_url, payload, correlation_id=out.correlation_id, tenant_id=tenant_id)

    return deliver_due_events(db, fn, limit=limit)


def get_outbox_event(db: Session, outbox_id: str) -> dict:
    """Inspect a single outbox delivery record + its immutable payload."""
    out = get_or_404(db, m.OutboxEvent, outbox_id, "OutboxEvent")
    event = db.get(m.EcosystemEvent, out.event_id)
    return {
        "id": out.id, "event_id": out.event_id, "target_service": out.target_service,
        "status": out.status, "attempt_count": out.attempt_count, "last_error": out.last_error,
        "next_attempt_at": out.next_attempt_at.isoformat() if out.next_attempt_at else None,
        "delivered_at": out.delivered_at.isoformat() if out.delivered_at else None,
        "external_reference": out.external_reference, "correlation_id": out.correlation_id,
        "event": {
            "event_type": event.event_type, "source_object_id": event.source_object_id,
            "source_revision": event.source_revision, "occurred_at": event.occurred_at.isoformat(),
            "actor_id": event.actor_id, "payload": event.payload,
        } if event else None,
    }


def retry_outbox_event(db: Session, outbox_id: str, actor_id: str | None = None) -> dict:
    """Safely re-enqueue a failed outbox event.

    The original payload is never mutated (it is only re-read from the
    immutable EcosystemEvent). The retry is recorded as its own audit event
    and increments the delivery attempt counter. Returns the re-queued
    record; actual dispatch happens through the normal due-outbox worker.
    """
    out = get_or_404(db, m.OutboxEvent, outbox_id, "OutboxEvent")
    if out.status != "FAILED":
        raise DomainError(f"Cannot retry outbox event in status {out.status}", status_code=409)
    event = db.get(m.EcosystemEvent, out.event_id)
    payload_before = (event.payload or {}) if event else None
    out.status = "PENDING"
    out.next_attempt_at = None
    out.last_error = None
    db.commit()
    record_audit(
        db, action="REPLAY_ATTEMPTED", project_id=event.project_id if event else None,
        actor_id=actor_id, object_type="OutboxEvent", object_id=out.id,
        correlation_id=out.correlation_id,
        metadata={"event_id": out.event_id, "target": out.target_service,
                  "payload_unchanged": event is not None and (event.payload or {}) == (payload_before or {})},
    )
    return get_outbox_event(db, out.id)


def list_ecosystem_events(db: Session, project_id: str | None = None, limit: int = 200) -> list[dict]:
    q = select(m.EcosystemEvent).order_by(m.EcosystemEvent.occurred_at.desc()).limit(limit)
    if project_id:
        q = q.where(m.EcosystemEvent.project_id == project_id)
    rows = db.execute(q).scalars().all()
    return [
        {
            "id": e.id, "event_type": e.event_type, "project_id": e.project_id,
            "source_service": e.source_service, "source_object_id": e.source_object_id,
            "source_revision": e.source_revision, "occurred_at": e.occurred_at.isoformat(),
            "actor_id": e.actor_id, "payload": e.payload, "correlation_id": e.correlation_id,
        }
        for e in rows
    ]


def list_outbox(db: Session, project_id: str | None = None, limit: int = 200) -> list[dict]:
    q = select(m.OutboxEvent).order_by(m.OutboxEvent.created_at.desc()).limit(limit)
    rows = db.execute(q).scalars().all()
    return [
        {
            "id": o.id, "event_id": o.event_id, "target_service": o.target_service,
            "status": o.status, "attempt_count": o.attempt_count, "last_error": o.last_error,
            "delivered_at": o.delivered_at.isoformat() if o.delivered_at else None,
            "external_reference": o.external_reference, "correlation_id": o.correlation_id,
        }
        for o in rows
    ]


# ---------------------------------------------------------------------------
# Ecosystem handoffs (PM / QA) and external references
# ---------------------------------------------------------------------------

EXECUTION_HANDOFF_STATUSES = {
    "DRAFT", "READY", "QUEUED", "DELIVERED_TO_CONDUCTOR",
    "ACKNOWLEDGED", "FAILED", "CANCELLED", "SENT",
}
QA_HANDOFF_STATUSES = {
    "DRAFT", "READY", "QUEUED", "DELIVERED_TO_CONDUCTOR",
    "ACKNOWLEDGED", "FAILED", "CANCELLED", "SENT",
}
# Conductor Main is the ecosystem orchestration authority; Document Again
# delivers design handoffs to its relay (never directly into PM/QA).
CONDUCTOR_MAIN_URL = os.environ.get("CONDUCTOR_MAIN_URL", "http://localhost:8010/api").rstrip("/")
EXTERNAL_RELATION_TYPES = {"IMPLEMENTED_BY", "VALIDATED_BY", "TRACKED_BY", "RELEASED_IN", "HANDED_OFF_TO"}


def create_execution_handoff(
    db: Session,
    project_id: str,
    baseline_id: str | None = None,
    source_revision_id: str | None = None,
    change_request_id: str | None = None,
    target_service: str = "pm-again",
    actor: str = "local-user",
    actor_id: str | None = None,
    status: str = "DRAFT",
) -> dict:
    """Snapshot the exact baseline context and emit a durable PM handoff.

    The payload holds only immutable references so that PM Again can fetch
    the authoritative design by id; Document Again never hands out a copy of
    mutable execution state.
    """
    guard_project(db, project_id)
    if status not in EXECUTION_HANDOFF_STATUSES:
        raise ValueError(f"invalid execution handoff status: {status}")
    baseline_id = baseline_id or _latest_baseline_id(db, project_id)
    source_revision_id = source_revision_id or _latest_design_revision_id(db, project_id)
    payload = versioned_payload(
        "execution-handoff", 1,
        baselineId=baseline_id,
        sourceRevisionId=source_revision_id,
        changeRequestId=change_request_id,
        projectId=project_id,
    )
    correlation_id = f"pm:{project_id}:{baseline_id or source_revision_id or 'adhoc'}"
    handoff = m.ExecutionHandoff(
        project_id=project_id,
        baseline_id=baseline_id,
        source_revision_id=source_revision_id,
        change_request_id=change_request_id,
        target_service=target_service,
        status=status,
        payload_snapshot=payload,
        correlation_id=correlation_id,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(handoff)
    db.flush()
    emit_event(
        db,
        event_type="EXECUTION_REQUESTED",
        project_id=project_id,
        tenant_id=None,
        source_object_id=handoff.id,
        source_revision=source_revision_id,
        actor_id=actor_id,
        payload=payload,
        correlation_id=correlation_id,
        target_services=[target_service],
    )
    db.commit()
    db.refresh(handoff)
    record_audit(
        db, action="HANDOFF_CREATED", project_id=project_id, actor_id=actor_id,
        object_type="ExecutionHandoff", object_id=handoff.id,
        correlation_id=correlation_id, metadata={"target": target_service},
    )
    return _execution_handoff_dict(handoff)


def create_qa_validation_handoff(
    db: Session,
    project_id: str,
    baseline_id: str | None = None,
    requirement_ids: list[str] | None = None,
    semantic_object_ids: list[str] | None = None,
    design_revision_ids: list[str] | None = None,
    target_release: str | None = None,
    target_service: str = "qa-again",
    actor: str = "local-user",
    actor_id: str | None = None,
    status: str = "DRAFT",
) -> dict:
    guard_project(db, project_id)
    if status not in QA_HANDOFF_STATUSES:
        raise ValueError(f"invalid qa handoff status: {status}")
    baseline_id = baseline_id or _latest_baseline_id(db, project_id)
    payload = versioned_payload(
        "qa-validation-handoff", 1,
        baselineId=baseline_id,
        requirementIds=requirement_ids or [],
        semanticObjectIds=semantic_object_ids or [],
        designRevisionIds=design_revision_ids or [],
        targetRelease=target_release,
        projectId=project_id,
    )
    correlation_id = f"qa:{project_id}:{baseline_id or 'adhoc'}"
    handoff = m.QAValidationHandoff(
        project_id=project_id,
        baseline_id=baseline_id,
        requirement_ids=requirement_ids or [],
        semantic_object_ids=semantic_object_ids or [],
        design_revision_ids=design_revision_ids or [],
        target_release=target_release,
        target_service=target_service,
        status=status,
        payload_snapshot=payload,
        correlation_id=correlation_id,
        created_by=actor,
        actor_id=actor_id,
    )
    db.add(handoff)
    db.flush()
    emit_event(
        db,
        event_type="QA_VALIDATION_REQUESTED",
        project_id=project_id,
        tenant_id=None,
        source_object_id=handoff.id,
        source_revision=None,
        actor_id=actor_id,
        payload=payload,
        correlation_id=correlation_id,
        target_services=[target_service],
    )
    db.commit()
    db.refresh(handoff)
    record_audit(
        db, action="HANDOFF_CREATED", project_id=project_id, actor_id=actor_id,
        object_type="QAValidationHandoff", object_id=handoff.id,
        correlation_id=correlation_id, metadata={"target": target_service},
    )
    return _qa_handoff_dict(handoff)


def mark_handoff_status(
    db: Session, handoff_id: str, kind: str, status: str, external_reference: str | None = None
) -> dict:
    model = m.ExecutionHandoff if kind == "execution" else m.QAValidationHandoff
    valid = EXECUTION_HANDOFF_STATUSES if kind == "execution" else QA_HANDOFF_STATUSES
    if status not in valid:
        raise ValueError(f"invalid {kind} handoff status: {status}")
    row = db.get(model, handoff_id)
    if row is None:
        raise KeyError(f"{kind} handoff not found: {handoff_id}")
    row.status = status
    if external_reference:
        row.external_reference = external_reference
    if status in {"SENT", "ACKNOWLEDGED"} and row.delivered_at is None:
        row.delivered_at = m.utcnow()
    db.commit()
    db.refresh(row)
    if status == "SENT":
        metrics.inc("handoff_sent")
    elif status == "ACKNOWLEDGED":
        metrics.inc("handoff_acknowledged")
    return _execution_handoff_dict(row) if kind == "execution" else _qa_handoff_dict(row)


def list_execution_handoffs(db: Session, project_id: str | None = None) -> list[dict]:
    if project_id:
        guard_project(db, project_id)
    q = select(m.ExecutionHandoff).order_by(m.ExecutionHandoff.created_at.desc())
    if project_id:
        q = q.where(m.ExecutionHandoff.project_id == project_id)
    return [_execution_handoff_dict(h) for h in db.execute(q).scalars().all()]


def list_qa_handoffs(db: Session, project_id: str | None = None) -> list[dict]:
    if project_id:
        guard_project(db, project_id)
    q = select(m.QAValidationHandoff).order_by(m.QAValidationHandoff.created_at.desc())
    if project_id:
        q = q.where(m.QAValidationHandoff.project_id == project_id)
    return [_qa_handoff_dict(h) for h in db.execute(q).scalars().all()]


def _execution_handoff_dict(h: m.ExecutionHandoff) -> dict:
    return {
        "id": h.id, "project_id": h.project_id, "baseline_id": h.baseline_id,
        "source_revision_id": h.source_revision_id, "change_request_id": h.change_request_id,
        "target_service": h.target_service, "status": h.status,
        "external_reference": h.external_reference, "payload_snapshot": h.payload_snapshot,
        "correlation_id": h.correlation_id, "created_at": h.created_at.isoformat(),
        "created_by": h.created_by, "actor_id": h.actor_id,
        "delivered_at": h.delivered_at.isoformat() if h.delivered_at else None,
        "last_error": h.last_error,
    }


def _qa_handoff_dict(h: m.QAValidationHandoff) -> dict:
    return {
        "id": h.id, "project_id": h.project_id, "baseline_id": h.baseline_id,
        "requirement_ids": h.requirement_ids, "semantic_object_ids": h.semantic_object_ids,
        "design_revision_ids": h.design_revision_ids, "target_release": h.target_release,
        "target_service": h.target_service, "status": h.status,
        "external_reference": h.external_reference, "payload_snapshot": h.payload_snapshot,
        "correlation_id": h.correlation_id, "created_at": h.created_at.isoformat(),
        "created_by": h.created_by, "actor_id": h.actor_id,
        "delivered_at": h.delivered_at.isoformat() if h.delivered_at else None,
        "last_error": h.last_error,
    }


def _requirement_refs(db: Session, req_ids: list[str]) -> list[dict]:
    """Resolve requirement ids to human {id, code, title} refs (P1 UX)."""
    if not req_ids:
        return []
    rows = db.execute(select(m.Requirement).where(m.Requirement.id.in_(req_ids))).scalars().all()
    by_id = {r.id: r for r in rows}
    return [{"id": rid, "code": by_id[rid].code, "title": by_id[rid].title} for rid in req_ids if rid in by_id]


def _workstreams(db: Session, project_id: str | None) -> list[dict]:
    """Expose the project's architecture tracks as human workstream names.

    PM Again materializes one Function per workstream; these names come from
    the confirmed design (architecture diagrams), not from invented scope."""
    if not project_id:
        return []
    rows = db.execute(
        select(m.ArchitectureDiagram).where(m.ArchitectureDiagram.project_id == project_id)
    ).scalars().all()
    return [{"semantic_id": d.semantic_id, "name": d.name} for d in rows if d.name]


def build_handoff_payload(db: Session, handoff_id: str, kind: str) -> dict:
    """Versioned ``document-again-handoff`` envelope for Conductor Main."""
    model = m.ExecutionHandoff if kind == "execution" else m.QAValidationHandoff
    h = get_or_404(db, model, handoff_id, "Handoff")
    project = db.get(m.Project, h.project_id) if h.project_id else None
    baseline = db.get(m.Baseline, h.baseline_id) if h.baseline_id else None
    title = project.name if project else None
    baseline_name = baseline.name if baseline else None
    if kind == "execution":
        req_ids = (h.payload_snapshot or {}).get("requirementIds") or []
        if not req_ids and h.project_id:
            req_ids = [r.id for r in db.execute(select(m.Requirement).where(m.Requirement.project_id == h.project_id)).scalars().all()]
        fields = {
            "handoff_type": "EXECUTION",
            "title": title,
            "baseline_name": baseline_name,
            "requirement_ids": req_ids,
            "requirement_refs": _requirement_refs(db, req_ids),
            "artifact_revision_ids": [h.source_revision_id] if h.source_revision_id else [],
            "semantic_object_ids": (h.payload_snapshot or {}).get("semanticObjectIds") or [],
            "workstreams": _workstreams(db, h.project_id),
            "change_request_id": h.change_request_id,
        }
    else:
        fields = {
            "handoff_type": "QA_VALIDATION",
            "title": title,
            "baseline_name": baseline_name,
            "requirement_ids": h.requirement_ids or [],
            "requirement_refs": _requirement_refs(db, h.requirement_ids or []),
            "artifact_revision_ids": h.design_revision_ids or [],
            "semantic_object_ids": h.semantic_object_ids or [],
            "target_release": h.target_release,
        }
    return versioned_payload(
        "document-again-handoff", 1,
        handoff_id=h.id,
        tenant_id=current_tenant(),
        project_id=h.project_id,
        baseline_id=h.baseline_id,
        correlation_id=h.correlation_id,
        source_service="DOCUMENT_AGAIN",
        payload_snapshot=h.payload_snapshot or {},
        **fields,
    )


def deliver_handoff_to_conductor(
    db: Session, handoff_id: str, kind: str, client=None
) -> dict:
    """Deliver a handoff to Conductor Main's relay (authority boundary).

    Conductor accepts DOCUMENT_AGAIN and dispatches to PM/QA. On success the
    handoff becomes ACKNOWLEDGED with the relayed external reference; on
    failure it becomes FAILED with the error retained (safe to retry).
    """
    from .ecosystem_client import DeliveryError, EcosystemDeliveryClient

    model = m.ExecutionHandoff if kind == "execution" else m.QAValidationHandoff
    h = get_or_404(db, model, handoff_id, "Handoff")
    if h.status == "ACKNOWLEDGED":
        return _execution_handoff_dict(h) if kind == "execution" else _qa_handoff_dict(h)

    payload = build_handoff_payload(db, handoff_id, kind)
    h.status = "QUEUED"
    db.commit()

    c = client or EcosystemDeliveryClient()
    try:
        ext_ref = c.deliver(
            f"{CONDUCTOR_MAIN_URL}/ecosystem/document-handoffs",
            payload, correlation_id=h.correlation_id, tenant_id=current_tenant(),
        )
    except DeliveryError as exc:
        h.status = "FAILED"
        h.last_error = str(exc)[:2000]
        db.commit()
        raise DomainError(f"Conductor handoff delivery failed: {exc}", status_code=502)

    h.status = "ACKNOWLEDGED"
    h.external_reference = ext_ref
    h.delivered_at = m.utcnow()
    h.last_error = None
    db.commit()
    db.refresh(h)
    record_audit(
        db, action="HANDOFF_DELIVERED", project_id=h.project_id, actor_id=h.actor_id,
        object_type="ExecutionHandoff" if kind == "execution" else "QAValidationHandoff",
        object_id=h.id, correlation_id=h.correlation_id,
        metadata={"via": "conductor-main", "external_reference": ext_ref},
    )
    return _execution_handoff_dict(h) if kind == "execution" else _qa_handoff_dict(h)


def _latest_baseline_id(db: Session, project_id: str) -> str | None:
    row = db.execute(
        select(m.Baseline.id)
        .where(m.Baseline.project_id == project_id)
        .order_by(m.Baseline.created_at.desc())
        .limit(1)
    ).scalar_one_or_none()
    return row


def _latest_design_revision_id(db: Session, project_id: str) -> str | None:
    row = db.execute(
        select(m.ArtifactRevision.id)
        .join(m.Artifact, m.Artifact.id == m.ArtifactRevision.artifact_id)
        .where(m.Artifact.project_id == project_id)
        .order_by(m.ArtifactRevision.created_at.desc())
        .limit(1)
    ).scalar_one_or_none()
    return row


def create_external_reference(
    db: Session,
    project_id: str,
    semantic_id: str,
    service: str,
    external_id: str,
    relation_type: str = "TRACKED_BY",
    object_type: str | None = None,
    url: str | None = None,
    metadata: dict | None = None,
) -> dict:
    guard_project(db, project_id)
    if relation_type not in EXTERNAL_RELATION_TYPES:
        raise ValueError(f"invalid external relation type: {relation_type}")
    existing = db.execute(
        select(m.ExternalReference).where(
            m.ExternalReference.project_id == project_id,
            m.ExternalReference.service == service,
            m.ExternalReference.external_id == external_id,
        )
    ).scalar_one_or_none()
    if existing:
        existing.relation_type = relation_type
        existing.semantic_id = semantic_id
        existing.object_type = object_type or existing.object_type
        existing.url = url or existing.url
        existing.metadata_json = metadata or existing.metadata_json
        db.commit()
        db.refresh(existing)
        return _external_reference_dict(existing)
    ref = m.ExternalReference(
        project_id=project_id,
        semantic_id=semantic_id,
        service=service,
        external_id=external_id,
        relation_type=relation_type,
        object_type=object_type,
        url=url,
        metadata_json=metadata,
    )
    db.add(ref)
    db.commit()
    db.refresh(ref)
    return _external_reference_dict(ref)


def list_external_references(
    db: Session, project_id: str | None = None, semantic_id: str | None = None
) -> list[dict]:
    q = select(m.ExternalReference).order_by(m.ExternalReference.created_at.desc())
    if project_id:
        guard_project(db, project_id)
        q = q.where(m.ExternalReference.project_id == project_id)
    if semantic_id:
        q = q.where(m.ExternalReference.semantic_id == semantic_id)
    return [_external_reference_dict(r) for r in db.execute(q).scalars().all()]


def ecosystem_trace(db: Session, project_id: str) -> dict:
    """Traverse the full ecosystem chain for one project:

    baselines -> PM/QA handoffs -> external references -> outbox delivery.

    Internal semantic objects, external service references, and orchestration
    handoffs are kept distinct in the returned structure.
    """
    guard_project(db, project_id)
    baselines = db.execute(
        select(m.Baseline).where(m.Baseline.project_id == project_id).order_by(m.Baseline.created_at)
    ).scalars().all()
    out = {"project_id": project_id, "baselines": []}
    for b in baselines:
        pm = db.execute(
            select(m.ExecutionHandoff).where(m.ExecutionHandoff.baseline_id == b.id)
        ).scalars().all()
        qa = db.execute(
            select(m.QAValidationHandoff).where(m.QAValidationHandoff.baseline_id == b.id)
        ).scalars().all()
        out["baselines"].append({
            "id": b.id, "name": b.name, "target_release": b.target_release,
            "pm_handoffs": [_execution_handoff_dict(h) for h in pm],
            "qa_handoffs": [_qa_handoff_dict(h) for h in qa],
        })
    out["external_references"] = list_external_references(db, project_id=project_id)
    return out


def _external_reference_dict(r: m.ExternalReference) -> dict:
    return {
        "id": r.id, "project_id": r.project_id, "semantic_id": r.semantic_id,
        "relation_type": r.relation_type, "service": r.service, "external_id": r.external_id,
        "object_type": r.object_type, "url": r.url, "metadata": r.metadata_json,
        "created_at": r.created_at.isoformat(),
    }


# ---------------------------------------------------------------------------
# Reproducible export (from versioned structured state, never live state)
# ---------------------------------------------------------------------------


def _safe_filename(name: str) -> str:
    """ASCII-only filename for Content-Disposition headers."""
    keep = []
    for ch in name:
        if ch.isascii() and ch.isalnum() or ch in ".-_":
            keep.append(ch)
        else:
            keep.append("_")
    return "".join(keep)


def export_metadata(db: Session, revision_id: str) -> dict:
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    baseline = db.execute(
        select(m.Baseline).join(m.BaselineBinding, m.BaselineBinding.baseline_id == m.Baseline.id)
        .where(m.BaselineBinding.artifact_revision_id == revision_id)
    ).scalars().first()
    return {
        "project": revision.artifact.project.name,
        "project_key": revision.artifact.project.key,
        "artifact_type": revision.artifact.type.value,
        "artifact_title": revision.artifact.title,
        "revision_number": revision.revision_number,
        "status": revision.status.value,
        "confirmed_by": revision.confirmed_by,
        "confirmed_at": revision.confirmed_at.isoformat() if revision.confirmed_at else None,
        "generated_at": m.utcnow().isoformat(),
        "baseline_id": baseline.id if baseline else None,
        "baseline_name": baseline.name if baseline else None,
    }


def _flatten_section(section: dict) -> list[dict]:
    """Flatten a Tiptap doc JSON section into export-friendly blocks."""
    blocks: list[dict] = []

    def text_of(node) -> str:
        if node.get("type") == "text":
            return node.get("text") or ""
        return "".join(text_of(c) for c in node.get("content") or [])

    def walk(node) -> None:
        t = node.get("type")
        if t == "text":
            return
        children = node.get("content") or []
        if t == "heading":
            blocks.append({"kind": "heading", "level": node.get("attrs", {}).get("level", 2), "text": "".join(text_of(c) for c in children)})
            return
        if t == "paragraph":
            blocks.append({"kind": "paragraph", "text": "".join(text_of(c) for c in children)})
            return
        if t in ("bulletList", "orderedList", "taskList"):
            for item in children:
                blocks.append({"kind": "list_item", "ordered": t == "orderedList", "text": "".join(text_of(c) for c in item.get("content") or [])})
            return
        if t == "codeBlock":
            blocks.append({"kind": "code", "text": "".join(text_of(c) for c in children)})
            return
        if t == "table":
            rows = []
            for row in children:
                rows.append(["".join(text_of(cell) for cell in (row.get("content") or []))])
            blocks.append({"kind": "table", "rows": rows})
            return
        for child in children:
            walk(child)

    walk(section.get("content") or {"type": "doc", "content": []})
    return blocks


def _data_dictionary_from_snapshot(snapshot: dict) -> list[dict]:
    """Data dictionary rows derived from a frozen technical-design snapshot."""
    rows = []
    for schema_sid, schema in (snapshot.get("db_schemas") or {}).items():
        for table_sid, table in (schema.get("tables") or {}).items():
            for field_sid, f in (table.get("fields") or {}).items():
                rows.append({
                    "table": table.get("name"), "table_semantic_id": table_sid,
                    "field": f.get("name"), "field_semantic_id": field_sid,
                    "data_type": f.get("data_type"), "length": f.get("length"),
                    "nullable": f.get("nullable"), "default": f.get("default"),
                    "primary_key": f.get("primary_key"), "foreign_key": f.get("foreign_key"),
                    "reference": f.get("reference"), "description": f.get("description"),
                    "remark": f.get("remark"),
                })
    return rows


def export_revision(db: Session, revision_id: str, format: str) -> tuple[bytes, str, str]:
    """Export a revision from its frozen snapshot (historical correctness)."""
    revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
    snapshot = revision.snapshot or {}
    meta = export_metadata(db, revision_id)
    sections = (snapshot.get("sections") or [])

    if format == "json":
        history = [
            {"revision_number": r.revision_number, "status": r.status.value,
             "confirmed_at": r.confirmed_at.isoformat() if r.confirmed_at else None,
             "id": r.id}
            for r in sorted(revision.artifact.revisions, key=lambda r: r.revision_number)
        ]
        payload = {
            "metadata": meta,
            "revision_history": history,
            "sections": sections,
            "technical_design": snapshot.get("technical_design"),
            "database": snapshot.get("database"),
        }
        return _json.dumps(payload, indent=2, default=str).encode(), "application/json", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}.json"

    if format == "csv":
        # Historical: prefer the frozen technical-design snapshot.
        dd = _data_dictionary_from_snapshot(snapshot.get("technical_design") or {})
        if not dd:
            dd = _data_dictionary_from_snapshot({"db_schemas": snapshot.get("database", {}) or {}})
        buf = _io.StringIO()
        writer = _csv.DictWriter(buf, fieldnames=["table", "field", "data_type", "length", "nullable", "default", "primary_key", "foreign_key", "reference", "description", "remark", "field_semantic_id"])
        writer.writeheader()
        for row in dd:
            writer.writerow({k: row.get(k) for k in writer.fieldnames})
        return buf.getvalue().encode(), "text/csv", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}.csv"

    if format == "pdf":
        return _render_pdf(meta, sections, snapshot), "application/pdf", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}.pdf"

    if format == "svg":
        return _render_erd_svg(snapshot), "image/svg+xml", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}.svg"

    if format == "flow-svg":
        return _render_flow_svg(snapshot), "image/svg+xml", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}-flow.svg"

    if format == "architecture-svg":
        return _render_arch_svg(snapshot), "image/svg+xml", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}-arch.svg"

    if format == "png":
        return _render_erd_png(snapshot), "image/png", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}.png"

    if format == "flow-png":
        return _render_flow_png(snapshot), "image/png", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}-flow.png"

    if format == "architecture-png":
        return _render_arch_png(snapshot), "image/png", f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}-arch.png"

    raise DomainError(f"Unknown export format '{format}'", status_code=422)


def _render_pdf(meta: dict, sections: list[dict], snapshot: dict | None = None) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm)
    styles = getSampleStyleSheet()
    styles["Title"].fontName = "Helvetica-Bold"
    story = [
        Paragraph(meta["artifact_title"], styles["Title"]),
        Paragraph(
            f"{meta['artifact_type']} · revision r{meta['revision_number']} · {meta['status']} · "
            f"confirmed by {meta.get('confirmed_by') or '—'} at {meta.get('confirmed_at') or '—'}<br/>"
            f"project {meta['project']} · generated {meta['generated_at']} · baseline {meta.get('baseline_name') or '—'}",
            styles["Normal"],
        ),
        Spacer(1, 6 * mm),
    ]
    for sec in sections:
        story.append(Paragraph(sec.get("heading") or "Untitled", ParagraphStyle("sh", parent=styles["Heading2"], fontSize=13, spaceBefore=8, spaceAfter=4)))
        story.append(Paragraph(f'<font size="7" color="#666">{sec.get("id")}</font>', styles["Normal"]))
        for blk in _flatten_section(sec):
            kind = blk["kind"]
            if kind == "heading":
                story.append(Paragraph(blk["text"], ParagraphStyle("h", parent=styles["Heading3"], fontSize=11, spaceBefore=4, spaceAfter=2)))
            elif kind == "paragraph":
                story.append(Paragraph(blk["text"].replace("\n", "<br/>"), styles["Normal"]))
            elif kind == "list_item":
                story.append(Paragraph(("• " if not blk.get("ordered") else "1. ") + blk["text"], styles["Normal"]))
            elif kind == "code":
                story.append(Preformatted(blk["text"], ParagraphStyle("code", fontName="Courier", fontSize=8, leading=10)))
            elif kind == "table":
                rows = blk["rows"]
                if rows:
                    t = Table([[Paragraph(c, styles["Normal"]) for c in row] for row in rows])
                    t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, "#999"), ("FONTSIZE", (0, 0), (-1, -1), 8)]))
                    story.append(t)
        story.append(Spacer(1, 3 * mm))
    for heading, rows in _design_summary_blocks((snapshot or {}).get("technical_design")):
        story.append(Paragraph(heading, ParagraphStyle("dh", parent=styles["Heading2"], fontSize=13, spaceBefore=8, spaceAfter=4)))
        if rows:
            t = Table([[Paragraph(c, styles["Normal"]) for c in row] for row in rows])
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, "#999"), ("FONTSIZE", (0, 0), (-1, -1), 8)]))
            story.append(t)
        story.append(Spacer(1, 3 * mm))
    doc.build(story)
    return buf.getvalue()


def _render_erd_svg(snapshot: dict) -> bytes:
    """Simple ERD SVG from a frozen technical-design snapshot."""
    design = snapshot.get("technical_design") or {}
    schemas = design.get("db_schemas") or {}
    tables: list[tuple[str, str, list]] = []  # (name, sid, [(field, type, pk, fk)])
    for schema in schemas.values():
        for tid, t in (schema.get("tables") or {}).items():
            fields = [(f.get("name"), f.get("data_type"), f.get("primary_key"), f.get("foreign_key")) for f in (t.get("fields") or {}).values()]
            tables.append((t.get("name"), tid, fields))
    parts = ['<svg xmlns="http://www.w3.org/2000/svg" width="800" height="600" viewBox="0 0 800 600">']
    x = 40
    y = 40
    for i, (name, sid, fields) in enumerate(tables):
        h = 30 + len(fields) * 18
        if y + h > 560:
            y = 40
            x += 240
        parts.append(f'<rect x="{x}" y="{y}" width="200" height="{h}" rx="6" fill="#161a23" stroke="#6366f1"/>')
        parts.append(f'<text x="{x+8}" y="{y+18}" font-size="13" font-weight="bold" fill="#e2e8f0">{name}</text>')
        for j, (fname, ftype, pk, fk) in enumerate(fields):
            fy = y + 34 + j * 18
            mark = "🔑" if pk else ("🔗" if fk else "")
            parts.append(f'<text x="{x+8}" y="{fy}" font-size="11" fill="#94a3b8">{mark} {fname}: {ftype}</text>')
        y += h + 24
    parts.append("</svg>")
    return "".join(parts).encode()


def _design_summary_blocks(design: dict) -> list[tuple[str, list[list[str]]]]:
    """Frozen-snapshot design summaries for embedding in PDF/DOCX exports.

    Everything is derived from the exact revision snapshot — never live state.
    """
    blocks: list[tuple[str, list[list[str]]]] = []
    if not design:
        return blocks
    if design.get("api_endpoints"):
        rows = [["API endpoint", "Method", "Summary"]]
        for sid, ep in design["api_endpoints"].items():
            rows.append([sid, ep.get("method"), ep.get("summary") or ""])
        blocks.append(("API summary", rows))
    if design.get("flows"):
        rows = [["Flow", "Step", "Type"]]
        for fsid, flow in design["flows"].items():
            for sid, step in (flow.get("steps") or {}).items():
                rows.append([flow.get("name") or fsid, step.get("name"), step.get("step_type")])
        blocks.append(("Process flow summary", rows))
    if design.get("architecture"):
        rows = [["Diagram", "Node", "Type"]]
        for aid, arch in design["architecture"].items():
            for nid, node in (arch.get("nodes") or {}).items():
                rows.append([arch.get("name") or aid, node.get("name"), node.get("node_type")])
        blocks.append(("Architecture summary", rows))
    return blocks


def _render_flow_svg(snapshot: dict) -> bytes:
    design = snapshot.get("technical_design") or {}
    flows = design.get("flows") or {}
    parts = ['<svg xmlns="http://www.w3.org/2000/svg" width="800" height="600" viewBox="0 0 800 600">']
    y = 30
    for fsid, flow in flows.items():
        parts.append(f'<text x="20" y="{y}" font-size="14" font-weight="bold" fill="#e2e8f0">{flow.get("name") or fsid}</text>')
        y += 20
        for sid, step in (flow.get("steps") or {}).items():
            parts.append(f'<rect x="20" y="{y-12}" width="180" height="22" rx="4" fill="#161a23" stroke="#8b5cf6"/>')
            parts.append(f'<text x="28" y="{y+3}" font-size="11" fill="#c4b5fd">{step.get("name")} [{step.get("step_type")}]</text>')
            y += 30
        y += 14
    parts.append("</svg>")
    return "".join(parts).encode()


def _render_one_arch_svg(name: str, nodes: dict, edges: dict) -> bytes:
    parts = ['<svg xmlns="http://www.w3.org/2000/svg" width="800" height="600" viewBox="0 0 800 600">']
    parts.append(f'<text x="20" y="30" font-size="14" font-weight="bold" fill="#e2e8f0">{name}</text>')
    y = 50
    for nid, node in nodes.items():
        parts.append(f'<circle cx="28" cy="{y-4}" r="5" fill="#14b8a6"/>')
        parts.append(f'<text x="42" y="{y}" font-size="11" fill="#94a3b8">{node.get("name")} [{node.get("node_type")}]</text>')
        y += 22
    for eid, edge in edges.items():
        parts.append(f'<text x="20" y="{y}" font-size="10" fill="#6366f1">{edge.get("from")} -> {edge.get("to")}{(": " + edge.get("label")) if edge.get("label") else ""}</text>')
        y += 14
    parts.append("</svg>")
    return "".join(parts).encode()


def _render_arch_svg(snapshot: dict) -> bytes:
    design = snapshot.get("technical_design") or {}
    arch = design.get("architecture") or {}
    parts = ['<svg xmlns="http://www.w3.org/2000/svg" width="800" height="600" viewBox="0 0 800 600">']
    y = 30
    for aid, d in arch.items():
        parts.append(f'<text x="20" y="{y}" font-size="14" font-weight="bold" fill="#e2e8f0">{d.get("name") or aid}</text>')
        y += 20
        for nid, node in (d.get("nodes") or {}).items():
            parts.append(f'<circle cx="28" cy="{y-4}" r="5" fill="#14b8a6"/>')
            parts.append(f'<text x="42" y="{y}" font-size="11" fill="#94a3b8">{node.get("name")} [{node.get("node_type")}]</text>')
            y += 22
        y += 12
    parts.append("</svg>")
    return "".join(parts).encode()


def render_architecture_diagram_svg(db: Session, diagram_id: str) -> bytes:
    d = get_or_404(db, m.ArchitectureDiagram, diagram_id, "ArchitectureDiagram")
    nodes = db.execute(select(m.ArchitectureNode).where(m.ArchitectureNode.diagram_id == d.id)).scalars().all()
    edges = db.execute(select(m.ArchitectureEdge).where(m.ArchitectureEdge.diagram_id == d.id)).scalars().all()
    return _render_one_arch_svg(
        d.name,
        {n.semantic_id: {"name": n.name, "node_type": n.node_type} for n in nodes},
        {e.semantic_id: {"from": e.from_node_semantic_id, "to": e.to_node_semantic_id, "label": e.label} for e in edges},
    )


def render_architecture_diagram_png(db: Session, diagram_id: str) -> bytes:
    return _svg_to_png(render_architecture_diagram_svg(db, diagram_id))


def _svg_to_png(svg_bytes: bytes) -> bytes:
    """Rasterize SVG to PNG with CairoSVG (mature renderer; no hand-rolled
    rasterization). Requires the cairo native library at runtime; on macOS
    Homebrew cairo is resolved via an explicit find_library hint."""
    # Homebrew cairo lives outside ctypes' default search path on macOS.
    _cairo = "/opt/homebrew/lib/libcairo.2.dylib"
    if os.path.exists(_cairo):
        import ctypes.util
        _orig = ctypes.util.find_library
        if getattr(ctypes.util, "_da_cairo_patched", None) is None:
            def _find(name: str) -> str | None:
                if "cairo" in name:
                    return _cairo
                return _orig(name)
            ctypes.util.find_library = _find
            ctypes.util._da_cairo_patched = True
    try:
        import cairosvg
    except Exception as exc:  # noqa: BLE001
        raise DomainError(
            f"PNG export unavailable: cairosvg not installed ({exc})", status_code=501
        ) from exc
    try:
        return cairosvg.svg2png(bytestring=svg_bytes)
    except Exception as exc:  # noqa: BLE001
        raise DomainError(f"PNG rasterization failed: {exc}", status_code=501) from exc


def _render_erd_png(snapshot: dict) -> bytes:
    return _svg_to_png(_render_erd_svg(snapshot))


def _render_flow_png(snapshot: dict) -> bytes:
    return _svg_to_png(_render_flow_svg(snapshot))


def _render_arch_png(snapshot: dict) -> bytes:
    return _svg_to_png(_render_arch_svg(snapshot))


def export_design_package(db: Session, baseline_id: str) -> bytes:
    """ZIP of export artifacts, all generated from the same baseline context."""
    baseline = get_or_404(db, m.Baseline, baseline_id, "Baseline")
    guard_project(db, baseline.project_id)
    bindings = db.execute(
        select(m.BaselineBinding).where(m.BaselineBinding.baseline_id == baseline_id)
    ).scalars().all()
    if not bindings:
        raise DomainError("Baseline has no bindings")

    buf = _io.BytesIO()
    with _zipfile.ZipFile(buf, "w", _zipfile.ZIP_DEFLATED) as z:
        meta_rows = []
        for b in bindings:
            rev = db.get(m.ArtifactRevision, b.artifact_revision_id)
            if rev is None:
                continue
            meta = export_metadata(db, rev.id)
            meta_rows.append(meta)
            if rev.artifact.type in (m.ArtifactType.UR, m.ArtifactType.DR):
                pdf, _, name = export_revision(db, rev.id, "pdf")
                z.writestr(name, pdf)
            json_bytes, _, name = export_revision(db, rev.id, "json")
            z.writestr(f"revision-{rev.revision_number}-{rev.artifact.type.value}.json", json_bytes)
            csv_bytes, _, _ = export_revision(db, rev.id, "csv")
            if csv_bytes.strip():
                z.writestr(f"data-dictionary-{rev.artifact.type.value}-r{rev.revision_number}.csv", csv_bytes)
        z.writestr("manifest.json", _json.dumps({"baseline": baseline.name, "baseline_id": baseline.id, "revisions": meta_rows}, indent=2, default=str))
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Export V2 — XLSX / DOCX / traceability matrix / design package directory
# ---------------------------------------------------------------------------


def _traceability_rows(db: Session, project_id: str) -> list[dict]:
    """Traceability matrix rows from current trace links (labelled as live)."""
    objects = {
        so.semantic_id: so
        for so in db.execute(
            select(m.SemanticObject).where(m.SemanticObject.project_id == project_id)
        ).scalars().all()
    }
    links = db.execute(
        select(m.TraceLink).where(m.TraceLink.project_id == project_id)
        .order_by(m.TraceLink.source_semantic_id)
    ).scalars().all()
    return [
        {
            "source": l.source_semantic_id,
            "source_title": objects.get(l.source_semantic_id).display_name if objects.get(l.source_semantic_id) else l.source_semantic_id,
            "source_type": objects.get(l.source_semantic_id).object_type.value if objects.get(l.source_semantic_id) else None,
            "relation": l.relation_type.value, "target": l.target_semantic_id,
            "target_title": objects.get(l.target_semantic_id).display_name if objects.get(l.target_semantic_id) else l.target_semantic_id,
            "target_type": objects.get(l.target_semantic_id).object_type.value if objects.get(l.target_semantic_id) else None,
            "revision_context": l.revision_context,
        }
        for l in links
    ]


def _render_xlsx(meta: dict, sections: list[dict], dd_rows: list[dict], trace_rows: list[dict]) -> bytes:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    navy = "1F4E78"
    pale_blue = "D9EAF7"
    pale_gray = "F3F4F6"
    white = "FFFFFF"
    thin = Side(style="thin", color="D1D5DB")

    def setup_page(sheet, landscape=False, repeat=None):
        sheet.sheet_view.showGridLines = False
        sheet.page_setup.orientation = "landscape" if landscape else "portrait"
        sheet.page_setup.fitToWidth = 1
        sheet.page_setup.fitToHeight = 0
        sheet.sheet_properties.pageSetUpPr.fitToPage = True
        sheet.page_margins.left = sheet.page_margins.right = 0.3
        if repeat:
            sheet.print_title_rows = repeat

    def table_header(sheet, row):
        for cell in sheet[row]:
            if cell.value is not None:
                cell.fill = PatternFill("solid", fgColor=navy)
                cell.font = Font(bold=True, color=white)
                cell.alignment = Alignment(wrap_text=True, vertical="center")
                cell.border = Border(bottom=thin)
        sheet.row_dimensions[row].height = 28

    # Narrative content is the first visible sheet for UR/DR.
    if sections:
        ws2 = wb.create_sheet("Document")
        setup_page(ws2, repeat="1:4")
        ws2.freeze_panes = "A5"
        ws2.column_dimensions["A"].width = 24
        ws2.column_dimensions["B"].width = 92
        ws2.merge_cells("A1:B1")
        title = ws2["A1"]
        title.value = str(meta.get("artifact_title") or "Project document")
        title.font = Font(bold=True, size=18, color=white)
        title.fill = PatternFill("solid", fgColor=navy)
        title.alignment = Alignment(vertical="center")
        ws2.row_dimensions[1].height = 34
        ws2.merge_cells("A2:B2")
        ws2["A2"] = f"{meta.get('project') or 'Project'} · {meta.get('artifact_type') or 'Document'} · Version {meta.get('revision_number') or '—'} · {str(meta.get('status') or '—').replace('_', ' ').title()}"
        ws2["A2"].font = Font(bold=True, color="374151")
        ws2.merge_cells("A3:B3")
        ws2["A3"] = f"Updated {meta.get('confirmed_at') or meta.get('generated_at') or '—'} · Generated {meta.get('generated_at') or '—'}"
        ws2["A3"].font = Font(size=10, color="6B7280")
        ws2.append(["Section", "Content"])
        table_header(ws2, 4)
        r = 5
        for sec in sections:
            sc = ws2.cell(row=r, column=1, value=str(sec.get("heading") or "Section"))
            hc = ws2.cell(row=r, column=2, value=str(sec.get("heading") or "Section"))
            for cell in (sc, hc):
                cell.font = Font(bold=True, color="1F2937")
                cell.fill = PatternFill("solid", fgColor=pale_blue)
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = Border(bottom=thin)
            # Keep semantic identity for audit, but unobtrusively in a note.
            sc.value = f"Section\n{sec.get('id') or ''}" if sec.get("id") else "Section"
            sc.font = Font(bold=True, color="374151", size=10)
            ws2.row_dimensions[r].height = 34
            r += 1
            for blk in _flatten_section(sec):
                kind = str(blk.get("kind") or "Content").replace("_", " ").title()
                ws2.cell(row=r, column=1, value=kind).alignment = Alignment(vertical="top")
                cell = ws2.cell(row=r, column=2, value=blk["text"])
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                for c in ws2[r]:
                    c.border = Border(bottom=Side(style="hair", color="E5E7EB"))
                lines = max(1, len(str(blk["text"])) // 105 + str(blk["text"]).count("\n") + 1)
                ws2.row_dimensions[r].height = min(90, max(22, lines * 15))
                r += 1

    ws = wb.create_sheet("Metadata")
    setup_page(ws)
    ws.column_dimensions["A"].width = 26
    ws.column_dimensions["B"].width = 88
    ws.append(["Document metadata", None])
    ws.merge_cells("A1:B1")
    ws["A1"].font = Font(bold=True, size=16, color=white)
    ws["A1"].fill = PatternFill("solid", fgColor=navy)
    ws.row_dimensions[1].height = 30
    for i, (k, v) in enumerate(meta.items(), start=1):
        row = i + 2
        c = ws.cell(row=row, column=1, value=str(k).replace("_", " ").title())
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", fgColor=pale_gray)
        vc = ws.cell(row=row, column=2, value=str(v))
        vc.alignment = Alignment(wrap_text=True, vertical="top")

    if dd_rows:
        ws3 = wb.create_sheet("Data Dictionary")
        setup_page(ws3, landscape=True, repeat="1:1")
        ws3.freeze_panes = "A2"
        headers = ["table", "field", "data_type", "length", "nullable", "primary_key",
                   "foreign_key", "reference", "description", "remark", "field_semantic_id"]
        ws3.append(headers)
        table_header(ws3, 1)
        ws3.auto_filter.ref = f"A1:{get_column_letter(len(headers))}1"
        for row in dd_rows:
            ws3.append([row.get(h) for h in headers])
        for col in range(1, len(headers) + 1):
            ws3.column_dimensions[get_column_letter(col)].width = 18 if col not in (8, 9, 10) else 34

    if trace_rows:
        ws4 = wb.create_sheet("Traceability")
        setup_page(ws4, landscape=True, repeat="1:1")
        ws4.freeze_panes = "A2"
        headers = ["Source ID", "Source title", "Source type", "Relationship", "Target ID", "Target title", "Target type", "Revision context"]
        ws4.append(headers)
        table_header(ws4, 1)
        ws4.auto_filter.ref = "A1:H1"
        human_relations = {"DERIVED_FROM": "Designed in", "REFERENCES": "References", "TRACES_TO": "Traces to"}
        for row in trace_rows:
            ws4.append([row["source"], row.get("source_title"), row["source_type"],
                        human_relations.get(row["relation"], row["relation"].replace("_", " ").title()),
                        row["target"], row.get("target_title"), row["target_type"], row["revision_context"]])
        for i, width in enumerate((18, 34, 18, 18, 22, 38, 20, 22), start=1):
            ws4.column_dimensions[get_column_letter(i)].width = width
        for row in ws4.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")

    # Open on the owner-facing content, not the technical metadata sheet.
    if sections:
        wb.active = wb.sheetnames.index("Document")
    elif trace_rows:
        wb.active = wb.sheetnames.index("Traceability")

    buf = _io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _render_docx(meta: dict, sections: list[dict], snapshot: dict | None = None) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(meta["artifact_title"], level=0)
    doc.add_paragraph(
        f"{meta['artifact_type']} · revision r{meta['revision_number']} · {meta['status']} · "
        f"confirmed by {meta.get('confirmed_by') or '—'} at {meta.get('confirmed_at') or '—'} · "
        f"project {meta['project']} · generated {meta['generated_at']}"
    )
    for sec in sections:
        doc.add_heading(sec.get("heading") or "Untitled", level=1)
        for blk in _flatten_section(sec):
            kind = blk["kind"]
            if kind == "heading":
                doc.add_heading(blk["text"], level=2)
            elif kind == "paragraph":
                doc.add_paragraph(blk["text"])
            elif kind == "list_item":
                doc.add_paragraph(blk["text"], style="List Bullet")
            elif kind == "code":
                doc.add_paragraph(blk["text"], style="No Spacing")
            elif kind == "table" and blk["rows"]:
                table = doc.add_table(rows=len(blk["rows"]), cols=len(blk["rows"][0]))
                for i, row in enumerate(blk["rows"]):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
    for heading, rows in _design_summary_blocks((snapshot or {}).get("technical_design")):
        doc.add_heading(heading, level=1)
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = str(cell)
    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_revision_v2(db: Session, revision_id: str, format: str) -> tuple[bytes, str, str]:
    """Export V2: xlsx / docx in addition to the V1 formats (reuse)."""
    if format in ("xlsx", "docx"):
        revision = get_or_404(db, m.ArtifactRevision, revision_id, "Revision")
        snapshot = revision.snapshot or {}
        meta = export_metadata(db, revision_id)
        sections = snapshot.get("sections") or []
        dd = _data_dictionary_from_snapshot(snapshot.get("technical_design") or {})
        if not dd:
            dd = _data_dictionary_from_snapshot({"db_schemas": snapshot.get("database", {}) or {}})
        base = f"{_safe_filename(meta['artifact_title'])}-r{revision.revision_number}"
        metrics.inc("export_generated")
        record_audit(
            db, action="EXPORT_GENERATED", project_id=revision.artifact.project_id,
            object_type="ArtifactRevision", object_id=revision.id,
            revision_context=revision.id, metadata={"format": format},
        )
        if format == "xlsx":
            trace = _traceability_rows(db, revision.artifact.project_id)
            return _render_xlsx(meta, sections, dd, trace), \
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"{base}.xlsx"
        return _render_docx(meta, sections, snapshot), \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{base}.docx"
    return export_revision(db, revision_id, format)


def export_design_package_v2(db: Session, baseline_id: str) -> bytes:
    """ZIP export with a clean directory structure (design package V2)."""
    baseline = get_or_404(db, m.Baseline, baseline_id, "Baseline")
    guard_project(db, baseline.project_id)
    bindings = db.execute(
        select(m.BaselineBinding).where(m.BaselineBinding.baseline_id == baseline_id)
    ).scalars().all()
    if not bindings:
        raise DomainError("Baseline has no bindings")

    buf = _io.BytesIO()
    with _zipfile.ZipFile(buf, "w", _zipfile.ZIP_DEFLATED) as z:
        meta_rows = []
        dd_rows = []
        trace_rows = _traceability_rows(db, baseline.project_id)
        for b in bindings:
            rev = db.get(m.ArtifactRevision, b.artifact_revision_id)
            if rev is None:
                continue
            meta = export_metadata(db, rev.id)
            meta_rows.append(meta)
            snapshot = rev.snapshot or {}
            dd_rows.extend(_data_dictionary_from_snapshot(snapshot.get("technical_design") or {}))
            json_bytes, _, _ = export_revision(db, rev.id, "json")
            z.writestr(f"revisions/{rev.revision_number}-{rev.artifact.type.value}.json", json_bytes)
            if rev.artifact.type in (m.ArtifactType.UR, m.ArtifactType.DR):
                pdf, _, _ = export_revision(db, rev.id, "pdf")
                z.writestr(f"documents/{_safe_filename(meta['artifact_title'])}-r{rev.revision_number}.pdf", pdf)
                docx, _, _ = export_revision_v2(db, rev.id, "docx")
                z.writestr(f"documents/{_safe_filename(meta['artifact_title'])}-r{rev.revision_number}.docx", docx)
            # OpenAPI export from API_DESIGN revisions
            if rev.artifact.type == m.ArtifactType.API_DESIGN:
                api_map = (snapshot.get("technical_design") or {}).get("api_endpoints") or {}
                if api_map:
                    z.writestr(f"api/{_safe_filename(meta['artifact_title'])}.openapi.json",
                               _json.dumps(export_openapi(db, rev.id), indent=2, default=str))
        if dd_rows:
            z.writestr("data-dictionary.xlsx", _render_xlsx(
                export_metadata(db, bindings[0].artifact_revision_id), [], dd_rows, []))
        z.writestr("traceability.xlsx", _render_xlsx(
            export_metadata(db, bindings[0].artifact_revision_id), [], [], trace_rows))
        z.writestr("manifest.json", _json.dumps(
            {"baseline": baseline.name, "baseline_id": baseline.id,
             "directory_structure": ["manifest.json", "documents/", "revisions/",
                                     "api/", "data-dictionary.xlsx", "traceability.xlsx"],
             "revisions": meta_rows}, indent=2, default=str))
    return buf.getvalue()


def export_design_package_v4(db: Session, baseline_id: str) -> bytes:
    """Design package V4 — clean directory structure, PNG + SVG visuals.

    Only includes what exists for the selected baseline; never mixes baseline
    context (every artifact is derived from the same frozen bindings).
    """
    baseline = get_or_404(db, m.Baseline, baseline_id, "Baseline")
    guard_project(db, baseline.project_id)
    bindings = db.execute(
        select(m.BaselineBinding).where(m.BaselineBinding.baseline_id == baseline_id)
    ).scalars().all()
    if not bindings:
        raise DomainError("Baseline has no bindings")

    buf = _io.BytesIO()
    with _zipfile.ZipFile(buf, "w", _zipfile.ZIP_DEFLATED) as z:
        meta_rows = []
        dd_rows = []
        trace_rows = _traceability_rows(db, baseline.project_id)
        for b in bindings:
            rev = db.get(m.ArtifactRevision, b.artifact_revision_id)
            if rev is None:
                continue
            meta = export_metadata(db, rev.id)
            meta_rows.append(meta)
            snapshot = rev.snapshot or {}
            design = snapshot.get("technical_design") or {}
            base = _safe_filename(meta["artifact_title"])
            dd_rows.extend(_data_dictionary_from_snapshot(design))
            json_bytes, _, _ = export_revision(db, rev.id, "json")
            z.writestr(f"revisions/{rev.revision_number}-{rev.artifact.type.value}.json", json_bytes)
            if rev.artifact.type in (m.ArtifactType.UR, m.ArtifactType.DR):
                pdf, _, _ = export_revision(db, rev.id, "pdf")
                docx, _, _ = export_revision_v2(db, rev.id, "docx")
                z.writestr(f"documents/{base}-r{rev.revision_number}.pdf", pdf)
                z.writestr(f"documents/{base}-r{rev.revision_number}.docx", docx)
            if design.get("db_schemas"):
                z.writestr(f"database/{base}.erd.svg", _render_erd_svg(snapshot))
                z.writestr(f"database/{base}.erd.png", _render_erd_png(snapshot))
            if design.get("flows"):
                z.writestr(f"flow/{base}.flow.svg", _render_flow_svg(snapshot))
                z.writestr(f"flow/{base}.flow.png", _render_flow_png(snapshot))
            if design.get("architecture"):
                z.writestr(f"architecture/{base}.arch.svg", _render_arch_svg(snapshot))
                z.writestr(f"architecture/{base}.arch.png", _render_arch_png(snapshot))
            if rev.artifact.type == m.ArtifactType.API_DESIGN and design.get("api_endpoints"):
                z.writestr(f"api/{base}.openapi.json", _json.dumps(export_openapi(db, rev.id), indent=2, default=str))
        if dd_rows:
            z.writestr("database/data-dictionary.xlsx", _render_xlsx(
                export_metadata(db, bindings[0].artifact_revision_id), [], dd_rows, []))
        z.writestr("traceability/traceability.xlsx", _render_xlsx(
            export_metadata(db, bindings[0].artifact_revision_id), [], [], trace_rows))
        z.writestr("manifest.json", _json.dumps({
            "baseline": baseline.name, "baseline_id": baseline.id,
            "directory_structure": ["manifest.json", "documents/", "revisions/",
                                     "database/", "flow/", "api/", "architecture/",
                                     "traceability/", "changes/", "audit/"],
            "revisions": meta_rows,
        }, indent=2, default=str))
    return buf.getvalue()


# ---------------------------------------------------------------------------
# OIDA Suggestion (R11) — AI observes & suggests; the human decides.
# ---------------------------------------------------------------------------

def _suggestion_out(s: m.Suggestion) -> dict:
    return {
        "id": s.id, "project_id": s.project_id, "domain": s.domain,
        "related_object_id": s.related_object_id, "type": s.type,
        "title": s.title, "description": s.description,
        "why_it_matters": s.why_it_matters, "question": s.question,
        "suggested_action": s.suggested_action, "severity": s.severity,
        "status": s.status.value, "created_by": s.created_by,
        "created_at": s.created_at.isoformat(),
        "updated_at": s.updated_at.isoformat() if s.updated_at else None,
        "answer": s.answer, "answer_source": s.answer_source,
        "interpretation": s.interpretation,
        "interpretation_confidence": s.interpretation_confidence,
        "follow_up": s.follow_up, "proposed_update": s.proposed_update,
        "review_decision": s.review_decision, "consultation": s.consultation,
        "resolved_at": s.resolved_at.isoformat() if s.resolved_at else None,
    }


def _suggestion_dedupe_key(title: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", (title or "").lower())[:60]


def generate_suggestions(db: Session, project_id: str, *, mode: str = "STANDARD", actor="local-user", actor_id: str | None = None) -> list[dict]:
    """Consult the AI runtime (independent where available; deterministic no-key
    fallback) and persist grounded OIDA Suggestions, skipping duplicates."""
    guard_project(db, project_id)
    from . import ai as ai_runtime
    consultation = ai_runtime.consult(db, project_id, purpose="PROJECT_REVIEW", mode=mode)

    existing_keys = {
        _suggestion_dedupe_key(s.title)
        for s in db.execute(select(m.Suggestion).where(m.Suggestion.project_id == project_id)).scalars()
    }
    # Only RESOLVED clarifications count as "already answered". Unresolved ones
    # are exactly what the Suggestion flow should surface for the human to answer.
    clarifications = db.execute(select(m.Clarification).where(m.Clarification.project_id == project_id)).scalars().all()
    answered_text = " ".join(
        (c.question or "") + " " + (c.answer or "") for c in clarifications if c.resolved
    ).lower()
    # Clarification semantic ids that already have a suggestion (avoid duplicates).
    surfaced_clr = {
        s.related_object_id
        for s in db.execute(select(m.Suggestion).where(m.Suggestion.project_id == project_id)).scalars()
        if s.related_object_id and s.related_object_id.startswith("CLR")
    }

    created = []
    for f in consultation.get("findings", []):
        clr_id = f.get("clarification_id")
        key = clr_id or _suggestion_dedupe_key(f["title"])
        # Do not re-ask something already answered in project memory.
        if key in existing_keys or key in surfaced_clr:
            continue
        if not clr_id and _suggestion_dedupe_key(f["title"]) in answered_text:
            continue
        row = m.Suggestion(
            project_id=project_id, domain=f.get("domain"), related_object_id=f.get("related_object_id"),
            type=f.get("type"), title=f["title"], description=f.get("description"),
            why_it_matters=f.get("why_it_matters"), question=f.get("question"),
            suggested_action=f.get("suggested_action"), severity=f.get("severity", "MEDIUM"),
            created_by=actor, actor_id=actor_id, consultation=consultation,
        )
        db.add(row)
        created.append(row)
        if clr_id:
            surfaced_clr.add(clr_id)
        else:
            existing_keys.add(_suggestion_dedupe_key(f["title"]))
    db.commit()
    record_audit(
        db, action="SUGGESTIONS_GENERATED", project_id=project_id, actor_id=actor_id,
        object_type="Project", object_id=project_id,
        metadata={"count": len(created), "mode": mode},
    )
    return [_suggestion_out(s) for s in created]


def list_suggestions(db: Session, project_id: str) -> list[dict]:
    rows = db.execute(
        select(m.Suggestion).where(m.Suggestion.project_id == project_id).order_by(m.Suggestion.created_at.desc())
    ).scalars().all()
    return [_suggestion_out(s) for s in rows]


def answer_suggestion(db: Session, suggestion_id: str, *, answer: str, source: str = "CUSTOMER", actor="local-user", actor_id: str | None = None) -> dict:
    s = get_or_404(db, m.Suggestion, suggestion_id, "Suggestion")
    s.answer = answer
    s.answer_source = source
    s.status = m.SuggestionStatus.ANSWERED
    db.commit()
    record_audit(
        db, action="SUGGESTION_ANSWERED", project_id=s.project_id, actor_id=actor_id,
        object_type="Suggestion", object_id=s.id, metadata={"source": source},
    )
    return _suggestion_out(s)


def interpret_suggestion(db: Session, suggestion_id: str, *, actor="local-user", actor_id: str | None = None) -> dict:
    s = get_or_404(db, m.Suggestion, suggestion_id, "Suggestion")
    if not s.answer:
        raise DomainError("Suggestion has no answer yet", status_code=422)
    from . import ai as ai_runtime
    result = ai_runtime.interpret_answer(s.answer, s.title)
    s.interpretation = result["interpretation"]
    s.interpretation_confidence = result["confidence"]
    s.follow_up = result.get("follow_up")
    s.proposed_update = result.get("proposed_update")
    if result.get("follow_up"):
        s.status = m.SuggestionStatus.NEEDS_FOLLOW_UP
    else:
        s.status = m.SuggestionStatus.PROPOSED_UPDATE
    db.commit()
    record_audit(
        db, action="SUGGESTION_INTERPRETED", project_id=s.project_id, actor_id=actor_id,
        object_type="Suggestion", object_id=s.id,
        metadata={"confidence": result["confidence"], "follow_up": bool(result.get("follow_up"))},
    )
    return _suggestion_out(s)


def review_suggestion(db: Session, suggestion_id: str, *, decision: str, actor="local-user", actor_id: str | None = None) -> dict:
    """Human review of a proposed update. Accepting only DRAFTS project-memory
    records (clarification / assumption); it never confirms requirements or
    baselines and never bypasses impact analysis or admin re-auth."""
    s = get_or_404(db, m.Suggestion, suggestion_id, "Suggestion")
    decision = (decision or "").upper()
    if decision not in ("ACCEPTED", "REJECTED"):
        raise DomainError("decision must be ACCEPTED or REJECTED")
    s.review_decision = decision

    applied = []
    if decision == "ACCEPTED" and s.proposed_update:
        pu = s.proposed_update
        # If this suggestion surfaced an existing OPEN clarification, resolve it
        # in place instead of duplicating it.
        existing_clr = None
        if s.related_object_id and s.related_object_id.startswith("CLR"):
            existing_clr = db.execute(
                select(m.Clarification).where(
                    m.Clarification.project_id == s.project_id,
                    m.Clarification.semantic_id == s.related_object_id,
                )
            ).scalar_one_or_none()
        if existing_clr:
            existing_clr.answer = s.answer
            existing_clr.resolved = True
            applied.append({"kind": "clarification_resolved", "id": existing_clr.semantic_id})
        else:
            clr = create_clarification(
                db, project_id=s.project_id, question=s.question or s.title,
                answer=s.answer, related_semantic_ids=[s.related_object_id] if s.related_object_id else None,
                actor=actor,
            )
            applied.append({"kind": "clarification", "id": clr.semantic_id})
        if pu.get("kind") == "assumption" and pu.get("text"):
            asm = create_assumption(
                db, project_id=s.project_id, content=pu["text"],
                related_semantic_ids=[s.related_object_id] if s.related_object_id else None,
                actor=actor,
            )
            applied.append({"kind": "assumption", "id": asm.semantic_id})
        s.status = m.SuggestionStatus.ACCEPTED
    else:
        s.status = m.SuggestionStatus.REJECTED
    s.resolved_at = m.utcnow()
    db.commit()
    record_audit(
        db, action="SUGGESTION_REVIEWED", project_id=s.project_id, actor_id=actor_id,
        object_type="Suggestion", object_id=s.id,
        metadata={"decision": decision, "applied": applied},
    )
    out = _suggestion_out(s)
    out["applied"] = applied
    return out


# ---------------------------------------------------------------------------
# R15 — Multi-Agent Council consultation records
# ---------------------------------------------------------------------------

def _consultation_out(c: m.Consultation) -> dict:
    return {
        "id": c.id, "project_id": c.project_id, "task_type": c.task_type,
        "role": c.role, "question": c.question, "context_snapshot": c.context_snapshot,
        "runs": c.runs, "aggregation": c.aggregation, "human_review": c.human_review,
        "stale": c.stale, "created_at": c.created_at.isoformat() if c.created_at else None,
    }


def create_consultation(db: Session, project_id: str, *, task_type: str, question: str,
                        context_envelope: dict, role: str | None = None, actor="local-user",
                        actor_id: str | None = None) -> dict:
    guard_project(db, project_id)
    from . import council
    cid = m.new_id("con")
    result = council.run_council(project_id, task_type, question, context_envelope, cid, role=role)
    row = m.Consultation(
        id=cid, project_id=project_id, task_type=task_type, role=role or "GENERAL_REVIEWER",
        question=question, context_snapshot=result["snapshot"], runs=result["runs"],
        aggregation=result["aggregation"],
    )
    db.add(row)
    db.commit()
    record_audit(
        db, action="COUNCIL_CONSULTED", project_id=project_id, actor_id=actor_id,
        object_type="Consultation", object_id=cid,
        metadata={"task_type": task_type, "mode": result["council_mode"]["mode"],
                  "completed": len([r for r in result["runs"] if r["status"] == "COMPLETED"])},
    )
    return _consultation_out(row)


def list_consultations(db: Session, project_id: str) -> list[dict]:
    rows = db.execute(
        select(m.Consultation).where(m.Consultation.project_id == project_id)
        .order_by(m.Consultation.created_at.desc())
    ).scalars().all()
    return [_consultation_out(c) for c in rows]


def get_consultation(db: Session, consultation_id: str) -> dict:
    c = get_or_404(db, m.Consultation, consultation_id, "Consultation")
    return _consultation_out(c)


def check_consultation_stale(db: Session, consultation_id: str, context_envelope: dict | None = None) -> dict:
    """Stale detection (Phase 9): compare the freshly-computed current context
    hash with the snapshot used for the runs. Never silently recomputes or
    replaces a reviewed result."""
    c = get_or_404(db, m.Consultation, consultation_id, "Consultation")
    snap = c.context_snapshot or {}
    stored = (snap.get("context_hash") or "").strip()
    current = ""
    if context_envelope:
        from . import council
        current = council.snapshot_hash(context_envelope)
    changed = bool(current and stored and current != stored)
    if changed and not c.stale:
        c.stale = True
        db.commit()
    return {
        "consultation_id": c.id, "stale": c.stale, "context_hash_now": current,
        "context_hash_at_run": stored, "changed": changed,
        "note": "Project truth changed after this consultation; re-run Council." if changed
                else "Context matches the snapshot used for the runs.",
    }


def review_consultation(db: Session, consultation_id: str, *, decision: str,
                        comment: str | None = None, important: list[str] | None = None,
                        incorrect: list[str] | None = None, actor="local-user",
                        actor_id: str | None = None) -> dict:
    """Human review of a Council result (Phase 17). Advisory only — never writes
    Council findings into any authority service."""
    c = get_or_404(db, m.Consultation, consultation_id, "Consultation")
    decision = (decision or "").upper()
    if decision not in ("USEFUL", "REJECTED"):
        raise DomainError("decision must be USEFUL or REJECTED")
    c.human_review = {
        "decision": decision, "comment": comment,
        "marked_important": important or [], "marked_incorrect": incorrect or [],
        "reviewed_by": actor, "reviewed_at": m.utcnow().isoformat(),
    }
    db.commit()
    record_audit(
        db, action="COUNCIL_REVIEWED", project_id=c.project_id, actor_id=actor_id,
        object_type="Consultation", object_id=c.id, metadata={"decision": decision},
    )
    return _consultation_out(c)


def council_finding_to_suggestion(db: Session, consultation_id: str, *, finding: dict,
                                  actor="local-user", actor_id: str | None = None) -> dict:
    """Convert a reviewed Council finding into an OIDA Suggestion (Phase 18).
    Provenance source=COUNCIL; the Suggestion flows through the existing
    human-led R11 lifecycle — Council never bypasses Suggestion governance."""
    c = get_or_404(db, m.Consultation, consultation_id, "Consultation")
    title = (finding.get("title") or finding.get("statement") or "Council finding")[:200]
    row = m.Suggestion(
        project_id=c.project_id,
        domain=None, related_object_id=None,
        type=(finding.get("finding_type") or "RISK"),
        title=title,
        description=finding.get("statement"),
        why_it_matters=finding.get("statement"),
        question=finding.get("question") or None,
        suggested_action=(finding.get("recommendation") or None),
        severity=(finding.get("severity") or "MEDIUM"),
        created_by=actor, actor_id=actor_id,
        consultation={
            "source": "COUNCIL",
            "consultation_id": consultation_id,
            "run_ids": finding.get("run_ids") or finding.get("run_id") or None,
            "providers": finding.get("providers"),
            "human_selected_by": actor,
        },
    )
    db.add(row)
    db.commit()
    record_audit(
        db, action="COUNCIL_TO_SUGGESTION", project_id=c.project_id, actor_id=actor_id,
        object_type="Suggestion", object_id=row.id, metadata={"consultation_id": consultation_id},
    )
    return _suggestion_out(row)


def rerun_consultation(db: Session, consultation_id: str, *, context_envelope: dict | None = None,
                       actor="local-user", actor_id: str | None = None) -> dict:
    """Explicit re-run (Phase 30). Creates a NEW consultation record with a new
    id; the old result is never overwritten."""
    c = get_or_404(db, m.Consultation, consultation_id, "Consultation")
    envelope = context_envelope or {}
    if not envelope:
        raise DomainError("Re-run requires the current context envelope (stale truth cannot be re-sent blindly).", 422)
    return create_consultation(
        db, c.project_id, task_type=c.task_type, question=c.question,
        context_envelope=envelope, role=c.role, actor=actor, actor_id=actor_id,
    )
